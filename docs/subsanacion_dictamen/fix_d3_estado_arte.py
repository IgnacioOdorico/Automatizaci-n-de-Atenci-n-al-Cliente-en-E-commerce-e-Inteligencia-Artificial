# -*- coding: utf-8 -*-
"""Tercer dictamen — recomendada 1: ampliar el estado del arte e incorporar
literatura arbitrada latinoamericana y argentina.

Los siete antecedentes que se agregan fueron verificados UNO POR UNO contra la
API de Crossref (autores, anio, revista, volumen, numero, paginas y DOI) o
contra el PDF publicado. Ninguna cita se escribio de memoria.

  1. Aguirre Mayorga (2022)  Cuadernos de Administracion 35, 1-22
     DOI 10.11144/Javeriana.cao35.amitd            [Crossref OK]  Colombia
  2. Alderete, Jones y Motta (2017)  Redes 23(45), 63-95
     DOI 10.48160/18517072re45.111                 [Crossref OK]  Argentina
  3. Alderete y Porris (2023)  Ciencias Administrativas (22), e122
     DOI 10.24215/23143738e122                     [Crossref OK]  Argentina
  4. Bravo Maruri et al. (2025)  GADE 5(3), 206-219
     DOI 10.63549/rg.v5i3.705                      [Crossref OK]  Ecuador
  5. Fondevila-Gascon et al. (2024)  Correspondencias & Analisis (19), 47-70
     DOI 10.24265/cian.2024.n19.02                 [Crossref OK]  Peru
  6. Pachas-Santos et al. (2023)  Computacion y Sistemas 27(2), 4119
     DOI 10.13053/cys-27-2-4119                    [Crossref OK]  Mexico
  7. Ramos De Santis (2024)  Retos 14(27), 115-130
     DOI 10.17163/ret.n27.2024.08                  [Crossref OK]  Ecuador

Con esto el estado del arte pasa de 6 a 13 antecedentes revisados.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document
import copy

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def par(pref):
    return d.paragraphs[idx(pref)]


# ============================================================ procedimiento
print('=== procedimiento de busqueda: se declaran las bases regionales ===')
n = replace_in_paragraph(par('Procedimiento de búsqueda.'),
    'Se consultaron ACL Anthology, arXiv, Crossref y los catálogos en línea de ScienceDirect y '
    'Springer Nature.',
    'Se consultaron ACL Anthology, arXiv, Crossref y los catálogos en línea de ScienceDirect y '
    'Springer Nature para la literatura internacional, y SciELO, Redalyc, Dialnet y los '
    'repositorios institucionales de universidades nacionales argentinas —entre ellos el '
    'repositorio del CONICET— para la literatura regional arbitrada.')
print('  bases regionales declaradas: %d' % n)
n = replace_in_paragraph(par('Procedimiento de búsqueda.'),
    'La ventana temporal se fijó entre 2018 y 2026, con la excepción de las obras canónicas ya '
    'incorporadas al marco teórico.',
    'La ventana temporal se fijó entre 2018 y 2026, con dos excepciones declaradas: las obras '
    'canónicas ya incorporadas al marco teórico, y el estudio de Alderete et al. (2017) sobre '
    'adopción de comercio electrónico en pymes de Córdoba, que se incluye por ser el trabajo '
    'empírico de referencia sobre el caso nacional específico que este trabajo sitúa como '
    'contexto. La identificación bibliográfica de cada antecedente —autores, revista, volumen, '
    'páginas y DOI— fue verificada contra el registro de Crossref o contra el texto publicado.')
print('  excepcion de ventana declarada: %d' % n)

# ============================================================ nueva 2.5.4
print()
print('=== nueva Seccion 2.5.4: antecedentes regionales ===')
p_ancla = par('2.5.3 Automatización de procesos en pequeñas')
# el ancla real es el ultimo parrafo de 2.5.3
i = idx('2.5.3 Automatización de procesos en pequeñas')
while not d.paragraphs[i + 1].style.name.startswith('Heading'):
    i += 1
p_ancla = d.paragraphs[i]

bloque = [
    ('2.5.4 Antecedentes latinoamericanos y del caso argentino', 'Heading 3'),
    ('Los antecedentes revisados hasta aquí describen el estado internacional de cada componente, '
     'pero ninguno se sitúa en el contexto que este trabajo declara como problema: el de una '
     'pequeña o mediana empresa de comercio electrónico argentina. Corresponde por lo tanto '
     'revisar la literatura arbitrada de la región, tanto la que caracteriza ese contexto como la '
     'que reporta implementaciones de asistentes conversacionales en él.',
     'First Paragraph'),
    ('Sobre la adopción de comercio electrónico en pymes argentinas, Alderete, Jones y Motta '
     '(2017) estiman un modelo probit ordenado sobre 119 pymes de Córdoba y encuentran que la '
     'preparación digital —tanto objetiva como percibida—, la educación de los empleados, los '
     'beneficios esperados, la calidad de la conexión de banda ancha y el grado de '
     'internacionalización inciden significativamente en la probabilidad de adoptar el canal. '
     'Alderete y Porris (2023) retoman la cuestión sobre 154 pymes de Bahía Blanca y la vinculan '
     'al entorno institucional: las empresas asociadas a la cámara sectorial alcanzan un nivel '
     'medio de adopción de 2,11 frente a 1,46 en las radicadas en el parque industrial. Ambos '
     'trabajos importan aquí por una razón precisa: sostienen con evidencia local el supuesto que '
     'este trabajo asume sobre su destinatario, esto es que la barrera de adopción en el segmento '
     'no es de voluntad sino de preparación técnica y de acompañamiento. Es exactamente el '
     'supuesto que justifica elegir una herramienta de bajo costo, auto-alojable y de '
     'configuración visual.', 'Body Text'),
    ('Sobre la automatización de procesos en la región, Aguirre Mayorga (2022) propone un marco '
     'metodológico de cuatro fases que integra gestión de procesos de negocio con design '
     'thinking, y lo valida sobre un proceso de certificación del sector eléctrico colombiano, '
     'donde documenta reducción de tiempos y digitalización del flujo de trabajo. Su aporte a '
     'este trabajo es metodológico antes que técnico: muestra que en la región la automatización '
     'de un proceso se aborda habitualmente desde el rediseño del proceso y no desde la '
     'instrumentación de su desempeño, que es el ángulo que aquí se adopta.', 'Body Text'),
    ('Sobre asistentes conversacionales con datos regionales, Ramos De Santis (2024) analiza 1.250 '
     'usuarios de chatbots de diez empresas líderes de logística en Colombia, Perú y Ecuador y, '
     'mediante regresión múltiple, identifica cinco predictores de la satisfacción del cliente '
     '—capacidad de resolver el problema, conocimiento del producto, autonomía en la resolución, '
     'corrección gramatical y disposición a recomendar el servicio—, con un coeficiente de '
     'determinación ajustado de 0,7512. Ese resultado es directamente pertinente para la lectura '
     'del Capítulo 5 de este trabajo, y en un sentido incómodo: cuatro de los cinco predictores '
     'que allí explican la satisfacción refieren a la calidad del contenido de la respuesta, que '
     'es precisamente la dimensión que este trabajo no midió. Fondevila-Gascón, Huamanchumo, '
     'Martín-Guart y Gutiérrez-Aragón (2024), publicados en una revista arbitrada peruana aunque '
     'con datos relevados en España —condición que corresponde consignar—, llegan por vía de '
     'encuesta a una conclusión convergente: los clientes están dispuestos a interactuar con '
     'agentes conversacionales, pero la disposición se revierte cuando la calidad de la respuesta '
     'y de la comprensión es pobre.', 'Body Text'),
    ('Sobre implementaciones concretas, Pachas-Santos, Calderón-Vilca y Cárdenas-Mariño (2023) '
     'construyen un chatbot de recomendación de productos con una arquitectura multimodal basada '
     'en ViLBERT, entrenada sobre Flickr30K y evaluada sobre un conjunto de 20.000 productos, con '
     'una exactitud del 80,6 % sobre consultas por imagen y del 75 % sobre consultas por texto. '
     'Bravo Maruri, Ramírez Reina, Orozco Lara y Espinoza Martínez (2025) documentan, mediante '
     'estudio de caso en una empresa tecnológica ecuatoriana, la integración de un chatbot con '
     'inteligencia artificial al sistema de gestión de clientes para automatizar el soporte de '
     'primer nivel. Este último es el antecedente regional más próximo a lo que aquí se hace '
     '—automatizar la atención de primer nivel de una empresa concreta— y por eso conviene marcar '
     'la diferencia de alcance: se trata de un estudio de caso descriptivo, sin baseline '
     'cronometrado ni contraste contra un umbral fijado de antemano.', 'Body Text'),
    ('Del conjunto regional se desprenden dos observaciones que ordenan el resto del capítulo. La '
     'primera es que la literatura latinoamericana sobre chatbots está fuertemente orientada a la '
     'percepción del usuario y a la satisfacción declarada, y mucho menos a la instrumentación '
     'del desempeño: se mide qué opina el cliente del asistente, no cuánto tarda el sistema ni '
     'con qué exactitud clasifica. La segunda es que los trabajos que sí describen '
     'implementaciones lo hacen en clave de estudio de caso descriptivo, sin término de '
     'comparación medido. Esa combinación deja disponible el hueco que este trabajo ocupa, y '
     'también explica por qué su contribución es de método antes que de magnitud.', 'Body Text'),
]
ancla = p_ancla
for texto, estilo in bloque:
    ancla = insert_paragraph_after(ancla, texto, estilo=estilo)
print('  insertados %d parrafos' % len(bloque))

n = replace_everywhere(d, '2.5.4 Vacío identificado y contribución de este trabajo',
                          '2.5.5 Vacío identificado y contribución de este trabajo')
print('  antigua 2.5.4 renumerada a 2.5.5: %d' % n)

# el vacio identificado incorpora lo regional
n = replace_in_paragraph(par('La revisión permite delimitar la contribución con precisión'),
    'y el estudio de caso de fulfillment opera a una escala organizacional que no es la del '
    'segmento aquí considerado.',
    'el estudio de caso de fulfillment opera a una escala organizacional que no es la del '
    'segmento aquí considerado; y la literatura regional, que sí sitúa el problema en el '
    'contexto correcto, mide percepción del usuario o describe implementaciones sin '
    'instrumentar tiempos ni contrastarlos contra un término de comparación medido.')
print('  vacio identificado actualizado: %d' % n)

# ============================================================ Tabla 2.1
print()
print('=== Tabla 2.1: se incorporan los antecedentes regionales ===')
t21 = None
for t in d.tables:
    if t.rows[0].cells[0].text.strip() == 'Antecedente':
        t21 = t
        break
assert t21 is not None
i_este = [i for i, r in enumerate(t21.rows) if r.cells[0].text.strip() == 'Este trabajo'][0]
molde = t21.rows[1]._tr
NUEVAS = [
    ('Alderete et al. (2017) · AR', 'Adopción de e-commerce en pymes', 'Probit ordenado sobre 119 pymes',
     'Determinantes de la adopción del canal', 'Encuesta (Córdoba)'),
    ('Alderete y Porris (2023) · AR', 'Adopción de e-commerce en pymes', 'Análisis descriptivo y ANOVA sobre 154 pymes',
     'Nivel de adopción según vínculo institucional', 'Relevamiento (Bahía Blanca)'),
    ('Aguirre Mayorga (2022) · CO', 'Procesos de negocio', 'BPM integrado con design thinking',
     'Rediseño del proceso y tiempos de ciclo', 'Estudio de caso (sector eléctrico)'),
    ('Ramos De Santis (2024) · CO/PE/EC', 'Chatbots en logística', 'Regresión múltiple sobre 1.250 usuarios',
     'Predictores de la satisfacción del cliente', 'Encuesta a usuarios B2C'),
    ('Fondevila-Gascón et al. (2024) · PE/ES', 'Chatbots en atención al cliente', 'Encuesta cuantitativa',
     'Percepción y actitud del cliente', 'Encuesta (datos de España)'),
    ('Pachas-Santos et al. (2023) · MX/PE', 'Recomendación en e-commerce', 'Modelo multimodal ViLBERT',
     'Exactitud sobre consultas por imagen y texto', 'Corpus públicos'),
    ('Bravo Maruri et al. (2025) · EC', 'Soporte al cliente de nivel 1', 'Chatbot con IA integrado al CRM',
     'Descripción de la implementación', 'Estudio de caso (empresa real)'),
]
for fila in NUEVAS:
    tr = copy.deepcopy(molde)
    t21.rows[i_este]._tr.addprevious(tr)
    j = [k for k, r in enumerate(t21.rows) if r._tr is tr][0]
    for c, txt in enumerate(fila):
        set_cell(t21, j, c, txt)
print('  filas agregadas: %d  (la tabla queda con %d antecedentes)'
      % (len(NUEVAS), len(t21.rows) - 2))

# ============================================================ bibliografia
print()
print('=== bibliografia: 7 entradas, en orden alfabetico ===')
REFS = [
    ('Axelos. (2019).',
     'Aguirre Mayorga, H. S. (2022). Aproximación metodológica para la innovación y '
     'transformación digital de los procesos de negocio. Un caso de estudio. Cuadernos de '
     'Administración, 35, 1–22. https://doi.org/10.11144/Javeriana.cao35.amitd', 'antes'),
    ('Axelos. (2019).',
     'Alderete, M. V., Jones, C., & Motta, J. J. (2017). Los factores organizacionales y del '
     'entorno en la adopción del comercio electrónico en pymes de Córdoba, Argentina. Redes. '
     'Revista de Estudios Sociales de la Ciencia y la Tecnología, 23(45), 63–95. '
     'https://doi.org/10.48160/18517072re45.111', 'antes'),
    ('Axelos. (2019).',
     'Alderete, M. V., & Porris, M. S. (2023). Análisis de la adopción del comercio electrónico '
     'en Pymes y su vínculo con instituciones locales. Ciencias Administrativas, (22), e122. '
     'https://doi.org/10.24215/23143738e122', 'antes'),
    ('Boettiger, C. (2015).',
     'Bravo Maruri, N. R., Ramírez Reina, Á. G., Orozco Lara, F. R., & Espinoza Martínez, M. P. '
     '(2025). Automatización del soporte al cliente mediante un chatbot con IA integrado al CRM: '
     'caso de estudio empresa EDITRATECH. GADE: Revista Científica, 5(3), 206–219. '
     'https://doi.org/10.63549/rg.v5i3.705', 'despues'),
    ('Dumas, M., La Rosa, M.',
     'Fondevila-Gascón, J.-F., Huamanchumo, A., Martín-Guart, R., & Gutiérrez-Aragón, Ó. (2024). '
     'El chatbot como factor de éxito comunicativo, de marketing y empresarial: análisis '
     'empírico. Correspondencias & Análisis, (19), 47–70. '
     'https://doi.org/10.24265/cian.2024.n19.02', 'despues'),
    ('Parikh, S., Tiwari, M.',
     'Pachas-Santos, L. A., Calderón-Vilca, H. D., & Cárdenas-Mariño, F. C. (2023). Chatbot '
     'basado en el aprendizaje profundo para recomendar productos relevantes. Computación y '
     'Sistemas, 27(2), 4119. https://doi.org/10.13053/cys-27-2-4119', 'antes'),
    ('Turban, E., Outland, J.',
     'Ramos De Santis, P. (2024). Satisfacción del cliente en la logística: un análisis de '
     'chatbots en las empresas líderes de Colombia, Perú y Ecuador. Retos. Revista de Ciencias '
     'de la Administración y Economía, 14(27), 115–130. '
     'https://doi.org/10.17163/ret.n27.2024.08', 'antes'),
]
for ancla_pref, texto, donde in REFS:
    p = par(ancla_pref)
    if donde == 'despues':
        insert_paragraph_after(p, texto, estilo=p.style.name)
    else:
        nuevo = insert_paragraph_after(p, texto, estilo=p.style.name)
        # se inserto despues del ancla: hay que intercambiarlos
        p._element.addprevious(nuevo._element)
    print('  %s' % texto.split('(')[0].strip())

d.save(RUTA)
print()
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
P = [p.text.strip() for p in d2.paragraphs]
i = next(k for k, t in enumerate(P) if t.startswith('CAPÍTULO 8'))
j = next(k for k, t in enumerate(P) if t.startswith('CAPÍTULO 9'))
refs = [P[k] for k in range(i + 1, j) if P[k]]
print()
print('=== bibliografia: %d entradas ===' % len(refs))
import re
claves = [re.split(r'\.\s*\(|\,\s*[A-Z]\.', r)[0] for r in refs]
ordenada = claves == sorted(claves, key=lambda s: s.lower())
print('orden alfabetico:', 'OK' if ordenada else 'REVISAR')
if not ordenada:
    for a, b in zip(claves, sorted(claves, key=lambda s: s.lower())):
        if a != b:
            print('   esperado %r  esta %r' % (b, a))
for r in refs:
    if any(r.startswith(x) for x in ('Aguirre', 'Alderete', 'Axelos', 'Boettiger', 'Bravo',
                                     'Brown', 'Dumas', 'Fondevila', 'Grafana', 'OpenAI',
                                     'Pachas', 'Parikh', 'Pypłacz', 'Ramos', 'Turban')):
        print('  ' + r[:104])
