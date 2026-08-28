# -*- coding: utf-8 -*-
"""Tercer dictamen — bloque de ESTRUCTURA (obligatorias 1, 8b y 9).

Lo verificado antes de tocar nada:
  - Hay DOS indices reales. Uno es un campo TOC de Word dentro de un <w:sdt>
    (invisible para python-docx, por eso ninguna auditoria previa lo vio) que
    esta ubicado ANTES de la portada y titulado "Table of Contents", en ingles.
    El otro es la lista manual de parrafos Compact bajo el encabezado "INDICE",
    sin numeros de pagina.
  - El indice manual declara Listado de Figuras / Tablas / Glosario al final,
    pero fisicamente estan antes del Capitulo 1.

Se conserva el campo automatico —tiene numeros de pagina, hipervinculos y se
actualiza solo— y se da de baja el manual. Con eso la inconsistencia de orden
deja de ser posible: el indice pasa a derivarse de la estructura real.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document
from docx.oxml.ns import qn

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
body = d.element.body


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def par(pref):
    return d.paragraphs[idx(pref)]


def set_texto(p, texto):
    if not p.runs:
        p.add_run(texto)
        return
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ''


print('=' * 74)
print(' OBS 1 — abstract, keywords y declaracion de originalidad')
print('=' * 74)

# --- declaracion de originalidad, antes del resumen -------------------------
p_portada_fin = par('Mendoza – Argentina, 2026')
bloque = [
    ('DECLARACIÓN DE ORIGINALIDAD', 'Heading 1'),
    ('Quienes suscriben declaran que el presente Trabajo Integrador es de su autoría, que fue '
     'elaborado íntegramente para la Tecnicatura Universitaria en Programación de la Universidad '
     'Tecnológica Nacional, Facultad Regional Mendoza, y que no ha sido presentado con '
     'anterioridad para la obtención de ningún otro título o grado académico.', 'First Paragraph'),
    ('Todas las fuentes consultadas se encuentran debidamente citadas en el cuerpo del texto y '
     'consignadas en el Capítulo 8. Los datos experimentales que se reportan en el Capítulo 5 '
     'fueron obtenidos por los autores sobre el sistema descripto en el Capítulo 4; los guiones '
     'de ejecución, los archivos de resultados crudos y los procedimientos de cálculo se '
     'versionan en el repositorio del trabajo y se transcriben en los Anexos I y J, de modo que '
     'toda cifra publicada pueda ser recalculada de forma independiente.', 'Body Text'),
    ('Se deja constancia del uso de herramientas de asistencia basadas en modelos de lenguaje '
     'durante la redacción y la revisión del documento y durante el desarrollo del software. Su '
     'empleo se limitó a tareas de asistencia: el diseño experimental, la ejecución de las '
     'mediciones, la interpretación de los resultados y las conclusiones son responsabilidad '
     'exclusiva de los autores, que asumen la autoría plena del trabajo.', 'Body Text'),
    ('Santiago Sordi — Ignacio Odorico — Juan Cruz Ana', 'Body Text'),
    ('Mendoza, 2026', 'Body Text'),
]
ancla = p_portada_fin
for texto, estilo in bloque:
    ancla = insert_paragraph_after(ancla, texto, estilo=estilo)
print('declaracion de originalidad ..... insertada (%d parrafos)' % len(bloque))

# --- abstract y keywords, despues del resumen -------------------------------
p_kw = par('Palabras clave:')
bloque_en = [
    ('ABSTRACT', 'Heading 1'),
    ('This work addresses the manual handling of the post-sale cycle in Argentine e-commerce '
     'SMEs, where craft-based management produces delays, errors and fragmented customer '
     'service. The research question is stated as follows: to what extent does an automation '
     'pipeline orchestrated with n8n reduce the operational times of the post-sale cycle of a '
     'simulated e-commerce store, measured through the MTTD, MTTR and TMR metrics?',
     'First Paragraph'),
    ('A solution was designed and implemented as two workflows orchestrated with n8n and '
     'deployed in Docker containers. The first is an order-processing pipeline (15 nodes) that '
     'receives orders through a webhook, checks stock in PostgreSQL, updates order states and '
     'notifies the customer by email. The second is a multichannel chatbot implemented and '
     'validated over two channels —a simulated channel that adopts the WhatsApp Cloud API '
     'message format and delivers through a local SMTP capture server, and Telegram, with real '
     'end-to-end delivery— using GPT-4o-mini to classify intents and generate automated '
     'replies. The infrastructure comprises PostgreSQL 15 as the database, Mailpit as the test '
     'SMTP server and Grafana for operational dashboards.', 'Body Text'),
    ('The results show a mean MTTD of 0.009 seconds and a mean MTTR of 0.054 seconds, for a '
     'total end-to-end time of 0.063 seconds over 50 measured orders. The marginal cost of '
     'processing one order in the pipeline is therefore about 780 times lower than that of the '
     'manual process, timed by the authors over ten orders at 49.13 seconds per order (95 % CI: '
     '43.30 s to 54.95 s; Fieller 95 % CI of the ratio: 686× to 875×). The two series are '
     'completely separated, and the contrast is conclusive (Mann-Whitney U = 0, p = 2.7 × 10⁻¹¹; '
     'Welch t = 19.05, df = 9.0, p = 1.4 × 10⁻⁸). The factor compares marginal processing costs '
     'and not customer-perceived latencies, and must be read as an upper bound, since the '
     'automated term comes from a laboratory environment. The chatbot achieved a mean response '
     'time of 1.47 seconds over a corpus of 150 interactions and 3.07 seconds over 45 '
     'interactions on the real Telegram channel, with continuous availability, and an intent '
     'classification accuracy of 92.7 % (Wilson 95 % CI: 87.3 % to 95.9 %) in a zero-shot '
     'regime. These values confirm the hypothesis that automation substantially reduces the '
     'operational times of the post-sale cycle with respect to the manual reference process.',
     'Body Text'),
    ('Keywords: business process automation, e-commerce, post-sale service, n8n, chatbot, large '
     'language models, Docker, MTTD, MTTR, mean response time.', 'Body Text'),
]
ancla = p_kw
for texto, estilo in bloque_en:
    ancla = insert_paragraph_after(ancla, texto, estilo=estilo)
print('abstract + keywords ............. insertado (%d parrafos)' % len(bloque_en))
p_fin_abstract = ancla

print()
print('=' * 74)
print(' OBS 9 — indice duplicado, portada en el arbol, numeracion')
print('=' * 74)

# --- el campo TOC: se retitula y se reubica ---------------------------------
sdt = [c for c in body.iterchildren() if c.tag == qn('w:sdt')]
assert len(sdt) == 1, 'esperaba un unico <w:sdt> (el campo TOC), encontre %d' % len(sdt)
sdt = sdt[0]
contenido = sdt.find(qn('w:sdtContent'))
titulo_toc = list(contenido)[0]
from docx.text.paragraph import Paragraph
set_texto(Paragraph(titulo_toc, d), 'ÍNDICE')
print('titulo del campo TOC ............ "Table of Contents" -> "ÍNDICE"')

# --- baja del indice manual --------------------------------------------------
i_ind = idx('ÍNDICE')
victimas = [d.paragraphs[i_ind]._element]
for p in d.paragraphs[i_ind + 1:]:
    if p.style.name != 'Compact' and p.text.strip():
        break
    if p.style.name.startswith('Heading'):
        break
    victimas.append(p._element)
# se corta en el ultimo Compact real (no arrastra el parrafo vacio siguiente)
while victimas and not ''.join(victimas[-1].itertext()).strip():
    victimas.pop()
for el in victimas:
    el.getparent().remove(el)
print('indice manual ................... %d parrafos eliminados' % len(victimas))

# el campo TOC pasa a ocupar ese lugar: despues del abstract
sdt.getparent().remove(sdt)
p_fin_abstract._element.addnext(sdt)
print('campo TOC ....................... reubicado despues del ABSTRACT')

# --- la portada sale del arbol de contenidos ---------------------------------
p_port = par('PORTADA')
p_port.style = d.styles['Body Text']
for r in p_port.runs:
    r.bold = True
print('encabezado PORTADA .............. Heading 1 -> Body Text (fuera del TOC)')

# --- numeracion de encabezados ----------------------------------------------
for viejo, nuevo in [
    ('Servicios Docker', '4.1.1 Servicios Docker'),
    ('Paneles del dashboard', '4.5.1 Paneles de los dashboards'),
    ('Pruebas funcionales', '5.1.1 Pruebas funcionales'),
    ('Métricas de tiempo', '5.1.2 Métricas de tiempo'),
    ('Pruebas de carga y de concurrencia', '5.1.3 Pruebas de carga y de concurrencia'),
]:
    hecho = 0
    for p in d.paragraphs:
        if p.style.name == 'Heading 3' and p.text.strip() == viejo:
            set_texto(p, nuevo)
            hecho += 1
    print('encabezado sin numerar .......... "%s" -> "%s"  (%d)' % (viejo, nuevo, hecho))

# --- Figuras A1 y A2: son del cuerpo, pasan a 10 y 11 ------------------------
for viejo, nuevo in [('Figura A1:', 'Figura 10:'), ('Figura A2:', 'Figura 11:'),
                     ('Figura A1', 'Figura 10'), ('Figura A2', 'Figura 11')]:
    n = replace_everywhere(d, viejo, nuevo)
    if n:
        print('figuras de anexo ................ "%s" -> "%s"  (%d)' % (viejo, nuevo, n))
n = replace_everywhere(d, 'constituyen la hoja de ruta técnica para una implementación productiva '
                          '(ver Figura 10) (ver Figura 11).',
                          'constituyen la hoja de ruta técnica para una implementación productiva '
                          '(Figuras 10 y 11).')
print('doble parentesis en 7.1 ......... %d' % n)

# --- el trabajo se nombra de una sola manera ---------------------------------
n = replace_everywhere(d, 'durante el desarrollo de esta tesis',
                          'durante el desarrollo de este trabajo integrador')
print('autodenominacion ................ %d' % n)

# --- Word refresca el indice al abrir ----------------------------------------
sett = d.settings.element
uf = sett.find(qn('w:updateFields'))
if uf is None:
    uf = sett.makeelement(qn('w:updateFields'), {qn('w:val'): 'true'})
    sett.append(uf)
else:
    uf.set(qn('w:val'), 'true')
print('settings.xml .................... updateFields=true (Word recalcula el TOC al abrir)')

d.save(RUTA)
print()
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
body2 = d2.element.body
print()
print('=== frontispicio resultante ===')
from docx.table import Table
pi = 0
for k, ch in enumerate(body2.iterchildren()):
    tag = ch.tag.split('}')[1]
    if tag == 'p':
        p = Paragraph(ch, d2)
        if p.text.strip() or p.style.name.startswith('Heading'):
            print('  %-14s %s' % (p.style.name, p.text.strip()[:78]))
        pi += 1
    elif tag == 'sdt':
        print('  %-14s <<< CAMPO TOC AUTOMATICO (%d entradas) >>>'
              % ('sdt', len(ch.find(qn('w:sdtContent'))) - 1))
    elif tag == 'tbl':
        print('  %-14s [tabla]' % 'tbl')
    if k > 40:
        break
print()
print('parrafos: %d  |  tablas: %d' % (len(d2.paragraphs), len(d2.tables)))
print('sdt restantes: %d' % len([c for c in body2.iterchildren() if c.tag == qn('w:sdt')]))
print('parrafos Compact restantes: %d' % len([p for p in d2.paragraphs if p.style.name == 'Compact']))
TXT = '\n'.join(p.text for p in d2.paragraphs) + '\n' + '\n'.join(
    c.text for t in d2.tables for r in t.rows for c in r.cells)
for pat in ('Figura A1', 'Figura A2', 'Table of Contents'):
    print('  "%s": %d %s' % (pat, TXT.count(pat), '' if TXT.count(pat) == 0 else '<-- REVISAR'))
print('  headings sin numerar restantes:',
      [p.text.strip() for p in d2.paragraphs
       if p.style.name == 'Heading 3' and not p.text.strip()[:1].isdigit()])
