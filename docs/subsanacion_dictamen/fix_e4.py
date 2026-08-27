# -*- coding: utf-8 -*-
"""BLOQUEANTE #1 — el baseline manual entra al documento.

Datos: experiments/E4/resultados/e4_tiempos.csv (n=10 válidas de 12; 2 descartadas).
También cierra #17 (desvíos e IC para MTTD/MTTR) y completa #3 (throughput bajo concurrencia).
"""
import sys
import copy
sys.path.insert(0, '.')
from docxkit import *
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)

# ============================================================
#  1. Renumerar 5.4..5.9 -> 5.5..5.10 (descendente, sin colisión)
# ============================================================
for n in (9, 8, 7, 6, 5, 4):
    k = replace_everywhere(d, 'Tabla 5.%d' % n, 'Tabla 5.%d' % (n + 1))
    print('Tabla 5.%d -> 5.%d : %d' % (n, n + 1, k))

# ============================================================
#  2. #17 — Tabla 5.2 con desvío e intervalo de confianza
# ============================================================
t52 = d.tables[17]
set_cell(t52, 0, 1, 'Valor sobre las 50 órdenes medidas')
set_cell(t52, 1, 1, 'media 0,009 s · desvío 0,003 s · mediana 0,009 s · IC 95 % [0,008; 0,009] s')
set_cell(t52, 2, 1, 'media 0,054 s · desvío 0,003 s · mediana 0,054 s · IC 95 % [0,053; 0,055] s')
set_cell(t52, 3, 1, 'media 0,063 s · desvío 0,005 s · mediana 0,062 s · IC 95 % [0,061; 0,064] s')

d.save(RUTA)

# ============================================================
#  3. Insertar §5.1.4 (después de P299) — de mayor a menor índice
# ============================================================
d = Document(RUTA)

BLOQUE_51 = [
    ('El segundo resultado matiza al primero y corresponde declararlo con la misma claridad: '
     'bajo concurrencia el pipeline no perdió integridad, pero sí perdió capacidad de proceso. '
     'De las 120 órdenes generadas, 49 (el 40,8 %) quedaron registradas sin marca de '
     'processed_at, es decir que el motor de flujos aceptó la solicitud pero no completó su '
     'ejecución. Ninguna quedó en estado de error ni produjo un descuento de stock incorrecto: '
     'simplemente no avanzaron. El comportamiento observado es, en consecuencia, seguro pero no '
     'elástico. Las métricas también se degradan de forma apreciable respecto de la corrida '
     'secuencial: el MTTD medio pasa de 0,009 s a 0,092 s —un factor de diez— y el tiempo '
     'end-to-end de 0,063 s a 0,192 s, valores que de todos modos se mantienen tres órdenes de '
     'magnitud por debajo del criterio operativo de 30 segundos. Esta limitación de capacidad '
     'motiva las recomendaciones de encolado y control de admisión del Capítulo 7.', 'Body Text'),

    ('5.1.4 Baseline de atención manual', 'Heading 3'),

    ('Para dotar de un término de comparación medido, y no estimado, a la afirmación central del '
     'trabajo, se cronometró el procesamiento manual de una orden siguiendo el protocolo '
     'descripto en la Sección 3.5.5. La Tabla 5.4 presenta los descriptivos por fase y del tiempo '
     'total sobre las diez órdenes válidas.', 'First Paragraph'),

    ('Tabla 5.4: Descriptivos del tiempo de procesamiento manual, por fase y total.', 'Body Text'),

    ('El tiempo medio de procesamiento manual de una orden resultó de 49,13 s, con un intervalo '
     'de confianza al 95 % de [43,30 s; 54,95 s] calculado por la distribución t de Student con '
     'nueve grados de libertad. El coeficiente de variación fue del 16,6 %, dispersión baja que '
     'indica que el procedimiento se ejecutó de manera estable a lo largo de la serie.',
     'Body Text'),

    ('El reparto del esfuerzo entre fases es en sí mismo un hallazgo. La redacción de la '
     'notificación al cliente concentra el 73,8 % del tiempo total (36,27 s de 49,13 s), mientras '
     'que la lectura y verificación de stock insume el 14,3 % (7,05 s) y la actualización de la '
     'base el 11,8 % restante (5,80 s). Es decir que el grueso del costo del procesamiento manual '
     'no está en la consulta ni en la escritura de datos, que son las operaciones que un operador '
     'humano ejecuta con rapidez, sino en la comunicación con el cliente. Desagregado por rama, '
     'el tiempo fue de 50,96 s en las siete órdenes con stock disponible y de 44,84 s en las tres '
     'sin stock, diferencia atribuible a que la plantilla de confirmación exige cargar el importe '
     'total y la de rechazo no.', 'Body Text'),

    ('Corresponden dos advertencias sobre la interpretación de este valor. La primera surge de un '
     'control previsto en el protocolo: la correlación de Pearson entre el orden de ejecución y '
     'el tiempo empleado fue de r = −0,693, magnitud que indica que el efecto de aprendizaje '
     'siguió actuando dentro de la serie válida. En consecuencia, la media de 49,13 s subestima '
     'el tiempo que emplearía un operador no entrenado, y el baseline debe leerse como un piso y '
     'no como un valor central. La segunda es que la medición se realizó sobre un único operador, '
     'lo que impide estimar la variabilidad entre personas; se declara como limitación en la '
     'Sección 5.4.1.', 'Body Text'),

    ('Por último, este valor no incluye la latencia de detección, esto es, el tiempo que una orden '
     'permanece a la espera antes de que un operador advierta su llegada. Esa magnitud no es '
     'observable sin convertirla en un parámetro elegido por los autores, y por lo tanto no se '
     'suma al número medido. A título ilustrativo puede señalarse que, si un operador revisara la '
     'bandeja cada quince minutos, la espera media esperable sería de 7,5 minutos, valor que por '
     'sí solo excede en dos órdenes de magnitud al tiempo de procesamiento cronometrado. Se '
     'consigna como escenario declarado y no como medición.', 'Body Text'),
]

creados = insertar_bloque(d, 299, BLOQUE_51)
idx_caption = None
for i, p in enumerate(d.paragraphs):
    if p.text.strip().startswith('Tabla 5.4: Descriptivos'):
        idx_caption = i
        break
print('caption de la Tabla 5.4 en P%s' % idx_caption)

# --- crear la tabla y moverla debajo del epígrafe ---
FILAS_T54 = [
    ('Fase', 'n', 'Media', 'Mediana', 'Desvío', 'Mín.', 'Máx.'),
    ('T1 — Lectura y verificación', '10', '7,05 s', '7,12 s', '2,00 s', '4,37 s', '11,39 s'),
    ('T2 — Actualización en la base', '10', '5,80 s', '5,27 s', '1,67 s', '4,45 s', '9,09 s'),
    ('T3 — Redacción de la notificación', '10', '36,27 s', '34,91 s', '6,55 s', '30,45 s', '53,67 s'),
    ('Total por orden', '10', '49,13 s', '46,37 s', '8,15 s', '42,75 s', '69,80 s'),
]
nueva = d.add_table(rows=len(FILAS_T54), cols=7)
try:
    nueva.style = d.tables[18].style
except Exception:
    nueva.style = 'Table Grid'
for ri, fila in enumerate(FILAS_T54):
    for ci, val in enumerate(fila):
        set_cell(nueva, ri, ci, val)
d.paragraphs[idx_caption]._element.addnext(nueva._element)

d.save(RUTA)
print('§5.1.4 insertada (%d párrafos + 1 tabla)' % len(creados))

# ============================================================
#  4. Verificación
# ============================================================
d2 = Document(RUTA)
for i in range(296, 312):
    print('P%-4d [%-16s] %s' % (i, d2.paragraphs[i].style.name, d2.paragraphs[i].text.strip()[:120]))
print()
for ti, t in enumerate(d2.tables):
    if t.rows[0].cells[0].text.strip() == 'Fase':
        print('Tabla nueva en índice %d:' % ti)
        for r in t.rows:
            print('   ' + ' | '.join(c.text.strip() for c in r.cells))
