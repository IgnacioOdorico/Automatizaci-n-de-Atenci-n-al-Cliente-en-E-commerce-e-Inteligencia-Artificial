# -*- coding: utf-8 -*-
"""El 780× se reencuadra como razón de COSTO MARGINAL, no de latencia percibida.
Y se anexa la Tabla I.2 con los registros de la base."""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def reescribir(i, texto):
    p = d.paragraphs[i]
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ''


# --- Resumen ---
reescribir(idx('Los resultados obtenidos muestran un MTTD promedio'),
    'Los resultados obtenidos muestran un MTTD promedio de 0,009 segundos y un MTTR promedio de '
    '0,054 segundos, para un tiempo total end-to-end de 0,063 segundos sobre 50 órdenes medidas. '
    'El costo marginal de procesar una orden en el pipeline resulta así unas 780 veces menor que '
    'el del proceso manual, cronometrado por el propio equipo sobre diez órdenes en 49,13 segundos '
    'por orden (IC 95 %: 43,3 s a 55,0 s); el factor compara costos marginales de procesamiento y '
    'no latencias percibidas por el cliente, y debe leerse como una cota superior por provenir el '
    'término automatizado de un entorno de laboratorio. El TMR promedio del chatbot fue de 1,47 '
    'segundos sobre un corpus de 150 interacciones y de 3,07 segundos sobre 45 interacciones del '
    'canal Telegram real, con disponibilidad continua, y la precisión de clasificación alcanzó el '
    '92,7 % (IC 95 % de Wilson: 87,3 % a 95,9 %). Estos valores permiten confirmar la hipótesis de '
    'que la automatización reduce significativamente los tiempos operativos del ciclo post-venta '
    'respecto del proceso manual de referencia.')

# --- §5.3, contrastación de H1 ---
reescribir(idx('• H1 se confirma en sus dos términos'),
    '• H1 se confirma en sus dos términos. En el término comparativo, que es el sustantivo: el '
    'costo marginal de procesar una orden en el pipeline (0,063 s; n = 50) resulta unas 780 veces '
    'menor que el del procesamiento manual, medido en 49,13 s por orden (IC 95 % del factor: 687× '
    'a 872×). Ambos términos miden la misma magnitud —el tiempo que insume una orden adicional— '
    'conforme se precisa en la Sección 3.5.5. En el criterio operativo secundario: ese mismo '
    'tiempo representa el 0,2 % del umbral de 30 segundos. La Sección 5.4 precisa por qué el '
    'factor debe leerse como una cota superior.')

# --- §5.4, discusión ---
reescribir(idx('Los resultados confirman la premisa central del trabajo'),
    'Los resultados confirman la premisa central del trabajo: la automatización reduce '
    'significativamente los tiempos operativos del ciclo post-venta. La reducción del costo '
    'marginal de procesamiento de una orden es de un factor aproximado de 780 veces respecto del '
    'baseline manual medido en la Sección 5.1.4 (49,13 s → 0,063 s). Corresponde precisar el '
    'alcance de esa comparación en tres puntos. Primero, ambos términos miden la misma magnitud: '
    'el tiempo que insume procesar una orden adicional. No se comparan contra el tiempo de espera '
    'observado en la cola manual (313,37 s de media), porque esa cifra depende del tamaño del lote '
    'y no del proceso, y utilizarla habría inflado el factor hasta cerca de 5.000×. Segundo, el '
    'factor debe leerse como una cota superior y no como el valor esperable en un despliegue '
    'productivo: el término automatizado proviene de un entorno de laboratorio en el que la '
    'notificación se entrega a un capturador SMTP alojado en el mismo host, sin tránsito de correo '
    'real ni latencia de servicios de terceros, mientras que el término manual se midió sobre un '
    'operador que ya conocía el procedimiento y trabajaba con una plantilla fija, dos decisiones '
    'deliberadamente conservadoras que reducen el numerador. Tercero, el alcance instrumentado no '
    'es idéntico: los 0,063 s cubren el intervalo received_at → notified_at, es decir escrituras '
    'sobre la base local, mientras que el baseline manual cubre la tarea humana completa de '
    'lectura, verificación, actualización y redacción. En atención al cliente, el chatbot responde '
    'en 1,47 s en promedio sobre el corpus evaluado y en 3,07 s sobre el canal Telegram real, con '
    'disponibilidad continua.')

# --- §6.1 ---
reescribir(idx('Los resultados demuestran que la automatización reduce los tiempos'),
    'Los resultados demuestran que la automatización reduce los tiempos de forma sustancial: el '
    'costo marginal de procesar una orden en el pipeline (0,063 s end-to-end; n = 50) es '
    'aproximadamente 780 veces menor que el del procesamiento manual, cronometrado en 49,13 s por '
    'orden sobre diez órdenes (Sección 3.5.5) y validado de forma independiente contra las marcas '
    'temporales de la base, que arrojan 51,28 s para la misma magnitud. La Sección 5.4 acota '
    'explícitamente ese factor como cota superior. El TMR del chatbot es de 1,47 s en promedio '
    'sobre el corpus evaluado, con disponibilidad continua. Las tres hipótesis de trabajo (H1, H2a '
    'y H2b) fueron confirmadas por los datos experimentales (ver Sección 5.3).')

d.save(RUTA)

# ============================================================
#  Tabla I.2 — los registros de la base
# ============================================================
d = Document(RUTA)
i = idx('El texto completo de las notificaciones redactadas por el operador')
insertar_bloque(d, i, [
    ('Se transcriben además los registros que el operador escribió en la base durante la medición. '
     'Las doce órdenes comparten la misma marca de recepción porque fueron creadas por una única '
     'sentencia de preparación; en consecuencia la columna de tiempo acumulado mide permanencia en '
     'la cola y no latencia individual, mientras que la columna de incremento reproduce el tiempo '
     'marginal por otra vía instrumental (Sección 5.1.4).', 'Body Text'),
    ('Tabla I.2: Registros en la base de datos de las órdenes procesadas manualmente.', 'Body Text'),
])

FILAS = [
    ('Orden', 'Estado', 'processed_at', 'notified_at', 'Acumulado', 'Incremento'),
    ('ORD-E4-001', 'pending', '—', '—', '—', '— (descartada)'),
    ('ORD-E4-002', 'pending', '—', '—', '—', '— (descartada)'),
    ('ORD-E4-003', 'confirmed', '16:47:12', '16:48:06', '72,38 s', '—'),
    ('ORD-E4-004', 'confirmed', '16:48:29', '16:49:04', '130,46 s', '58,08 s'),
    ('ORD-E4-005', 'no_stock', '16:49:33', '16:50:04', '190,13 s', '59,67 s'),
    ('ORD-E4-006', 'confirmed', '16:50:19', '16:50:55', '241,80 s', '51,67 s'),
    ('ORD-E4-007', 'confirmed', '16:51:12', '16:51:50', '297,04 s', '55,24 s'),
    ('ORD-E4-008', 'confirmed', '16:52:06', '16:52:38', '344,58 s', '47,54 s'),
    ('ORD-E4-009', 'no_stock', '16:52:53', '16:53:28', '394,29 s', '49,71 s'),
    ('ORD-E4-010', 'confirmed', '16:53:39', '16:54:15', '441,75 s', '47,46 s'),
    ('ORD-E4-011', 'confirmed', '16:54:28', '16:55:01', '487,30 s', '45,55 s'),
    ('ORD-E4-012', 'no_stock', '16:55:12', '16:55:47', '533,94 s', '46,64 s'),
    ('Media', '—', '—', '—', '313,37 s', '51,28 s'),
    ('Desvío', '—', '—', '—', '154,57 s', '5,22 s'),
]
j = idx('Tabla I.2: Registros en la base')
t = d.add_table(rows=len(FILAS), cols=6)
try:
    t.style = d.tables[18].style
except Exception:
    t.style = 'Table Grid'
for ri, fila in enumerate(FILAS):
    for ci, v in enumerate(fila):
        set_cell(t, ri, ci, v)
d.paragraphs[j]._element.addnext(t._element)

# --- alta en el listado de tablas ---
t1 = d.tables[1]
import copy
for ri, r in enumerate(t1.rows):
    if r.cells[0].text.strip() == 'Tabla I.1':
        copia = copy.deepcopy(r._tr)
        r._tr.addnext(copia)
        set_cell(t1, ri + 1, 0, 'Tabla I.2')
        set_cell(t1, ri + 1, 1, 'Registros en la base de datos de las órdenes procesadas manualmente')
        set_cell(t1, ri + 1, 2, 'Anexo I')
        break

d.save(RUTA)
print('780x reencuadrado como marginal; Tabla I.2 agregada')

d2 = Document(RUTA)
for ti, t in enumerate(d2.tables):
    if t.rows[0].cells[0].text.strip() == 'Orden' and len(t.columns) == 6:
        print('Tabla I.2 en índice %d (%d filas)' % (ti, len(t.rows)))
        for r in list(t.rows)[:3] + list(t.rows)[-2:]:
            print('   ' + ' | '.join(c.text.strip() for c in r.cells))
