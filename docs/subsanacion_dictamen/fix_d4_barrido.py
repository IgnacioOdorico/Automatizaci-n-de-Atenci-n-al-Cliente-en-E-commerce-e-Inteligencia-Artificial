# -*- coding: utf-8 -*-
"""Cuarto dictamen — barrido propio de la MISMA CLASE de defecto.

Las tres observaciones del dictamen son contradicciones que introdujo la ronda
anterior. Barrido el documento completo buscando mas de lo mismo, aparecieron
cuatro:

  4. §3.5.4 (plan de analisis) declara descriptivos + Wilson + kappa y NO
     declara el contraste de H1 que la §5.3 ahora ejecuta. Un capitulo de
     metodo tiene que anunciar las pruebas que despues corre; si no, el jurado
     lee en el Capitulo 5 una prueba que nadie planifico.
  5. La Tabla 5.5 conserva la columna "Resolución automática", que es
     exactamente el rotulo que la §5.2.1 declaro que sobreestima lo medido.
  6. La linea futura del baseline (§7.2) dice que replicar sobre varios
     operadores "habilitaria un contraste inferencial": el contraste ya esta
     hecho. Lo que ese diseño agregaria es otra cosa.
  7. El ABSTRACT reporta el contraste y el IC por Fieller; el RESUMEN en
     castellano no. Los dos deben ser espejo.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def par(pref):
    return d.paragraphs[idx(pref)]


def reescribir(pref, texto):
    p = par(pref)
    assert 'blip' not in p._element.xml, 'lleva imagen: no reescribir sus runs'
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ''
    return p


# ==================================================== 4 — el plan de analisis
p = par('Los tiempos de procesamiento (MTTD, MTTR y TMR) se reportan')
insert_paragraph_after(p,
    'Para el contraste de H1 —que compara la serie de tiempos del pipeline contra la del '
    'procesamiento manual— se adopta como prueba principal la U de Mann-Whitney, elegida por lo '
    'que acaba de declararse sobre la distribución: al no asumirse normalidad, la prueba de '
    'rangos es la que corresponde. Se la acompaña de la t de Welch, que no supone igualdad de '
    'varianzas, a título de comprobación complementaria, y del tamaño del efecto en ambas escalas '
    '—correlación rango-biserial y d de Cohen—, porque con muestras de este tamaño el valor p por '
    'sí solo dice poco. El nivel de significación se fija en 0,05 y todos los contrastes son '
    'bilaterales. El intervalo de confianza del factor de mejora, que es un cociente de dos '
    'medias, se calcula por el teorema de Fieller y no por propagación simplificada de la '
    'incertidumbre del numerador. Se anticipa aquí el límite que el Capítulo 5 retoma: una prueba '
    'de contraste establece si la diferencia observada es atribuible al azar del muestreo, pero '
    'no convierte el diseño en experimental, porque no hay asignación al azar a condiciones ni '
    'grupo de control.',
    estilo='Body Text')
print('4  §3.5.4 declara ahora el contraste que la §5.3 ejecuta')

# ==================================================== 5 — Tabla 5.5
t55 = None
for t in d.tables:
    if t.rows[0].cells[-1].text.strip() == 'Resolución automática':
        t55 = t
        break
assert t55 is not None, 'no encontre la Tabla 5.5'
col = len(t55.rows[0].cells) - 1
set_cell(t55, 0, col, 'No escalada')
print('5  Tabla 5.5: columna "Resolución automática" -> "No escalada"')

# ==================================================== 6 — linea futura §7.2
reescribir('Medición del baseline manual con un diseño de mayor alcance',
    'Medición del baseline manual con un diseño de mayor alcance: replicar el protocolo de la '
    'Sección 3.5.5 sobre varios operadores independientes y sin entrenamiento previo. El '
    'contraste de H1 ya está ejecutado (Sección 5.3), de modo que lo que ese diseño agregaría no '
    'es la prueba sino precisamente aquello que la prueba no puede dar: una estimación de la '
    'variabilidad entre operadores, la posibilidad de un grupo de control asignado al azar y, con '
    'ello, la atribución de la diferencia al procedimiento y no al caso particular medido.')
print('6  §7.2: la línea futura ya no promete una prueba que el trabajo hizo')

# ==================================================== 7 — resumen espejo del abstract
n = replace_in_paragraph(par('Los resultados obtenidos muestran un MTTD promedio'),
    'cronometrado por el propio equipo sobre diez órdenes en 49,13 segundos por orden '
    '(IC 95 %: 43,3 s a 55,0 s);',
    'cronometrado por el propio equipo sobre diez órdenes en 49,13 segundos por orden '
    '(IC 95 %: 43,30 s a 54,95 s; IC 95 % del factor por el teorema de Fieller: 686× a 875×). '
    'Ambas series están completamente separadas y el contraste es concluyente (U de Mann-Whitney '
    '= 0, p = 2,7 × 10⁻¹¹; t de Welch = 19,05, gl = 9,0, p = 1,4 × 10⁻⁸);')
print('7  RESUMEN en castellano: %d  (espejo del ABSTRACT)' % n)

d.save(RUTA)
print()
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
P2 = [p.text.strip() for p in d2.paragraphs]
TXT_P = '\n'.join(P2)
TODO = TXT_P + '\n' + '\n'.join(
    c.text for t in d2.tables for r in t.rows for c in r.cells)
fallos = []


def check(rot, cond):
    print('  [%s] %s' % ('OK   ' if cond else 'FALLA', rot))
    if not cond:
        fallos.append(rot)


print()
print('=== control ===')
i354 = next(k for k, t in enumerate(P2) if t.startswith('3.5.4'))
i355 = next(k for k, t in enumerate(P2) if t.startswith('3.5.5'))
bloque354 = '\n'.join(P2[i354:i355])
check('4a  §3.5.4 declara Mann-Whitney', 'Mann-Whitney' in bloque354)
check('4b  §3.5.4 declara Welch', 'Welch' in bloque354)
check('4c  §3.5.4 declara Fieller', 'Fieller' in bloque354)
check('4d  §3.5.4 fija el nivel de significación', 'se fija en 0,05' in bloque354)
check('4e  §3.5.4 justifica la elección no paramétrica',
      'al no asumirse normalidad, la prueba de rangos es la que corresponde' in bloque354)
check('5   sin la columna "Resolución automática"', 'Resolución automática' not in TODO)
check('5b  la columna se llama "No escalada"',
      any(r.cells[-1].text.strip() == 'No escalada'
          for t in d2.tables for r in t.rows[:1]))
check('6   §7.2 no promete una prueba ya hecha',
      'habilitaría un contraste inferencial' not in TODO
      and 'El contraste de H1 ya está ejecutado' in TODO)
i_res = next(k for k, t in enumerate(P2) if t == 'RESUMEN')
i_abs = next(k for k, t in enumerate(P2) if t == 'ABSTRACT')
resumen = '\n'.join(P2[i_res:i_abs])
abstract = '\n'.join(P2[i_abs:i_abs + 6])
check('7a  el RESUMEN reporta el contraste',
      'Mann-Whitney' in resumen and 'Welch' in resumen)
check('7b  el RESUMEN reporta el IC por Fieller', 'Fieller' in resumen)
check('7c  el ABSTRACT sigue reportándolos',
      'Mann-Whitney' in abstract and 'Fieller' in abstract)

print()
if fallos:
    print('FALLAS: %d' % len(fallos))
    for f in fallos:
        print('  - %s' % f)
    sys.exit(1)
print('SIN FALLAS')
