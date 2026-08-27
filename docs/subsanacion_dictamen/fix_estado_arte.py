# -*- coding: utf-8 -*-
"""§2.5 Estado del arte (#5) y §2.6 Ley 25.326 (#13), más las 5 referencias nuevas.

TODAS las referencias fueron verificadas contra fuente autorizada antes de escribirlas:
  Ngai et al. (2021)        -> Crossref, DOI 10.1016/j.elerap.2021.101098
  Parikh et al. (2023)      -> ACL Anthology 2023.acl-industry.71 (orden de autores confirmado)
  Luo et al. (2023)         -> arXiv 2309.14779
  Pypłacz y Žukovskis (2023)-> Crossref, DOI 10.1016/j.procs.2023.10.018
  Zhang et al. (2021)       -> Crossref, DOI 10.1016/j.ijinfomgt.2020.102304
  Ley 25.326                -> InfoLeg (texto oficial): arts. 5, 6, 9, 10, 12, 14 y 16
"""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def buscar_idx(doc, prefijo):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(prefijo):
            return i
    raise KeyError(prefijo)


# ============================================================
#  1. Bibliografía — de mayor a menor índice
# ============================================================
NUEVAS = [
    ('Zeithaml, V. A.',
     'Zhang, D., Pee, L. G., & Cui, L. (2021). Artificial intelligence in e-commerce fulfillment: '
     'A case study of resource orchestration at Alibaba’s Smart Warehouse. International Journal '
     'of Information Management, 57, 102304. https://doi.org/10.1016/j.ijinfomgt.2020.102304'),
    ('PostgreSQL Global',
     'Pypłacz, P., & Žukovskis, J. (2023). Implementing robotic process automation in small and '
     'medium-sized enterprises — Implications for organisations. Procedia Computer Science, 225, '
     '337–346. https://doi.org/10.1016/j.procs.2023.10.018'),
    ('OpenAI. (2024)',
     'Parikh, S., Tiwari, M., Tumbade, P., & Vohra, Q. (2023). Exploring zero and few-shot '
     'techniques for intent classification. En Proceedings of the 61st Annual Meeting of the '
     'Association for Computational Linguistics (Vol. 5: Industry Track) (pp. 744–751). '
     'Association for Computational Linguistics. https://doi.org/10.18653/v1/2023.acl-industry.71'),
    ('n8n GmbH. (2025)',
     'Ngai, E. W. T., Lee, M. C. M., Luo, M., Chan, P. S. L., & Liang, T. (2021). An intelligent '
     'knowledge-based chatbot for customer service. Electronic Commerce Research and '
     'Applications, 50, 101098. https://doi.org/10.1016/j.elerap.2021.101098'),
    ('Liu, P., Yuan, W.',
     'Luo, H., Liu, P., & Esping, S. (2023). Towards data-efficient customer intent recognition '
     'with prompt-based learning paradigm. arXiv. https://arxiv.org/abs/2309.14779'),
]
for ancla, texto in NUEVAS:
    i = buscar_idx(d, ancla)
    insertar_bloque(d, i, [(texto, d.paragraphs[i].style.name)])
    print('  + %s' % texto[:60])

# --- Ley 25.326 va al final del listado, después de Zhang ---
i = buscar_idx(d, 'Zhang, D., Pee')
insertar_bloque(d, i, [
    ('Ley 25.326 de Protección de los Datos Personales. (2000). Boletín Oficial de la República '
     'Argentina. https://servicios.infoleg.gob.ar/infolegInternet/anexos/60000-64999/64790/norma.htm',
     d.paragraphs[i].style.name)])

d.save(RUTA)
print('Bibliografía: 6 entradas nuevas')

# ============================================================
#  2. §2.5 y §2.6 — después de P146
# ============================================================
d = Document(RUTA)

BLOQUE = [
    ('2.5 Estado del arte', 'Heading 2'),

    ('Para situar la contribución de este trabajo corresponde revisar qué se ha publicado sobre el '
     'problema abordado. Se declara en primer término el procedimiento seguido, de modo que tanto '
     'el alcance de la revisión como sus omisiones resulten atribuibles.', 'First Paragraph'),

    ('Procedimiento de búsqueda. Se consultaron ACL Anthology, arXiv, Crossref y los catálogos en '
     'línea de ScienceDirect y Springer Nature. Las cadenas combinaron los términos chatbot, '
     'customer service, e-commerce, intent classification, large language model, zero-shot, '
     'business process automation, robotic process automation, SME y order fulfillment, en inglés, '
     'con sus equivalentes en español para las búsquedas de alcance regional. La ventana temporal '
     'se fijó entre 2018 y 2026, con la excepción de las obras canónicas ya incorporadas al marco '
     'teórico. Se incluyeron trabajos empíricos que reportaran métricas cuantitativas, revisiones '
     'sistemáticas y preprints de laboratorios o conferencias reconocidas; se excluyeron material '
     'de divulgación, documentación de producto y publicaciones sin proceso de revisión.',
     'Body Text'),

    ('Limitación declarada. La revisión que sigue no constituye una revisión sistemática en el '
     'sentido de un protocolo formal de cribado con doble evaluador: es una revisión de '
     'antecedentes acotada, proporcionada al nivel de un trabajo integrador de tecnicatura. Se la '
     'presenta como tal y no como hallazgo exhaustivo del campo.', 'Body Text'),

    ('2.5.1 Chatbots en la atención al cliente de comercio electrónico', 'Heading 3'),

    ('Ngai et al. (2021) presentan un chatbot basado en conocimiento para atención al cliente, '
     'construido sobre una representación explícita del dominio antes de la difusión de los '
     'modelos de lenguaje de gran escala. Su trabajo resulta pertinente como línea de base '
     'arquitectónica: muestra el costo de construir y mantener la base de conocimiento que un '
     'sistema de reglas exige, costo que es precisamente el que un modelo de lenguaje permite '
     'evitar. Misischia et al. (2022), por su parte, abordan la cuestión desde la calidad de '
     'servicio antes que desde la técnica: identifican las funciones del chatbot relevantes para '
     'el cliente y sostienen que atributos como un estilo de interacción socialmente orientado y '
     'un trato empático inciden sobre la calidad percibida de la atención. Ambos trabajos '
     'coinciden en un punto que este trabajo asume como supuesto de diseño: la utilidad del '
     'chatbot no se agota en la exactitud de su clasificación, sino que depende también de la '
     'calidad de la respuesta que entrega.', 'First Paragraph'),

    ('2.5.2 Clasificación de intenciones con modelos de lenguaje', 'Heading 3'),

    ('Parikh et al. (2023) evalúan cuatro estrategias de bajo consumo de datos para clasificación '
     'de intenciones —adaptación de dominio, aumento de datos, prompting zero-shot con modelos de '
     'lenguaje y ajuste fino de parámetros eficientes— y concluyen que el ajuste fino eficiente '
     'sobre Flan-T5 obtiene el mejor desempeño incluso con un único ejemplo por intención, al '
     'tiempo que confirman que el prompting zero-shot basado en descripciones de las intenciones '
     'resulta efectivo para la tarea. Luo et al. (2023) llegan a una conclusión convergente por '
     'otro camino: muestran que el aprendizaje basado en prompts permite a modelos de lenguaje '
     'pequeños alcanzar un desempeño competitivo en reconocimiento de intenciones de cliente con '
     'una cantidad mínima de datos de entrenamiento, y que un prompt bien diseñado habilita '
     'desempeño útil en escenarios zero-shot.', 'First Paragraph'),

    ('Estos antecedentes son los que dan sentido al régimen zero-shot que emplea este trabajo '
     '(Sección 2.2.2) y permiten interpretar su resultado: el accuracy obtenido no es un valor '
     'aislado sino un punto dentro de un rango que la literatura ya documenta como alcanzable sin '
     'ajuste fino. Corresponde señalar, sin embargo, una diferencia de alcance que impide la '
     'comparación directa de cifras: los trabajos citados evalúan sobre corpus públicos '
     'normalizados —del tipo de CLINC150, Banking77 o SNIPS—, con decenas o centenas de '
     'intenciones, mientras que este trabajo opera sobre cuatro categorías de un dominio '
     'deliberadamente acotado. Un accuracy más alto sobre menos clases no constituye un mejor '
     'resultado, y este trabajo no lo presenta como tal.', 'Body Text'),

    ('2.5.3 Automatización de procesos en pequeñas y medianas empresas', 'Heading 3'),

    ('Pypłacz y Žukovskis (2023) estudian la implementación de automatización robótica de procesos '
     'en pequeñas y medianas empresas y sus implicancias organizacionales, y documentan que la '
     'adopción efectiva permanece muy por debajo del interés declarado. Su diagnóstico sobre las '
     'barreras —infraestructura digital limitada, brechas de competencias y procedimientos poco '
     'formalizados— es coherente con el problema que motiva este trabajo y respalda empíricamente '
     'la elección de una herramienta de bajo costo y auto-alojable. En el extremo opuesto de la '
     'escala organizacional, Zhang et al. (2021) analizan mediante estudio de caso la orquestación '
     'de recursos de inteligencia artificial en el centro logístico de Alibaba, y muestran cómo la '
     'capacidad emerge de la articulación entre los recursos técnicos y los humanos y '
     'organizacionales, y no del componente tecnológico en aislamiento. Ese hallazgo es pertinente '
     'aquí en sentido admonitorio: sugiere que las mejoras de tiempo que un prototipo mide en '
     'condiciones de laboratorio no se trasladan automáticamente a una organización real.',
     'First Paragraph'),

    ('2.5.4 Vacío identificado y contribución de este trabajo', 'Heading 3'),

    ('La revisión permite delimitar la contribución con precisión, y también su modestia. Los '
     'trabajos sobre clasificación de intenciones miden el componente de inteligencia artificial '
     'de manera aislada, sobre corpus normalizados y sin integración con un sistema operativo real; '
     'los trabajos sobre automatización de procesos en pequeñas y medianas empresas describen la '
     'adopción y sus barreras, pero no instrumentan métricas de tiempo extremo a extremo sobre un '
     'artefacto concreto; y el estudio de caso de fulfillment opera a una escala organizacional que '
     'no es la del segmento aquí considerado. No se identificaron trabajos que midan, sobre un '
     'mismo artefacto y con instrumentación embebida, tanto el tiempo de procesamiento de órdenes '
     'como el de respuesta conversacional, y que además comparen ese resultado contra un baseline '
     'manual cronometrado por los propios autores.', 'First Paragraph'),

    ('La Tabla 2.1 sitúa este trabajo frente a los antecedentes revisados. Su contribución no '
     'reside en superar a ninguno de ellos en su propia métrica, cosa que no hace ni pretende, '
     'sino en articular dos componentes que la literatura suele tratar por separado dentro de un '
     'artefacto reproducible, y en medirlos con instrumentación propia contra un término de '
     'comparación medido y no supuesto.', 'Body Text'),

    ('Tabla 2.1: Posición de este trabajo respecto de los antecedentes revisados.', 'Body Text'),

    ('2.6 Tratamiento de datos personales y marco legal aplicable', 'Heading 2'),

    ('El dominio de aplicación de este trabajo es la atención al cliente, y en consecuencia el '
     'sistema manipula datos personales: nombre, dirección de correo electrónico y número de '
     'teléfono de quien realiza una compra, además del contenido íntegro de los mensajes que esa '
     'persona envía. Corresponde por lo tanto explicitar el marco legal aplicable, aun cuando el '
     'prototipo opere sobre datos ficticios y en un entorno local.', 'First Paragraph'),

    ('En la República Argentina el tratamiento de datos personales se rige por la Ley 25.326 de '
     'Protección de los Datos Personales, sancionada en el año 2000. Cinco de sus disposiciones '
     'resultan directamente pertinentes al sistema construido. El artículo 5 exige el '
     'consentimiento libre, expreso e informado del titular como condición de licitud del '
     'tratamiento. El artículo 6 impone el deber de informarle la finalidad de la recolección y '
     'los derechos que le asisten. Los artículos 9 y 10 establecen, respectivamente, la obligación '
     'de adoptar medidas técnicas y organizativas que garanticen la seguridad de los datos y el '
     'deber de confidencialidad de quienes los tratan. Los artículos 14 y 16 reconocen al titular '
     'los derechos de acceso, rectificación, actualización y supresión. Y el artículo 12 regula la '
     'transferencia internacional de datos, prohibiéndola hacia países que no proporcionen niveles '
     'adecuados de protección.', 'Body Text'),

    ('Esta última disposición interpela directamente a la arquitectura del Flujo 2. El sistema '
     'remite el texto completo del mensaje del cliente a la interfaz de programación de OpenAI, '
     'alojada fuera del territorio nacional, en cada interacción. Se trata de una transferencia '
     'internacional de datos personales en el sentido del artículo 12, y un despliegue productivo '
     'debería resolverla mediante alguna de las vías que la propia norma y su reglamentación '
     'contemplan, o bien evitarla ejecutando el modelo de lenguaje en infraestructura propia, '
     'alternativa que este trabajo recomienda evaluar en el Capítulo 7. La contradicción entre '
     'esta dependencia y el criterio de soberanía de datos invocado en la Sección 2.1.2 se declara '
     'allí de manera explícita.', 'Body Text'),

    ('Existe además un deber que no deriva de la Ley 25.326 sino de la buena fe en la relación de '
     'consumo, y que un sistema conversacional automatizado debe atender: informar al cliente que '
     'está interactuando con un sistema y no con una persona. El prompt del asistente construido '
     'en este trabajo instruye al modelo a adoptar un tono humano y a no presentarse como un '
     'chatbot genérico (Anexo H), decisión que optimiza la calidad percibida de la interacción '
     'pero que, en un despliegue real, debería acompañarse de una identificación inequívoca del '
     'carácter automatizado del canal.', 'Body Text'),

    ('Qué implementa el prototipo y qué faltaría. En su estado actual el sistema opera '
     'exclusivamente sobre datos ficticios, generados para las pruebas sobre dominios reservados '
     'para documentación, y persiste la totalidad de la información en una base de datos local, de '
     'modo que no existe tratamiento de datos personales reales ni riesgo asociado. Un despliegue '
     'productivo requeriría, como mínimo: obtener y registrar el consentimiento del titular en el '
     'momento de la compra; publicar la finalidad del tratamiento y el procedimiento para ejercer '
     'los derechos de acceso, rectificación y supresión; cifrar los datos en tránsito y en reposo; '
     'definir un plazo de retención y un procedimiento de supresión que hoy no existen, dado que '
     'el esquema conserva las interacciones de forma indefinida; minimizar los datos remitidos al '
     'proveedor de inferencia, que actualmente recibe el mensaje íntegro sin depuración previa; y '
     'resolver la base legal de la transferencia internacional. Estas obligaciones se incorporan a '
     'las recomendaciones del Capítulo 7.', 'Body Text'),
]

creados = insertar_bloque(d, 146, BLOQUE)
print('§2.5 y §2.6 insertadas: %d párrafos' % len(creados))

# --- Tabla 2.1 comparativa ---
FILAS = [
    ('Antecedente', 'Dominio', 'Enfoque técnico', 'Qué mide', 'Entorno'),
    ('Ngai et al. (2021)', 'Atención al cliente', 'Chatbot basado en conocimiento explícito',
     'Desempeño del asistente sobre consultas de clientes', 'Organización real'),
    ('Misischia et al. (2022)', 'Atención al cliente', 'Revisión de funciones del chatbot',
     'Calidad de servicio percibida', 'Revisión'),
    ('Parikh et al. (2023)', 'Clasificación de intenciones',
     'Zero-shot, aumento de datos y ajuste fino eficiente', 'Accuracy sobre corpus públicos',
     'Corpus normalizados'),
    ('Luo et al. (2023)', 'Intenciones de cliente', 'Aprendizaje basado en prompts',
     'Accuracy con datos escasos', 'Corpus de industria'),
    ('Pypłacz y Žukovskis (2023)', 'Procesos de PyME', 'Automatización robótica de procesos',
     'Adopción, barreras e implicancias organizacionales', 'Estudio de campo'),
    ('Zhang et al. (2021)', 'Fulfillment de e-commerce', 'Orquestación de recursos con IA',
     'Eficiencia operativa y exactitud de órdenes', 'Estudio de caso (Alibaba)'),
    ('Este trabajo', 'Ciclo post-venta de PyME',
     'Orquestación con n8n y clasificación zero-shot con LLM',
     'MTTD, MTTR, TMR y accuracy, contra baseline manual medido',
     'Prototipo local instrumentado'),
]
idx = buscar_idx(d, 'Tabla 2.1: Posición')
t = d.add_table(rows=len(FILAS), cols=5)
try:
    t.style = d.tables[18].style
except Exception:
    t.style = 'Table Grid'
for ri, fila in enumerate(FILAS):
    for ci, v in enumerate(fila):
        set_cell(t, ri, ci, v)
d.paragraphs[idx]._element.addnext(t._element)

d.save(RUTA)

# ============================================================
d2 = Document(RUTA)
print()
for i in range(145, 175):
    print('%4d [%-16s] %s' % (i, d2.paragraphs[i].style.name, d2.paragraphs[i].text.strip()[:105]))
