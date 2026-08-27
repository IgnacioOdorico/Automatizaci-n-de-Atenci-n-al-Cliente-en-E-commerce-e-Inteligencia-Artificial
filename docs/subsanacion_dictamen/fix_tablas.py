# -*- coding: utf-8 -*-
"""Barrida de tablas: hallazgos #6, #8h, #12, #15, #16, #23."""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)

# ===== #15 — Tabla 3.3: residuo de "107 mensajes" =====
t33 = d.tables[5]
set_cell(t33, 3, 2,
    'Comparación de la intención asignada por el modelo contra la intención de referencia '
    'para los 150 mensajes del corpus etiquetado')

# ===== #23 — Tabla 4.8: códigos internos y el panel de valor constante =====
t48 = d.tables[13]
set_cell(t48, 0, 3, 'Fuente de datos')
for r in (1, 2, 3):
    set_cell(t48, r, 3, 'orders (filtrada a la corrida de carga secuencial)')
set_cell(t48, 9, 3, 'Valor constante')
set_cell(t48, 9, 4,
    'Accuracy de clasificación del corpus. No se deriva de una consulta SQL: su cálculo '
    'requiere las etiquetas de referencia, que no residen en la base de datos operativa. '
    'El panel muestra el valor obtenido en el Anexo H.')

# ===== #16 — Tabla 5.4: la columna rotulaba la distribución predicha =====
t54 = d.tables[19]
set_cell(t54, 0, 1, 'Mensajes clasificados')

# ===== #6 — Tabla 5.9: intervalo de Wilson correcto =====
t59 = d.tables[24]
set_cell(t59, 3, 2, '92,7 % (139/150); IC 95 % de Wilson [87,3 %; 95,9 %]')
set_cell(t59, 1, 1, 'End-to-end menor que el baseline manual; < 30 s como criterio operativo')
set_cell(t59, 1, 2,
    '0,063 s frente a 49,13 s del baseline manual (factor ≈ 780×; IC 95 % 687× a 872×). '
    'Muy por debajo del criterio de 30 s.')

# ===== #12 y #15 — Tabla 6.1 =====
t61 = d.tables[25]
set_cell(t61, 2, 2,
    'TMR: 1,47 s en el canal simulado vía SMTP local (n = 150) y 3,07 s en el canal Telegram '
    'real (n = 45). Accuracy de clasificación: 92,7 % sobre el corpus etiquetado. Dos canales '
    'implementados y validados; Gmail queda como diseño (§7).')
set_cell(t61, 2, 3, 'CUMPLIDO (dos canales)')

set_cell(t61, 3, 1, 'Schema PostgreSQL con las tablas, vistas e índices necesarios')
set_cell(t61, 3, 2,
    '7 tablas (products, orders, order_items, interactions, tickets, faq_responses y '
    'stock_alerts), 6 vistas de métricas y 12 índices de rendimiento sobre las columnas de '
    'filtrado y de unión. El objetivo comprometía 5 tablas y 5 vistas: el esquema final las '
    'supera por incorporación de order_items, stock_alerts y la vista del corpus.')

set_cell(t61, 4, 2,
    '13 paneles distribuidos en dos tableros: 6 para el pipeline de órdenes y 7 para el '
    'chatbot, con conexión directa a las vistas de PostgreSQL.')

set_cell(t61, 5, 2,
    '10 pruebas funcionales (5 por flujo), 1 corrida de carga secuencial de 50 órdenes y '
    '1 prueba de concurrencia de 6 rondas × 20 solicitudes simultáneas, todas documentadas.')

d.save(RUTA)

# ===== Verificación =====
d2 = Document(RUTA)
for ti, lbl in ((5, 'Tabla 3.3'), (13, 'Tabla 4.8'), (19, 'Tabla 5.4'),
                (24, 'Tabla 5.9'), (25, 'Tabla 6.1')):
    print('=== %s (idx %d) ===' % (lbl, ti))
    for ri, r in enumerate(d2.tables[ti].rows):
        print('  r%-2d | %s' % (ri, ' | '.join(c.text.strip()[:78] for c in r.cells)))
    print()
