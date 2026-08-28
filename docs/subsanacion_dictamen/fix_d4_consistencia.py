# -*- coding: utf-8 -*-
"""Cuarto dictamen: tres contradicciones internas que introdujo la ronda anterior.

Las tres son consecuencia de las correcciones del tercer dictamen y las tres son
legitimas. Verificadas en el documento antes de tocar nada:

  1. §5.4.1 (d) sigue afirmando que el trabajo "no ejecuta un contraste formal" y
     que los intervalos "no constituyen una prueba de hipotesis en sentido
     inferencial estricto". La §5.3 ahora reporta U de Mann-Whitney y t de Welch:
     contradiccion frontal.
  2. §6.2 cierra diciendo "7 tablas y 6 vistas frente a las 5 y 5 previstas, y 13
     paneles frente a los 7 enunciados". Se habia corregido la CELDA OE3 de la
     Tabla 6.1 pero no el parrafo de prosa que la sigue.
  3. §4.3.3 remite a que la captura en la capa HTTP "queda planteada en el
     Capitulo 7". Se verifico: el Capitulo 7 no la contiene. Referencia colgada.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def reescribir(pref, texto):
    p = d.paragraphs[idx(pref)]
    assert 'blip' not in p._element.xml, 'el parrafo lleva una imagen: no reescribir sus runs'
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ''
    return p


# ============================================================ 1 — §5.4.1 (d)
reescribir('(d) Ausencia de contraste inferencial contra un grupo de control',
    '(d) Ausencia de grupo de control y de diseño experimental comparativo: el contraste de H1 se '
    'ejecutó y se reporta al cierre de la Sección 5.3 —U de Mann-Whitney sobre las dos series y t '
    'de Welch sobre sus descriptivos—, de modo que la diferencia observada no es atribuible al '
    'azar del muestreo. Lo que ese contraste no resuelve es de otro orden y conviene separarlo en '
    'tres puntos. Primero, el baseline manual se midió sobre un único operador: la prueba compara '
    'dos series de tiempos, no dos poblaciones de operadores, y por lo tanto no permite estimar '
    'la variabilidad entre personas ni descartar el efecto del evaluador. Segundo, no hubo '
    'asignación al azar a condiciones ni grupo de control; las dos series provienen de dos '
    'procedimientos distintos aplicados a una tarea equivalente, no de un experimento controlado, '
    'de modo que la atribución causal de la diferencia a la automatización descansa en el diseño '
    'del caso y no en la aleatorización. Tercero, y como se precisa en la propia Sección 5.3, la '
    'significación estadística confirma que la diferencia existe pero no resuelve la asimetría de '
    'constructo entre ambos términos: cada instrumento cubre un alcance distinto. Una comparación '
    'inferencial en sentido pleno exigiría varios operadores independientes y un diseño '
    'experimental que excede el alcance declarado de este trabajo, y queda planteada como línea '
    'futura en el Capítulo 7.')
print('1  §5.4.1 (d) reescrito: se conserva el nucleo valido y se retira lo desmentido por 5.3')

# ============================================================ 2 — cierre de §6.2
reescribir('Los cinco objetivos específicos fueron cumplidos.',
    'Los cinco objetivos específicos fueron cumplidos, con dos precisiones sobre el alcance de '
    'ese cumplimiento. OE3 y OE4 se formularon en términos cualitativos en la Sección 1.5.2 —las '
    'tablas, vistas de métricas e índices necesarios para instrumentar el ciclo, y los dashboards '
    'para visualizar MTTD, MTTR y TMR— y el sistema desplegado los satisface: siete tablas, seis '
    'vistas y doce índices sostienen el cálculo de las tres métricas sin cómputo externo, y trece '
    'paneles repartidos en dos tableros las visualizan en tiempo real. No se los contrasta contra '
    'una cifra comprometida de antemano porque los objetivos no la fijaron. OE1 se cumple en '
    'régimen secuencial: bajo concurrencia el pipeline conserva la integridad transaccional pero '
    'pierde el 40,8 % de las órdenes, según se consigna en su propia celda y se analiza en la '
    'Sección 5.1.3. OE2 se cumplió sobre los dos canales efectivamente implementados y medidos, '
    'quedando el canal de correo como diseño documentado en el Capítulo 7.')
print('2  cierre de §6.2 alineado con la celda OE3 de la Tabla 6.1')

# ============================================================ 3 — §7.1
p_admision = d.paragraphs[idx('Control de admisión y encolado de solicitudes')]
insert_paragraph_after(p_admision,
    'Capturar la marca de recepción en la capa HTTP: hoy received_at la escribe el nodo Registrar '
    'Orden, ya dentro del pipeline, de modo que el MTTD mide el intervalo entre dos escrituras '
    'sobre la base local y no el que media entre el arribo del pedido y su detección '
    '(Sección 4.3.3). Tomar la marca en el punto de entrada —desde el propio webhook, antes de la '
    'primera escritura— corrige la operacionalización sin alterar la definición de la métrica, y '
    'permite además separar la latencia de red de la de procesamiento. Es una corrección '
    'instrumental de bajo costo que vuelve al MTTD comparable con el de un despliegue real.',
    estilo=p_admision.style.name)
print('3  ítem de la capa HTTP incorporado a §7.1: la remisión de §4.3.3 deja de colgar')

d.save(RUTA)
print()
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
P2 = [p.text.strip() for p in d2.paragraphs]
TODO = '\n'.join(P2) + '\n' + '\n'.join(
    c.text for t in d2.tables for r in t.rows for c in r.cells)
fallos = []


def check(rot, cond):
    print('  [%s] %s' % ('OK   ' if cond else 'FALLA', rot))
    if not cond:
        fallos.append(rot)


print()
print('=== control ===')
check('1a  sin "no ejecuta un contraste formal"', 'no ejecuta un contraste formal' not in TODO)
check('1b  sin "no constituyen una prueba de hipótesis"',
      'no constituyen una prueba de hipótesis' not in TODO)
check('1c  (d) reconoce el contraste ejecutado',
      'el contraste de H1 se ejecutó y se reporta al cierre de la Sección 5.3' in TODO)
check('1d  (d) conserva el operador único',
      'el baseline manual se midió sobre un único operador' in TODO)
check('1e  (d) conserva la ausencia de aleatorización',
      'no hubo asignación al azar a condiciones ni grupo de control' in TODO)
check('2a  sin "frente a las 5 y 5 previstas"', 'frente a las 5 y 5 previstas' not in TODO)
check('2b  sin "frente a los 7 enunciados"', 'frente a los 7 enunciados' not in TODO)
check('2c  §6.2 dice que los objetivos no fijaron cifra',
      'No se los contrasta contra una cifra comprometida de antemano' in TODO)
check('2d  §6.2 incorpora el 40,8 % de OE1',
      'pierde el 40,8 % de las órdenes' in TODO)
check('3a  §7.1 contiene el ítem de la capa HTTP',
      'Capturar la marca de recepción en la capa HTTP' in TODO)

# el item tiene que estar DENTRO del Capitulo 7, no en otro lado
i7 = next(k for k, t in enumerate(P2) if t.startswith('CAPÍTULO 7'))
i8 = next(k for k, t in enumerate(P2) if t.startswith('CAPÍTULO 8'))
check('3b  y esta dentro del Capítulo 7',
      any('Capturar la marca de recepción en la capa HTTP' in P2[k] for k in range(i7, i8)))

# ninguna otra remision colgada del mismo tipo
import re
check('3c  sin otras cifras "5 tablas y 5 vistas" sueltas',
      '5 tablas y 5 vistas' not in TODO)

print()
if fallos:
    print('FALLAS: %d' % len(fallos))
    sys.exit(1)
print('SIN FALLAS')
