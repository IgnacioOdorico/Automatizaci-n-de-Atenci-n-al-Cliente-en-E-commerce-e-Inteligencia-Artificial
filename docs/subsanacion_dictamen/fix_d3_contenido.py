# -*- coding: utf-8 -*-
"""Tercer dictamen (APROBADO CON OBSERVACIONES, 8,5/10) — bloque de CONTENIDO.

Cubre las obligatorias 2, 3, 5, 6, 7, 8a y 10, y las recomendadas 2 (MTTD),
3 (IC del factor por Fieller), 4 (terminologia), 5 (canal simulado) y 6 (nodos
5 y 6 del Flujo 2). La estructura del frontispicio va en fix_d3_estructura.py.

Verificado contra el documento antes de escribir:
  - Tabla I.1: 10 validas = 7 con stock + 3 sin stock; 2 descartadas (una de
    ellas 'sin stock', la otra sin rama consignada). El dictamen tiene razon:
    "nueve con stock y tres sin stock" no se sostiene.
  - Tabla 6.1 OE4 NO dice "13 frente a los 7 enunciados" (eso del dictamen esta
    desactualizado). El defecto real esta solo en OE3.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def reescribir(i, texto):
    p = d.paragraphs[i]
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ''


def tabla_por_celda(fila0col0, contiene=None):
    """Devuelve la tabla cuyo encabezado empieza asi (y opcionalmente contiene X)."""
    for t in d.tables:
        if t.rows[0].cells[0].text.strip() == fila0col0:
            if contiene is None or contiene in '\n'.join(
                    c.text for r in t.rows for c in r.cells):
                return t
    raise KeyError(fila0col0)


print('=' * 74)
print(' OBLIGATORIAS')
print('=' * 74)

# ---------------------------------------------------------------- OBS 2
# La composicion declarada no coincide con la Tabla I.1. Se escribe la real.
n = replace_in_paragraph(d.paragraphs[idx('Muestra e instrumento.')],
    'Se procesaron doce órdenes, compuestas por nueve con stock disponible y tres sin stock, '
    'proporción próxima a la distribución observada en la corrida automatizada (35 y 15 sobre 50). '
    'Las dos ramas demandan trabajo manual distinto y la muestra debe reflejar ambas.',
    'Se procesaron doce órdenes. De las diez que resultaron válidas, siete correspondieron a la '
    'rama con stock disponible y tres a la rama sin stock —una proporción de 70 % y 30 % que '
    'reproduce la distribución observada en la corrida automatizada (35 y 15 sobre 50)—; de las '
    'dos descartadas por interrupción del operador, una (ORD-E4-002) pertenecía a la rama sin '
    'stock y la otra (ORD-E4-001) se interrumpió antes de que su rama quedara consignada. Las dos '
    'ramas demandan trabajo manual distinto y la muestra debe reflejar ambas. La composición '
    'completa, orden por orden y con su marca de descarte, se transcribe en la Tabla I.1 del '
    'Anexo I.')
print('OBS 2  composicion de la muestra en 3.5.5 ......... %d' % n)

# ---------------------------------------------------------------- OBS 3
# El compromiso numerico "5 tablas y 5 vistas" no figura en 1.5.2. Se elimina
# la comparacion de la Tabla 6.1 en lugar de retro-escribir los objetivos.
t61 = tabla_por_celda('Obj.')
fila_oe3 = [i for i, r in enumerate(t61.rows) if r.cells[0].text.strip() == 'OE3'][0]
set_cell(t61, fila_oe3, 2,
    '7 tablas (products, orders, order_items, interactions, tickets, faq_responses y '
    'stock_alerts), 6 vistas de métricas y 12 índices de rendimiento sobre las columnas de '
    'filtrado y de unión. El objetivo se formuló en términos cualitativos —las tablas, vistas e '
    'índices necesarios para instrumentar el ciclo— y el esquema desplegado los provee: MTTD, '
    'MTTR y TMR se calculan íntegramente desde las vistas, sin cómputo externo.')
print('OBS 3  trazabilidad de OE3 (sin compromiso numerico) ... hecho')

# ---------------------------------------------------------------- OBS 6
fila_oe1 = [i for i, r in enumerate(t61.rows) if r.cells[0].text.strip() == 'OE1'][0]
set_cell(t61, fila_oe1, 2,
    'MTTD: 0,009 s / MTTR: 0,054 s / Total: 0,063 s (n = 50, corrida E1.a). 50/50 órdenes '
    'procesadas sin errores. Bajo concurrencia (E1.b, 120 órdenes) el pipeline sostuvo la '
    'atomicidad del descuento de stock sin sobreventa, pero perdió capacidad de proceso: 49 de '
    '120 órdenes (40,8 %) quedaron sin marca de processed_at y las métricas se degradaron por un '
    'factor aproximado de diez (Sección 5.1.3). El objetivo se da por cumplido en régimen '
    'secuencial; en régimen concurrente el comportamiento es seguro pero no elástico.')
set_cell(t61, fila_oe1, 3, 'CUMPLIDO (régimen secuencial)')
print('OBS 6  perdida del 40,8 %% incorporada a OE1 ....... hecho')

# ---------------------------------------------------------------- OBS 5
reescribir(idx('Tasa de resolución automática:'),
    'Tasa de no escalada: 73,3 % (110 de 150 mensajes). Si se excluyen los reclamos —que escalan '
    'a ticket por diseño del sistema y no por falla del clasificador—, la tasa asciende al 100 %: '
    'el chatbot respondió sin derivar a un operador humano la totalidad de los mensajes que no '
    'eran reclamos. Corresponde precisar qué mide este indicador y qué no mide. Mide que el flujo '
    'se completó sin escalar: el sistema clasificó el mensaje, generó una respuesta y la entregó '
    'por el canal de origen. No mide que esa respuesta fuera correcta. La corrección del '
    'contenido entregado al cliente no fue evaluada en este trabajo, y la omisión es material en '
    'las consultas de tipo FAQ: como se documenta en la Sección 4.4.3, el contexto de la base de '
    'conocimiento no llega al prompt, de modo que esas respuestas provienen del conocimiento '
    'general del modelo y no de las políticas efectivamente vigentes en el comercio. Llamar '
    '«tasa de resolución» a este indicador sobreestimaría lo medido; se lo denomina en '
    'consecuencia tasa de no escalada.')
print('OBS 5  no escalada != resolucion (5.2.1) ......... hecho')

# ---------------------------------------------------------------- OBS 7
# (a) se retira el vocabulario inferencial donde no hubo inferencia
for viejo, nuevo, rot in [
    ('reduce significativamente el tiempo end-to-end', 'reduce de manera sustancial el tiempo end-to-end', 'H1'),
    ('la hipótesis de reducción significativa de los tiempos operativos',
     'la hipótesis de reducción sustancial de los tiempos operativos', 'objetivo general'),
    ('la automatización reduce significativamente los tiempos operativos',
     'la automatización reduce de manera sustancial los tiempos operativos', '5.4'),
]:
    n = replace_everywhere(d, viejo, nuevo)
    print('OBS 7a "%s" ... %d' % (rot, n))

# (b) y se ejecuta el contraste, que con estos datos es concluyente
p_ancla = d.paragraphs[idx('• H2b se confirma con un accuracy global')]
insert_paragraph_after(p_ancla,
    'Contraste estadístico de H1. El dictamen de evaluación observó, con razón, que el verbo '
    '«reducir de manera sustancial» describe una diferencia de magnitud pero no una inferencia, y '
    'que el trabajo no había ejecutado ninguna prueba de contraste entre ambas series. Se '
    'incorpora aquí esa prueba. Las dos series están completamente separadas: el máximo de la '
    'serie automatizada (0,095 s sobre n = 50) es inferior al mínimo de la serie manual '
    '(42,754 s sobre n = 10), de modo que ningún par de observaciones se cruza. La U de '
    'Mann-Whitney vale por lo tanto 0, con un valor p exacto bilateral de 2/C(60, 10) = '
    '2,7 × 10⁻¹¹ y un tamaño del efecto rango-biserial de 1,000, el máximo posible. La t de '
    'Welch sobre los descriptivos publicados arroja t = 19,05 con 9,0 grados de libertad y '
    'p = 1,4 × 10⁻⁸ (d de Cohen = 15,3). Ambos contrastes se calculan íntegramente a partir de '
    'las cifras de las Tablas 5.2, 5.4 y I.1, de modo que son reproducibles sin acceso a los '
    'datos crudos. Corresponde una salvedad de alcance: la prueba establece que la diferencia '
    'observada no es atribuible al azar del muestreo, pero no resuelve la asimetría de '
    'constructo entre ambos términos —el alcance instrumentado de cada uno— que se discute en la '
    'Sección 5.4. La significación estadística confirma que la diferencia existe; su '
    'interpretación sigue estando acotada por lo que cada instrumento midió.',
    estilo='Body Text')
print('OBS 7b contraste U de Mann-Whitney y t de Welch ... insertado en 5.3')

# (c) Tabla 5.10: el veredicto de H1 pasa a apoyarse en la prueba
t510 = tabla_por_celda('Hipótesis')
set_cell(t510, 1, 0, 'H1: Pipeline sustancialmente más rápido que el proceso manual')
set_cell(t510, 1, 2,
    '0,063 s frente a 49,13 s del baseline manual (factor ≈ 780×; IC 95 % por el teorema de '
    'Fieller: 686× a 875×). U de Mann-Whitney = 0 con separación completa de ambas series '
    '(p = 2,7 × 10⁻¹¹); t de Welch = 19,05, gl = 9,0, p = 1,4 × 10⁻⁸. Muy por debajo del '
    'criterio operativo de 30 s.')
print('OBS 7c Tabla 5.10, fila H1 ...................... hecho')

# ---------------------------------------------------------------- OBS 8a
n = replace_everywhere(d, 'El panel muestra el valor obtenido en el Anexo H.',
                          'El panel muestra el valor obtenido en el Anexo J.')
print('OBS 8a Tabla 4.8: Anexo H -> Anexo J ............. %d' % n)

# ---------------------------------------------------------------- OBS 10
n = replace_everywhere(d, 'Workflow de 18 nodos funcionales: 3 canales, IA, tickets',
                          'Workflow de 18 nodos funcionales: dos canales (WhatsApp simulado y '
                          'Telegram real), IA, tickets')
print('OBS 10 Tabla 3.1 Etapa 4: 3 canales -> dos ....... %d' % n)

print()
print('=' * 74)
print(' RECOMENDADAS')
print('=' * 74)

# ---------------------------------------------------------------- REC 2 (MTTD)
n = replace_in_paragraph(d.paragraphs[idx('Ambas métricas se calculan en la vista')],
    'Cabe señalar que el destinatario es un capturador SMTP local: el tiempo de tránsito hasta la '
    'casilla real del cliente queda fuera de lo observable en este entorno y no forma parte de la '
    'métrica.',
    'Cabe señalar que el destinatario es un capturador SMTP local: el tiempo de tránsito hasta la '
    'casilla real del cliente queda fuera de lo observable en este entorno y no forma parte de la '
    'métrica. Sobre received_at corresponde una precisión que acota el alcance del MTTD. Esa '
    'marca la escribe el propio pipeline, ya en posesión del pedido, y no la capa HTTP que lo '
    'recibe: el intervalo medido no es el que media entre el arribo del pedido al sistema y su '
    'detección, sino el que separa dos escrituras consecutivas sobre la base local. El nombre '
    'heredado de la práctica ITIL promete por lo tanto más de lo que el instrumento entrega, y '
    'así se declara. La comparación contra el baseline manual conserva no obstante la simetría, '
    'porque ese baseline también excluye la latencia de detección y así se establece en la '
    'Sección 3.5.5. Capturar la marca de arribo en la capa HTTP, anterior a la escritura en base, '
    'queda planteado en el Capítulo 7 como corrección instrumental de una futura versión.')
print('REC 2  operacionalizacion del MTTD declarada .... %d' % n)

# ---------------------------------------------------------------- REC 3 (Fieller)
n = replace_in_paragraph(d.paragraphs[idx('• H1 se confirma en sus dos términos')],
    '(IC 95 % del factor: 687× a 872×)',
    '(IC 95 % del factor: 686× a 875×, por el teorema de Fieller)')
print('REC 3a IC del factor en 5.3 .................... %d' % n)

n = replace_in_paragraph(d.paragraphs[idx('Los resultados confirman la premisa central')],
    'La reducción del costo marginal de procesamiento de una orden es de un factor aproximado de '
    '780 veces respecto del baseline manual medido en la Sección 5.1.4 (49,13 s → 0,063 s).',
    'La reducción del costo marginal de procesamiento de una orden es de un factor aproximado de '
    '780 veces respecto del baseline manual medido en la Sección 5.1.4 (49,13 s → 0,063 s), con '
    'un intervalo de confianza al 95 % de 686× a 875×. Ese intervalo se obtiene por el teorema de '
    'Fieller, que es el procedimiento exacto para el cociente de dos medias y propaga la '
    'incertidumbre de ambos términos. Se consigna, por transparencia del cálculo, que la '
    'propagación simplificada —tratar el denominador automatizado como constante y arrastrar solo '
    'la incertidumbre del baseline— arroja 687× a 872×: la diferencia entre ambos procedimientos '
    'es inferior al 0,4 % porque el desvío del término automatizado (0,005 s sobre n = 50) es '
    'cuatro órdenes de magnitud menor que el del término manual.')
print('REC 3b declaracion del metodo en 5.4 ........... %d' % n)

# ---------------------------------------------------------------- REC 4 (terminologia)
# "precision" pasa a ser la metrica por clase; la global es "exactitud (accuracy)".
for viejo, nuevo in [
    ('H2b: Precisión clasificación ≥ 85%', 'H2b: Exactitud (accuracy) de clasificación ≥ 85 %'),
    ('Criterio de evaluación de precisión', 'Criterio de evaluación de la exactitud'),
    ('Chatbot: TMR < 10 s y precisión ≥ 85%', 'Chatbot: TMR < 10 s y exactitud (accuracy) ≥ 85 %'),
]:
    n = replace_everywhere(d, viejo, nuevo)
    if n:
        print('REC 4  "%s" ... %d' % (viejo[:42], n))
n = replace_in_paragraph(d.paragraphs[idx('H2b: El chatbot clasifica correctamente')],
    'Este umbral se fija como criterio del equipo',
    'A lo largo del trabajo se reserva «exactitud (accuracy)» para esta métrica global y '
    '«precisión» para la métrica por clase de la Tabla 5.9, conforme la terminología de Jurafsky '
    'y Martin (2024). Este umbral se fija como criterio del equipo')
print('REC 4  distincion exactitud/precision declarada en 1.4.2 ... %d' % n)

# ---------------------------------------------------------------- REC 5 (canal simulado)
n = replace_everywhere(d, 'Envía la respuesta por email (simulación del canal WhatsApp)',
    'Envía la respuesta por email: canal simulado en formato WhatsApp Cloud API con entrega SMTP')
print('REC 5a Tabla 4.7 nodo 17 ....................... %d' % n)
n = replace_everywhere(d, 'A diferencia del canal WhatsApp (envío simulado a un capturador SMTP local)',
    'A diferencia del canal simulado —que adopta el formato de la WhatsApp Cloud API en la '
    'recepción y entrega por SMTP a un capturador local, sin tocar la red de WhatsApp en ningún '
    'extremo—')
print('REC 5b 5.2.1 ................................... %d' % n)
n = replace_everywhere(d, '5.2 Resultados del Flujo 2 — Chatbot (canales WhatsApp y Telegram)',
    '5.2 Resultados del Flujo 2 — Chatbot (canal simulado formato WhatsApp y canal Telegram real)')
print('REC 5c encabezado 5.2 .......................... %d' % n)

# ---------------------------------------------------------------- REC 6 (nodos 5 y 6)
n = replace_in_paragraph(d.paragraphs[idx('Se eligió GPT-4o-mini por su balance')],
    'El cableado efectivo de ese contexto queda como línea futura en el Capítulo 7.',
    'El cableado efectivo de ese contexto queda como línea futura en el Capítulo 7. De ello se '
    'sigue una salvedad sobre el recuento de nodos que se consigna aquí para que no induzca a '
    'error: los nodos 5 y 6 del Flujo 2 —Buscar FAQ en PostgreSQL y Preparar Contexto FAQ— '
    'integran los dieciocho nodos funcionales del workflow, pero no inciden en la salida del '
    'sistema en la configuración efectivamente medida. El recuento describe el workflow tal como '
    'está construido, no la cadena causal que produce la respuesta.')
print('REC 6  salvedad de los nodos 5 y 6 ............. %d' % n)

# ---------------------------------------------------------------- residuo propio
# P389 decia "ampliar los ejemplos de few-shot", pero no hay ninguno: es
# residuo de la afirmacion de few-shot que se dio de baja en la ronda anterior.
n = replace_everywhere(d, 'o en ampliar los ejemplos de few-shot prompting para esta clase',
                          'o en incorporar al prompt ejemplos etiquetados de esta clase '
                          '(few-shot), que hoy no tiene ninguno')
print('EXTRA  residuo de few-shot en 5.4.1 (a-bis) .... %d' % n)

d.save(RUTA)
print()
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
TXT = '\n'.join(p.text for p in d2.paragraphs) + '\n' + '\n'.join(
    c.text for t in d2.tables for r in t.rows for c in r.cells)
print()
print('=== control ===')
for rot, pat, esperado in [
    ('"nueve con stock"',              'nueve con stock',                    0),
    ('"reduce significativamente"',    'reduce significativamente',          0),
    ('"reducción significativa"',      'reducción significativa',            0),
    ('"Tasa de resolución automática"','Tasa de resolución automática',      0),
    ('"valor obtenido en el Anexo H"', 'valor obtenido en el Anexo H',       0),
    ('"3 canales"',                    '3 canales',                          0),
    ('"IC 95 % del factor: 687"',      'IC 95 % del factor: 687',            0),
    ('U de Mann-Whitney',              'Mann-Whitney',                       2),
    ('t de Welch',                     'de Welch',                           2),
    ('40,8 % en Tabla 6.1',            '40,8 %',                             None),
    ('Fieller',                        'Fieller',                            None),
]:
    n = TXT.count(pat)
    ok = 'OK' if (esperado is None and n > 0) or n == esperado else '<-- REVISAR'
    print('  %-32s %d  %s' % (rot, n, ok))
