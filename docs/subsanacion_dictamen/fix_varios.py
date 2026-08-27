# -*- coding: utf-8 -*-
"""Hallazgos #9, #10, #12, #15, #18, #19, #20, #24, #25, #26."""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
P = d.paragraphs


def set_par(i, texto, estilo=None):
    p = P[i]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(texto)
    if estilo:
        try:
            p.style = estilo
        except KeyError:
            pass


# ============ #12 — el canal WhatsApp nunca usó sandbox ============
set_par(107,
    'WhatsApp Business API: requiere una cuenta verificada por Meta, verificación que no se '
    'completó dentro del alcance de este trabajo. En consecuencia el canal de WhatsApp no se '
    'ejecutó contra la API real ni contra su entorno sandbox: las pruebas del Flujo 2 sobre ese '
    'canal se realizaron entregando la respuesta a un capturador SMTP local, de modo que las '
    'latencias reportadas para él no incluyen componente alguno de la red de WhatsApp. La '
    'validación sobre un canal de mensajería real se resolvió mediante Telegram (Sección 5.2.1).')

# ============ #10 — la soberanía de datos no cubre la inferencia ============
replace_in_paragraph(P[118],
    '(2) soberanía de datos: al ejecutarse en infraestructura propia, no se transmiten datos '
    'sensibles de clientes a servidores de terceros;',
    '(2) soberanía de datos: al ejecutarse en infraestructura propia, la persistencia de los '
    'datos de clientes y la totalidad del procesamiento del Flujo 1 permanecen bajo control del '
    'operador, sin transmitirse a servidores de terceros. Corresponde precisar el alcance de '
    'este criterio, que no se extiende a la inferencia del Flujo 2: ese flujo remite el texto '
    'del mensaje del cliente a la API de OpenAI en cada interacción, dependencia que se discute '
    'como limitación en la Sección 5.4.1 y que motiva la recomendación de evaluar modelos de '
    'lenguaje de ejecución local del Capítulo 7;')

# ============ #9 — omnicanalidad teorizada pero no implementada ============
set_par(134, '2.3 Comunicación multicanal en e-commerce', 'Heading 2')
set_par(135, '2.3.1 Multicanalidad y omnicanalidad: delimitación conceptual', 'Heading 3')

set_par(136,
    'Conviene distinguir con precisión dos estrategias que suelen emplearse como sinónimos y no '
    'lo son. La estrategia multicanal consiste en poner a disposición del cliente varios canales '
    'de comunicación —WhatsApp, Telegram, correo electrónico, redes sociales— atendidos por un '
    'mismo sistema de procesamiento, pero donde cada conversación se resuelve de forma '
    'independiente y no se conserva contexto entre canales. La estrategia omnicanal (omnichannel) '
    'va un paso más allá: integra esos canales en una experiencia unificada, de manera que el '
    'historial y el contexto del cliente se preserven con independencia del canal utilizado '
    '(Verhoef et al., 2015), en oposición al enfoque en el que cada canal opera en silos. La '
    'diferencia operativa entre ambas es concreta y verificable: la omnicanalidad exige identidad '
    'unificada de cliente entre canales y recuperación del historial conversacional previo.')

set_par(137,
    'Corresponde declarar con precisión el alcance de lo implementado en este trabajo. El sistema '
    'desarrollado es multicanal: recibe y responde por WhatsApp y por Telegram a través de un '
    'mismo flujo de procesamiento, con un único clasificador y una única base de datos. Sin '
    'embargo, la tabla interactions no registra identificador de conversación ni identidad '
    'unificada de cliente entre canales, y ningún nodo del Flujo 2 recupera interacciones previas '
    'al construir la respuesta: cada mensaje se atiende sin memoria del anterior. La '
    'omnicanalidad en sentido estricto —identidad unificada y memoria conversacional entre '
    'canales— no fue implementada, y se declara como línea futura en el Capítulo 7. En el '
    'segmento de PyMEs de e-commerce argentinas, donde la atención suele distribuirse entre '
    'WhatsApp y correo electrónico, esa evolución constituye el paso natural del prototipo.')

# ============ #4-estado del arte (atenuación honesta) ============
replace_in_paragraph(P[152],
    'exploratorio porque no existen benchmarks publicados para n8n en el contexto de PyMEs argentinas',
    'exploratorio porque la revisión de antecedentes de la Sección 2.5 no identificó benchmarks '
    'publicados de n8n aplicados al ciclo post-venta de PyMEs argentinas')

# ============ #15 — categorías del catálogo ============
replace_in_paragraph(P[153],
    '(Notebooks, Periféricos, Monitores, Audio, Almacenamiento, Tablets, Redes, Accesorios, '
    'Componentes, Sillas y Conectividad)',
    '(Notebooks, Periféricos, Monitores, Audio, Almacenamiento, Tablets, Redes, Accesorios, '
    'Componentes, Mobiliario e Impresoras)')

# ============ #3 — el tamaño de muestra ahora describe los dos ensayos ============
set_par(167,
    'Flujo 1 — Corridas de medición: se ejecutaron dos ensayos con propósitos distintos. El '
    'primero envió 50 órdenes secuenciales espaciadas 2 segundos, tamaño que provee las '
    'estimaciones de MTTD, MTTR y tiempo end-to-end con un intervalo de confianza acotado. El '
    'segundo disparó 20 solicitudes simultáneas contra un stock deliberadamente insuficiente, '
    'repetidas durante 6 rondas (120 órdenes), lo que representa un volumen de concurrencia '
    'razonable para una PyME de escala pequeña. Este segundo tamaño no pretende representatividad '
    'estadística sino verificar la integridad transaccional del sistema bajo concurrencia, '
    'condición de diseño crítica para cualquier sistema de procesamiento de órdenes.')

# ============ #19 — contradicción media / mediana ============
set_par(191,
    'Mitigación: Todas las pruebas se realizaron en una ventana de tiempo acotada para minimizar '
    'la variabilidad temporal del servicio externo. Los tiempos se reportan como promedios '
    'aritméticos acompañados de su desvío estándar, criterio uniforme con el establecido en la '
    'Sección 3.5.4; la mediana se informa además de forma complementaria en los casos en que la '
    'distribución resulta asimétrica.')

# ============ #15 y #24 — conteo de vistas, estados y el "log de eventos" inexistente ============
replace_in_paragraph(P[203], 'contiene 7 tablas, 5 vistas de métricas',
                     'contiene 7 tablas, 6 vistas de métricas')

set_par(205,
    'Estados admitidos por la restricción de integridad del campo status: pending, processing, '
    'confirmed, no_stock, error, cancelled, shipped y delivered.')

set_par(206,
    'Nota: stock_verified y stock_updated no son estados del campo status sino pasos internos del '
    'pipeline, observables en el historial de ejecuciones que el propio motor de flujos conserva. '
    'El esquema de este trabajo no incorpora una tabla de bitácora de eventos: el campo status '
    'refleja únicamente los estados de negocio visibles al cliente y al operador, y las marcas '
    'temporales received_at, processed_at y notified_at registran los instantes del procesamiento.',
    'Block Text')

# ============ #18 — received/notified no son estados ============
set_par(243,
    'El MTTD se calcula sobre el intervalo received_at → processed_at y el MTTR sobre el '
    'intervalo processed_at → notified_at; ambas son diferencias entre marcas temporales y no '
    'entre estados, conforme la distinción establecida en el párrafo anterior. En el caso de '
    'no_stock, el MTTR refleja el tiempo transcurrido hasta el envío de la notificación de '
    'rechazo.')

# ============ #20 — Liu et al. (2023) es un survey de prompting ============
replace_in_paragraph(P[333],
    'lo que es consistente con los resultados reportados por Liu et',
    'resultado que se apoya en las técnicas sistematizadas por Liu et')

d.save(RUTA)

# ===== Verificación =====
d2 = Document(RUTA)
for i in (107, 118, 134, 135, 136, 137, 152, 153, 167, 191, 203, 205, 206, 243, 333):
    print('===== P%d [%s] =====' % (i, d2.paragraphs[i].style.name))
    print(d2.paragraphs[i].text[:700])
    print()
