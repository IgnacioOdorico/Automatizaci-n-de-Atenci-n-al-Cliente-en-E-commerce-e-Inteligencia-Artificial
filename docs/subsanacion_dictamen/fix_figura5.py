# -*- coding: utf-8 -*-
"""HALLAZGO PROPIO — la Figura 5 no representaba el workflow implementado.

El diagrama decía "Chatbot Omnicanal", "~21 nodos", y dibujaba nodos que NO
existen en el flujo (Merge Canales, ¿Urgente?, Notificar Admin) además de tres
"Normalizar" separados cuando el real tiene uno solo. Faltaban Router Canal,
Enviar Telegram, Enviar Respuesta y Registrar Interacción.

Se reemplaza por un diagrama generado a partir de las conexiones reales del
archivo workflows/Flujo 2 — Chatbot WhatsApp + Telegram.json (18 nodos).
El marco de la imagen en el documento se reajusta para no deformarla.
"""
import zipfile
import shutil
import re
import io
import os

DOCX = 'docs/TESIS_FINAL_UTN_v6.docx'
NUEVA = 'docs/figuras_v6/diagrama_flujo2.png'
TMP = DOCX + '.tmp'

shutil.copy(DOCX, DOCX + '.bak-fig5')

from PIL import Image
w, h = Image.open(NUEVA).size
ratio = w / h
blob = open(NUEVA, 'rb').read()

zin = zipfile.ZipFile(DOCX, 'r')
zout = zipfile.ZipFile(TMP, 'w', zipfile.ZIP_DEFLATED)

cambios = {'imagen': 0, 'extent': 0}
for item in zin.infolist():
    data = zin.read(item.filename)

    if item.filename == 'word/media/image5.png':
        data = blob
        cambios['imagen'] += 1

    elif item.filename == 'word/document.xml':
        txt = data.decode('utf-8')
        # ubicar el bloque <w:drawing> que referencia a image5 vía su relación
        rels = zin.read('word/_rels/document.xml.rels').decode('utf-8')
        m = re.search(r'Id="(rId\d+)"[^>]*Target="media/image5\.png"', rels)
        rid = m.group(1)

        def ajustar(bloque):
            e = re.search(r'cx="(\d+)" cy="(\d+)"', bloque)
            if not e:
                return bloque
            cx = int(e.group(1))
            cy_nuevo = int(round(cx / ratio))
            cambios['extent'] += 1
            return re.sub(r'cx="\d+" cy="\d+"', 'cx="%d" cy="%d"' % (cx, cy_nuevo), bloque)

        # cada <w:drawing>...</w:drawing> que contenga r:embed="rid"
        partes = re.split(r'(<w:drawing>.*?</w:drawing>)', txt, flags=re.S)
        for i, parte in enumerate(partes):
            if parte.startswith('<w:drawing>') and ('r:embed="%s"' % rid) in parte:
                partes[i] = ajustar(parte)
        data = ''.join(partes).encode('utf-8')

    zout.writestr(item, data)

zin.close()
zout.close()
os.replace(TMP, DOCX)

print('imagen reemplazada : %d' % cambios['imagen'])
print('extents ajustados  : %d' % cambios['extent'])

# --- verificación ---
z = zipfile.ZipFile(DOCX)
im = Image.open(io.BytesIO(z.read('word/media/image5.png')))
print('image5 ahora       : %s' % (im.size,))
from docx import Document
d = Document(DOCX)
for i, p in enumerate(d.paragraphs):
    if p.text.strip().startswith('Figura 5:'):
        e = re.search(r'<wp:extent cx="(\d+)" cy="(\d+)"', d.paragraphs[i-1]._element.xml)
        cx, cy = int(e.group(1)), int(e.group(2))
        print('extent             : %.2f x %.2f cm  ratio %.4f (imagen %.4f)'
              % (cx/360000, cy/360000, cx/cy, ratio))
        print('epígrafe           : %s' % p.text.strip()[:120])
        break
print('parrafos %d | tablas %d' % (len(d.paragraphs), len(d.tables)))
