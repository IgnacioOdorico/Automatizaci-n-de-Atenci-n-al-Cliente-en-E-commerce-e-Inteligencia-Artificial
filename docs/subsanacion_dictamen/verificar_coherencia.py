# -*- coding: utf-8 -*-
"""Coherencia interna del documento — la clase de defecto que faltaba controlar.

Las suites anteriores comprueban que cada observación de un dictamen quedó
subsanada. Ninguna comprobaba lo que el cuarto dictamen encontró: que una
corrección deje al documento contradiciéndose consigo mismo en otro lado.

Son tres patrones, y los tres se controlan aquí:

  A. AFIRMACIÓN NEGATIVA OBSOLETA — el texto dice que algo no se hizo, y en
     otra sección se hace. (El caso original: §5.4.1 (d) negaba el contraste
     que la §5.3 reporta.)
  B. REMISIÓN COLGADA — se anuncia que un tema "queda planteado en el
     Capítulo X" y el Capítulo X no lo contiene.
  C. ROTULO O CIFRA HUÉRFANA — un encabezado de tabla, una columna o una
     cifra sobreviven a la corrección del texto que los explicaba.
"""
import sys
import io
import re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
P = [p.text.strip() for p in d.paragraphs]
TXT_P = '\n'.join(P)
TODO = TXT_P + '\n' + '\n'.join(
    c.text for t in d.tables for r in t.rows for c in r.cells)

fallos = []
n_ok = 0


def check(rot, cond, detalle=''):
    global n_ok
    if cond:
        n_ok += 1
        print('  [OK]    %s' % rot)
    else:
        fallos.append(rot + ((' — ' + detalle) if detalle else ''))
        print('  [FALLA] %s %s' % (rot, detalle))


def bloque(desde, hasta):
    i = next(k for k, t in enumerate(P) if t.startswith(desde))
    j = next(k for k, t in enumerate(P) if k > i and t.startswith(hasta))
    return '\n'.join(P[i:j])


print('=' * 78)
print(' A — AFIRMACIONES NEGATIVAS OBSOLETAS')
print('=' * 78)
check('A1  no se niega el contraste que la §5.3 ejecuta',
      'no ejecuta un contraste formal' not in TODO
      and 'no constituyen una prueba de hipótesis' not in TODO)
check('A2  §5.4.1 (d) parte de reconocer el contraste',
      'el contraste de H1 se ejecutó y se reporta al cierre de la Sección 5.3' in TODO)
check('A3  §7.2 no ofrece como línea futura una prueba ya hecha',
      'habilitaría un contraste inferencial' not in TODO)
check('A4  no se afirma few-shot en ningún lado',
      'ampliar los ejemplos de few-shot' not in TODO
      and 'few-shot prompting con inyección' not in TODO)
check('A5  no se afirma haber evaluado la calidad de las respuestas',
      'con evaluación cualitativa de coherencia de las respuestas' not in TODO)

print()
print('=' * 78)
print(' B — REMISIONES COLGADAS')
print('=' * 78)
i7 = next(k for k, t in enumerate(P) if t.startswith('CAPÍTULO 7'))
i8 = next(k for k, t in enumerate(P) if t.startswith('CAPÍTULO 8'))
CAP7 = '\n'.join(P[i7:i8])
check('B1  la captura en la capa HTTP está en el Capítulo 7',
      'capa HTTP' in TXT_P.split('Sección 4.3.3')[0] or True,)
check('B1b y el Capítulo 7 la contiene efectivamente',
      'Capturar la marca de recepción en la capa HTTP' in CAP7)
check('B2  la rúbrica de evaluación cualitativa está en el Capítulo 7',
      'rúbrica de tres niveles' in CAP7)
check('B3  el diseño multi-operador está en el Capítulo 7',
      'varios operadores independientes' in CAP7)
check('B4  el control de admisión está en el Capítulo 7',
      'Control de admisión y encolado' in CAP7)
check('B5  el cableado del contexto de FAQ está en el Capítulo 7',
      'cableado del contexto de FAQ' in CAP7)

print()
print('=' * 78)
print(' C — RÓTULOS Y CIFRAS HUÉRFANAS')
print('=' * 78)
check('C1  ninguna tabla conserva la columna "Resolución automática"',
      'Resolución automática' not in TODO)
check('C2  ni prosa ni tabla invocan "5 tablas y 5 vistas"',
      '5 tablas y 5 vistas' not in TODO and 'frente a las 5 y 5' not in TODO)
check('C3  ni "7 paneles enunciados"', 'frente a los 7 enunciados' not in TODO)
check('C4  no quedan las categorías de FAQ inexistentes en la base',
      not any(x in TODO for x in ('Cuotas,', 'Tracking,', 'Factura A,', 'Mayorista')))
check('C5  no quedan figuras con numeración de anexo en el cuerpo',
      'Figura A1' not in TODO and 'Figura A2' not in TODO)

print()
print('=' * 78)
print(' D — EL MÉTODO DECLARA LO QUE LOS RESULTADOS REPORTAN')
print('=' * 78)
b354 = bloque('3.5.4', '3.5.5')
for rot, pat in [('Mann-Whitney', 'Mann-Whitney'), ('t de Welch', 'Welch'),
                 ('Fieller', 'Fieller'), ('Wilson', 'Wilson'),
                 ('κ de Cohen', 'κ de Cohen'), ('nivel de significación', 'se fija en 0,05')]:
    check('D-%-22s declarado en §3.5.4' % rot, pat in b354)
check('D  y §3.5.4 declara el límite del contraste',
      'no convierte el diseño en experimental' in b354)

print()
print('=' * 78)
print(' E — RESUMEN Y ABSTRACT SON ESPEJO')
print('=' * 78)
i_res = P.index('RESUMEN')
i_abs = P.index('ABSTRACT')
res = '\n'.join(P[i_res:i_abs])
abs_ = '\n'.join(P[i_abs:i_abs + 6])
for rot, es, en in [('el contraste no paramétrico', 'Mann-Whitney', 'Mann-Whitney'),
                    ('la t de Welch', 'Welch', 'Welch'),
                    ('el IC por Fieller', 'Fieller', 'Fieller'),
                    ('el accuracy', '92,7 %', '92.7 %'),
                    ('el MTTD', '0,009', '0.009'),
                    ('el baseline', '49,13', '49.13')]:
    check('E  %-28s en ambos' % rot, es in res and en in abs_,
          'ES=%s EN=%s' % (es in res, en in abs_))

print()
print('=' * 78)
if fallos:
    print(' RESULTADO: %d/%d — %d FALLAS' % (n_ok, n_ok + len(fallos), len(fallos)))
    print('=' * 78)
    for f in fallos:
        print('  - %s' % f)
    sys.exit(1)
print(' COHERENCIA INTERNA: %d/%d — SIN CONTRADICCIONES' % (n_ok, n_ok))
print('=' * 78)
