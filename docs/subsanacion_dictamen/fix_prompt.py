# -*- coding: utf-8 -*-
"""HALLAZGO PROPIO — el documento afirma few-shot prompting e inyección de FAQ.

Verificado contra workflows/Flujo 2 — Chatbot WhatsApp + Telegram.json:
  * El prompt de sistema (1032 caracteres) NO contiene ningún ejemplo etiquetado.
  * El nodo "Preparar Contexto FAQ" construye faq_context, pero esa variable
    aparece UNA sola vez en todo el workflow: en el nodo que la crea. Nunca se
    referencia desde el prompt del nodo "IA - Motor Decision".
  => El 92,7 % se obtuvo con prompting ZERO-SHOT por instrucción.
"""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
P = d.paragraphs


def set_par(i, texto, estilo=None):
    p = P[i]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(texto)
    if estilo:
        try:
            p.style = estilo
        except KeyError:
            pass


# ===== §2.2.2 — marco teórico =====
set_par(130,
    'En este proyecto se utiliza GPT-4o-mini de OpenAI (OpenAI, 2024), un modelo optimizado para '
    'inferencia rápida con buen balance entre capacidad y costo. La interacción con el modelo se '
    'configura mediante prompt engineering estructurado: el prompt de sistema define el rol del '
    'modelo (agente de atención al cliente de un e-commerce), las categorías de intención válidas '
    '(FAQ, ESTADO_PEDIDO, RECLAMO y GENERAL), el formato de respuesta esperado (un objeto JSON '
    'estricto, sin texto ni marcado alrededor) y las reglas de idioma y de tono.')

set_par(131,
    'Corresponde precisar el régimen de prompting empleado, porque la distinción tiene '
    'consecuencias sobre la interpretación de los resultados. Brown et al. (2020) diferencian '
    'tres regímenes según la cantidad de ejemplos etiquetados que se incluyen en el prompt: '
    'zero-shot, donde el modelo recibe únicamente la instrucción y la definición de las clases; '
    'one-shot, con un ejemplo; y few-shot, con varios ejemplos de entrada y salida. Liu et al. '
    '(2023) sistematizan el conjunto de estas técnicas. El clasificador de este trabajo opera en '
    'régimen zero-shot: el prompt de sistema —que se transcribe íntegramente en el Anexo H— '
    'enumera las cuatro categorías y fija el formato de salida, pero no incluye ningún ejemplo '
    'etiquetado. La elección no fue deliberada en su origen, sino que se constató al auditar la '
    'implementación; se la reporta como tal porque condiciona la lectura del resultado: el '
    'accuracy obtenido corresponde al desempeño del modelo sin ejemplos ni ajuste fino, lo que '
    'constituye un piso y no un techo para la tarea.')

# ===== §4.4.3 — descripción del prompt =====
set_par(264,
    'El nodo de GPT-4o-mini recibe un prompt de sistema que define la identidad del asistente, la '
    'tarea de clasificación, las cuatro categorías de intención admitidas (FAQ, ESTADO_PEDIDO, '
    'RECLAMO y GENERAL) y el formato estricto de la respuesta: un objeto JSON con los campos '
    'intent, order_id, urgente y respuesta. El texto completo se transcribe en el Anexo H, de '
    'modo que el resultado de clasificación del Capítulo 5 sea auditable.')

set_par(265,
    'Se eligió GPT-4o-mini por su balance entre costo y capacidad: para la clasificación de '
    'intenciones en un dominio acotado y la generación de respuestas de soporte no se requiere la '
    'potencia completa de GPT-4o. El prompt no incorpora ejemplos etiquetados, de modo que opera '
    'en régimen zero-shot (véase la Sección 2.2.2). Cabe señalar una diferencia entre el diseño y '
    'la implementación que se detectó al auditar el workflow y que se declara antes que '
    'corregirse, porque corregirla invalidaría la medición ya realizada: el flujo incluye un nodo '
    'que recupera las FAQ de la base de datos y arma un bloque de contexto, pero esa variable no '
    'llega a referenciarse desde el prompt del nodo de inferencia. En consecuencia, el accuracy '
    'del 92,7 % reportado en el Capítulo 5 se obtuvo sin inyección de FAQ en el prompt. El '
    'cableado efectivo de ese contexto queda como línea futura en el Capítulo 7.')

# ===== §5.4 — discusión =====
replace_in_paragraph(P[348],
    'La estrategia de few-shot prompting con inyección dinámica de FAQ permitió alcanzar un 92,7% '
    'de precisión sin requerir fine-tuning del modelo, resultado que se apoya en las técnicas '
    'sistematizadas por Liu et al. (2023) para tareas de clasificación con LLMs en dominios '
    'especializados.',
    'El dato relevante es que ese 92,7 % de precisión se alcanzó en régimen zero-shot: sin ajuste '
    'fino, sin ejemplos etiquetados en el prompt y sin inyección de la base de FAQ, con un prompt '
    'de sistema de poco más de mil caracteres que se limita a enumerar las categorías y a fijar '
    'el formato de salida (Anexo H). El resultado debe leerse por lo tanto como un piso del '
    'desempeño alcanzable en la tarea, y no como su techo: las técnicas de prompting '
    'sistematizadas por Liu et al. (2023), así como el cableado del contexto de FAQ ya construido '
    'en el flujo, constituyen margen de mejora todavía sin explotar.')

# ===== §6.3 — conclusiones sustantivas =====
set_par(365,
    'Sobre la efectividad de GPT-4o-mini en atención al cliente: el modelo alcanzó un accuracy del '
    '92,7 % en la clasificación de intenciones sin ajuste fino y en régimen zero-shot, es decir '
    'guiado exclusivamente por una instrucción que enumera las cuatro categorías y fija el formato '
    'de salida, sin ejemplos etiquetados. Que un modelo de bajo costo alcance ese desempeño con un '
    'prompt de mil caracteres es, a juicio de los autores, el hallazgo más transferible del '
    'trabajo para una PyME: reduce la barrera de entrada de la clasificación automática de '
    'consultas a la redacción cuidadosa de una instrucción. El margen de mejora disponible —'
    'incorporar ejemplos al prompt y cablear el contexto de FAQ que el flujo ya recupera— queda '
    'documentado como línea futura.')

d.save(RUTA)

d2 = Document(RUTA)
for i in (130, 131, 264, 265, 348, 365):
    print('===== P%d =====' % i)
    print(d2.paragraphs[i].text)
    print()
buscar(d2, r'few.?shot', 'menciones de few-shot restantes')
