# -*- coding: utf-8 -*-
"""Auditoria de 2.a instancia (8,2/10, aprobado) — BLOQUE 1 del plan de correccion.

Cubre A-01, A-02, M-01, M-02 y los defectos de edicion B-01 a B-04.

Verificado antes de escribir, sobre los archivos versionados:
  A-01  Tabla 4.10, prueba C4: el resultado esperado incluye "notificacion admin".
        Recorridos los 18 nodos de la Tabla 4.7: la rama de RECLAMO va de Switch
        Intent a Crear Ticket, de ahi a Preparar Respuesta Ticket y al Router
        Canal. No hay segundo destino. El auditor tiene razon.
  A-02  Cotejadas las consultas de grafana/dashboards/*.json contra las vistas de
        init_simple.sql. De los 13 paneles, 10 se restringen a los registros
        medidos, 1 es constante y 2 NO se restringen:
          - "Ordenes totales"          -> v_metrics_summary   (SELECT COUNT(*) FROM orders)
          - "Ordenes procesadas por dia" -> v_daily_order_summary (FROM orders, sin WHERE)
        La afirmacion universal de §4.5 es falsa. Confirmado.
  M-01  "disponibilidad continua" aparece en el resumen, el abstract y §6.1.
  M-02  Ningun nodo de la Tabla 4.5 escribe order_items; el comentario del DDL
        lo admite pero el cuerpo del Capitulo 4 no.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document
from docx.oxml.ns import qn

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def par(pref):
    return d.paragraphs[idx(pref)]


def reescribir(pref, texto):
    p = par(pref)
    assert 'blip' not in p._element.xml, 'lleva imagen: no reescribir sus runs'
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ''
    return p


def tabla_de(epigrafe):
    """Devuelve la tabla que sigue al parrafo del epigrafe dado."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    ti = 0
    ult = None
    for ch in d.element.body.iterchildren():
        tag = ch.tag.split('}')[1]
        if tag == 'p':
            t = Paragraph(ch, d).text.strip()
            if t.startswith(epigrafe + ':'):
                ult = True
        elif tag == 'tbl':
            if ult:
                return d.tables[ti]
            ti += 1
    raise KeyError(epigrafe)


print('=' * 76)
print(' A-01 — la prueba C4 declara una notificacion que ningun nodo implementa')
print('=' * 76)
n = replace_everywhere(d, 'Ticket + notificación admin + is_urgent=true',
                          'Ticket creado + is_urgent=true + respuesta al cliente por el canal de origen')
print('  resultado esperado de C4 corregido: %d' % n)
n = replace_everywhere(d, 'notificación admin', 'notificación al administrador')
print('  residuos de "notificación admin": %d' % n)

print()
print('=' * 76)
print(' A-02 — §4.5 afirma un filtrado universal que las vistas no sostienen')
print('=' * 76)
n = replace_in_paragraph(par('Se configuraron dos dashboards en Grafana'),
    'Las consultas de todos los paneles se restringen a los registros medidos '
    "(data_source = 'measured'), de modo que los valores mostrados corresponden a las mediciones "
    'de los experimentos y no a los datos de carga inicial.',
    'El filtrado por procedencia no es uniforme entre paneles, y corresponde declarar qué panel '
    'aplica cuál criterio, porque de ello depende cómo se lee cada tablero. Diez de los trece '
    "paneles se restringen a los registros medidos (data_source = 'measured'): cuatro del tablero "
    'del Flujo 1 —los tres indicadores de tiempo y la distribución de estados— lo hacen mediante '
    'una cláusula explícita en la consulta del propio panel, y seis del tablero del Flujo 2 lo '
    'heredan de la vista v_chatbot_corpus, que además acota la población a la ventana temporal de '
    'la corrida evaluada. El panel de exactitud es un valor constante, según se declara en la '
    'Tabla 4.8. Los dos paneles restantes, ambos del tablero del Flujo 1, no aplican ese filtro: '
    '«Órdenes totales» consume la vista v_metrics_summary y «Órdenes procesadas por día» la vista '
    'v_daily_order_summary, y ninguna de las dos restringe por procedencia —sus definiciones, '
    'transcriptas en el Anexo A, agregan sobre la totalidad de la tabla de órdenes—. En '
    'consecuencia, esos dos paneles incluyen también las órdenes de carga inicial. La '
    'consecuencia sobre los resultados de este trabajo es nula, y conviene decir por qué: ninguna '
    'cifra del Capítulo 5 proviene de esas dos vistas. Las métricas reportadas se calculan sobre '
    'las corridas identificadas por su propio prefijo de número de orden, según se detalla en la '
    'Sección 3.5.1. Los dos paneles cumplen una función de contexto operativo y no de evidencia. '
    'Homogeneizar el filtro en las cuatro vistas que hoy no lo aplican queda planteado en el '
    'Capítulo 7.')
print('  §4.5 reescrita con el filtrado real: %d' % n)

n = replace_everywhere(d,
    'Figura 7: Dashboard “Pipeline Post-Venta” en Grafana, alimentado desde las vistas de '
    'PostgreSQL con los datos medidos del Flujo 1.',
    'Figura 7: Dashboard “Pipeline Post-Venta” en Grafana, alimentado desde las vistas de '
    'PostgreSQL con los datos medidos del Flujo 1. Los paneles de indicadores y de distribución '
    'de estados se restringen a los registros medidos; los de «Órdenes totales» y «Órdenes '
    'procesadas por día» no lo hacen, de modo que el primer punto de la serie diaria corresponde '
    'a las órdenes de carga inicial y no a una corrida experimental (Sección 4.5).')
print('  epígrafe de la Figura 7 con la salvedad: %d' % n)

# la homogeneizacion del filtro pasa a ser recomendacion de produccion
p_http = par('Capturar la marca de recepción en la capa HTTP')
insert_paragraph_after(p_http,
    'Homogeneizar el filtro de procedencia en las vistas de métricas: cuatro de las seis vistas '
    'agregan sobre la totalidad de la tabla de órdenes o de interacciones, sin restringir por '
    "data_source (Anexo A). Dos paneles del tablero del Flujo 1 consumen esas vistas y por lo "
    'tanto muestran también los datos de carga inicial (Sección 4.5). Incorporar la restricción a '
    'la definición de las vistas —y no a la consulta de cada panel— evita que la propiedad '
    'dependa de quién construya el tablero, y hace que la separación entre lo medido y lo '
    'precargado sea una garantía del esquema y no una convención de uso.',
    estilo=p_http.style.name)
print('  ítem de homogeneización agregado al Capítulo 7')

print()
print('=' * 76)
print(' M-01 — "disponibilidad continua" se afirma sin medicion que la sostenga')
print('=' * 76)
for viejo, nuevo, rot in [
    ('y de 3,07 segundos sobre 45 interacciones del canal Telegram real, con disponibilidad '
     'continua, y', 'y de 3,07 segundos sobre 45 interacciones del canal Telegram real, y', 'resumen'),
    ('and 3.07 seconds over 45 interactions on the real Telegram channel, with continuous '
     'availability, and', 'and 3.07 seconds over 45 interactions on the real Telegram channel, and',
     'abstract'),
    ('El TMR del chatbot es de 1,47 s en promedio sobre el corpus evaluado, con disponibilidad '
     'continua.',
     'El TMR del chatbot es de 1,47 s en promedio sobre el corpus evaluado. La disponibilidad '
     'ininterrumpida es una propiedad esperable de la arquitectura —el sistema no depende de un '
     'operador humano para responder— pero no fue medida: no se ejecutó ensayo de disponibilidad '
     'ni se observó el servicio durante una ventana prolongada, de modo que no se la reporta como '
     'resultado.', '§6.1'),
    ('En atención al cliente, el chatbot responde en 1,47 s en promedio sobre el corpus evaluado '
     'y en 3,07 s sobre el canal Telegram real, con disponibilidad continua.',
     'En atención al cliente, el chatbot responde en 1,47 s en promedio sobre el corpus evaluado '
     'y en 3,07 s sobre el canal Telegram real.', '§5.4'),
    ('Un pipeline automatizado garantiza confirmaciones inmediatas, notificaciones proactivas y '
     'respuestas 24/7 a través de un chatbot.',
     'Un pipeline automatizado permite confirmaciones inmediatas, notificaciones proactivas y '
     'respuestas sin dependencia del horario de un operador.', '§1.3'),
]:
    n = replace_everywhere(d, viejo, nuevo)
    print('  %-10s %d' % (rot, n))

print()
print('=' * 76)
print(' M-02 — order_items existe en el esquema y ningun nodo la escribe')
print('=' * 76)
p42 = par('4.2.1 Tablas principales')
# la salvedad va despues del parrafo introductorio de 4.2.1
i = idx('4.2.1 Tablas principales')
while not d.paragraphs[i + 1].text.strip():
    i += 1
insert_paragraph_after(d.paragraphs[i + 1],
    'Corresponde una salvedad sobre order_items, porque el esquema y el sistema implementado no '
    'coinciden en este punto. La tabla existe, tiene sus claves foráneas y sus índices, y '
    'normaliza correctamente la relación entre una orden y sus productos. Pero ninguno de los '
    'quince nodos del Flujo 1 (Tabla 4.5) inserta en ella: el pipeline implementado carga órdenes '
    'de un solo producto, y conserva a tal efecto los campos de producto y cantidad en la propia '
    'tabla de órdenes. El esquema resuelve por lo tanto el modelado de órdenes multiproducto, y '
    'el sistema todavía no lo ejercita. Se declara aquí y no solo en el comentario del script del '
    'Anexo A, porque es en el cuerpo del capítulo donde el lector lo busca.',
    estilo='Body Text')
print('  salvedad incorporada al cuerpo de §4.2.1')

t61 = None
for t in d.tables:
    if t.rows[0].cells[0].text.strip() == 'Obj.':
        t61 = t
        break
f3 = [i for i, r in enumerate(t61.rows) if r.cells[0].text.strip() == 'OE3'][0]
set_cell(t61, f3, 2, t61.rows[f3].cells[2].text.strip() +
         ' Se consigna que order_items está en el esquema pero ningún nodo del pipeline la '
         'escribe: el sistema implementado carga órdenes mono-producto (Sección 4.2.1).')
print('  salvedad incorporada a la celda de OE3')

print()
print('=' * 76)
print(' B-01 a B-04 — defectos de edicion')
print('=' * 76)

# B-01: los listados llevaban las descripciones cortadas a longitud fija.
import re as _re
epigrafes = {}
for p in d.paragraphs:
    m = _re.match(r'(Figura \d+):\s*(.+)', p.text.strip())
    if m:
        epigrafes[m.group(1)] = m.group(2).strip()
rep = 0
for t in d.tables:
    if t.rows[0].cells[0].text.strip() == 'Figura' and t.rows[0].cells[2].text.strip() == 'Sección':
        for i, r in enumerate(t.rows[1:], 1):
            rot = r.cells[0].text.strip()
            if rot in epigrafes and r.cells[1].text.strip() != epigrafes[rot]:
                set_cell(t, i, 1, epigrafes[rot])
                rep += 1
print('  B-01 descripciones repuestas desde los epígrafes reales: %d' % rep)

# B-02: la segunda fila del Listado de Tablas estaba marcada como encabezado repetido
quit_ = 0
for t in d.tables[:2]:
    for i, r in enumerate(t.rows):
        trPr = r._tr.find(qn('w:trPr'))
        if trPr is None:
            continue
        th = trPr.find(qn('w:tblHeader'))
        if th is not None and i > 0:
            trPr.remove(th)
            quit_ += 1
            print('  B-02 desmarcada como encabezado la fila %d (%r)' % (i, r.cells[0].text.strip()))
if not quit_:
    print('  B-02 (nada que desmarcar)')

# B-03: rotulos de la numeracion anterior en el Anexo G
n = replace_everywhere(d, 'Las figuras A1 y A2 corresponden', 'Las Figuras 10 y 11 corresponden')
print('  B-03 rótulos A1/A2 del Anexo G: %d' % n)

# B-04: los comentarios del DDL conservan identificadores de la auditoria previa
for viejo, nuevo in [
    ('-- Ítems de una orden: N productos por orden (A-12). orders.product_id y',
     '-- Ítems de una orden: normaliza la relación orden→productos. orders.product_id y'),
    ('-- Alertas de stock bajo (A-07): se registra cada vez que un producto',
     '-- Alertas de stock bajo: se registra cada vez que un producto'),
    ('--  ÍNDICES DE RENDIMIENTO (OE3 / A-12)', '--  ÍNDICES DE RENDIMIENTO'),
]:
    n = replace_everywhere(d, viejo, nuevo)
    print('  B-04 comentario del DDL: %d  (%s)' % (n, viejo[:46]))

d.save(RUTA)
print()
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
P2 = [p.text.strip() for p in d2.paragraphs]
TODO = '\n'.join(P2) + '\n' + '\n'.join(
    c.text for t in d2.tables for r in t.rows for c in r.cells)
fallos = []


def check(rot, cond, det=''):
    print('  [%s] %s %s' % ('OK   ' if cond else 'FALLA', rot, det))
    if not cond:
        fallos.append(rot)


print()
print('=== control ===')
check('A-01 sin "notificación admin" en la C4', 'notificación admin +' not in TODO)
check('A-01 el resultado esperado describe lo que existe',
      'Ticket creado + is_urgent=true + respuesta al cliente' in TODO)
check('A-02 sin la afirmación universal de filtrado',
      'Las consultas de todos los paneles se restringen' not in TODO)
check('A-02 §4.5 declara los dos paneles que no filtran',
      'v_metrics_summary' in TODO and 'v_daily_order_summary' in TODO
      and 'Diez de los trece paneles' in TODO)
check('A-02 el epígrafe de la Figura 7 lo explica',
      'el primer punto de la serie diaria corresponde' in TODO)
check('A-02 el Capítulo 7 recoge la homogeneización',
      'Homogeneizar el filtro de procedencia' in TODO)
check('M-01 sin "disponibilidad continua"', 'disponibilidad continua' not in TODO)
check('M-01 sin "continuous availability"', 'continuous availability' not in TODO)
check('M-01 §6.1 la degrada a propiedad esperable y no medida',
      'es una propiedad esperable de la arquitectura' in TODO and 'pero no fue medida' in TODO)
check('M-02 la salvedad está en el cuerpo de §4.2.1',
      'ninguno de los quince nodos del Flujo 1 (Tabla 4.5) inserta en ella' in TODO)
check('M-02 y en la celda de OE3',
      any('mono-producto' in ' '.join(c.text for c in r.cells)
          for t in d2.tables for r in t.rows if r.cells[0].text.strip() == 'OE3'))
trunc = []
for t in d2.tables:
    if t.rows[0].cells[0].text.strip() == 'Figura' and t.rows[0].cells[2].text.strip() == 'Sección':
        for r in t.rows[1:]:
            if not r.cells[1].text.strip().endswith(('.', ')')):
                trunc.append(r.cells[0].text.strip())
check('B-01 ninguna descripción truncada en los listados', not trunc, str(trunc))
enc = []
for t in d2.tables[:2]:
    for i, r in enumerate(t.rows):
        trPr = r._tr.find(qn('w:trPr'))
        if trPr is not None and trPr.find(qn('w:tblHeader')) is not None and i > 0:
            enc.append(i)
check('B-02 solo la fila 0 se repite como encabezado', not enc, str(enc))
check('B-03 sin rótulos A1/A2', 'figuras A1 y A2' not in TODO and 'Figura A1' not in TODO)
check('B-04 sin identificadores de auditoría en el DDL',
      not _re.search(r'\(A-\d\d\)|\(A-\d\d /|OE3 / A-12', TODO))

print()
if fallos:
    print('FALLAS: %d' % len(fallos))
    for f in fallos:
        print('  - %s' % f)
    sys.exit(1)
print('BLOQUE 1 — SIN FALLAS')
