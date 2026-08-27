# -*- coding: utf-8 -*-
"""§3.5.5 — el protocolo de medición del baseline manual (Bloqueante #1, recomendación 1)."""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)

BLOQUE = [
    ('3.5.5 Medición del baseline de atención manual', 'Heading 3'),

    ('La pregunta de investigación y el objetivo general están formulados en términos '
     'comparativos: se busca establecer en qué medida la automatización reduce los tiempos '
     'operativos respecto del proceso manual. Ese término de comparación exige una medición '
     'propia y no una estimación tomada de la experiencia informal del equipo. Por ese motivo se '
     'diseñó un experimento específico, destinado exclusivamente a cronometrar el procesamiento '
     'manual de una orden bajo condiciones equivalentes a las del pipeline automatizado.',
     'First Paragraph'),

    ('Tarea cronometrada. Un operador procesa una orden completa contra la misma base de datos y '
     'el mismo catálogo que utiliza el Flujo 1, sin asistencia del sistema automatizado. La tarea '
     'se descompone en tres fases que se miden por separado: T1, lectura de los datos de la orden '
     'y consulta del stock disponible del producto mediante una sentencia SQL; T2, descuento del '
     'stock y actualización del estado de la orden a confirmada, o bien marcado como sin stock '
     'cuando corresponde; y T3, redacción de la notificación al cliente completando los campos '
     'variables de una plantilla fija. El tiempo total por orden es la suma de las tres fases. El '
     'desglose importa tanto como el total, porque permite identificar dónde se concentra '
     'efectivamente el esfuerzo manual.', 'Body Text'),

    ('Decisiones de diseño y su sentido. Dos decisiones condicionan el valor obtenido y se '
     'declaran de manera explícita, porque ambas se resolvieron eligiendo deliberadamente el '
     'escenario que menos favorece a la hipótesis. La primera concierne a la redacción del '
     'correo: se optó por una plantilla fija a completar y no por la redacción desde cero, ya que '
     'la primera representa a una PyME con un proceso mínimo establecido, que es el caso realista. '
     'Esta elección reduce el tiempo del baseline y, en consecuencia, reduce el factor de mejora '
     'que el trabajo podrá reportar. La segunda concierne al alcance del cronómetro, que cubre '
     'únicamente el procesamiento: desde que el operador toma conocimiento de la orden hasta que '
     'termina de redactar la notificación. La latencia de detección, es decir cada cuánto un '
     'operador revisa efectivamente la bandeja de entrada, no se incluye dentro del valor medido, '
     'porque no resulta observable sin convertirla en un parámetro elegido por los autores; se '
     'trata por separado y como supuesto declarado en la Sección 5.1.4. La regla que ordena todo '
     'el procedimiento es que lo medido y lo supuesto no se combinan nunca dentro de una misma '
     'cifra.', 'Body Text'),

    ('Muestra e instrumento. Se procesaron doce órdenes, compuestas por nueve con stock disponible '
     'y tres sin stock, proporción próxima a la distribución observada en la corrida automatizada '
     '(35 y 15 sobre 50). Las dos ramas demandan trabajo manual distinto y la muestra debe '
     'reflejar ambas. Las órdenes se crearon en la base con un identificador de procedencia '
     'propio, distinto del que llevan las órdenes medidas del pipeline, de modo que sus marcas '
     'temporales a escala humana no contaminen el cálculo de MTTD y MTTR; todas las consultas de '
     'resultados del Capítulo 5 filtran por ese identificador. La medición se realizó con un '
     'cronómetro por fases desarrollado para este fin, que registra cada transición y exporta los '
     'tiempos a un archivo separado por comas sin intervención manual sobre los valores.',
     'Body Text'),

    ('Reglas fijadas antes de comenzar. Se estableció por anticipado que el operador trabajaría a '
     'ritmo normal, sin preparar consultas de antemano, y que toda orden durante cuya ejecución '
     'se produjera una interrupción sería descartada y consignada como tal. Dos de las doce '
     'órdenes se descartaron por aplicación de esa regla; quedan registradas en el archivo de '
     'resultados con su marca de descarte y no se eliminaron del archivo. El resultado primario '
     'se calcula sobre las diez órdenes restantes.', 'Body Text'),

    ('Criterio de análisis. Se reportan media, mediana, desvío estándar, mínimo y máximo del '
     'tiempo total y de cada fase. Dado el tamaño reducido de la muestra, el intervalo de '
     'confianza de la media se calcula mediante la distribución t de Student con nueve grados de '
     'libertad y no por aproximación normal, que subestimaría su amplitud. Se controla además el '
     'efecto de aprendizaje mediante el coeficiente de correlación de Pearson entre el orden de '
     'ejecución y el tiempo empleado, cuyo resultado se reporta junto con los descriptivos en la '
     'Sección 5.1.4. Los tiempos crudos, las etiquetas de descarte y el texto de las '
     'notificaciones redactadas se transcriben en el Anexo I.', 'Body Text'),
]

creados = insertar_bloque(d, 177, BLOQUE)
print('§3.5.5 insertada: %d párrafos' % len(creados))

# --- Listado de Tablas: alta de la 5.4 y descripción de la 5.3 ---
t1 = d.tables[1]
import copy
for ri, r in enumerate(t1.rows):
    if r.cells[0].text.strip() == 'Tabla 5.3':
        set_cell(t1, ri, 1, 'Resultados de las pruebas de carga y de concurrencia')
        nueva_tr = copy.deepcopy(r._tr)
        r._tr.addnext(nueva_tr)
        set_cell(t1, ri + 1, 0, 'Tabla 5.4')
        set_cell(t1, ri + 1, 1, 'Descriptivos del tiempo de procesamiento manual, por fase y total')
        set_cell(t1, ri + 1, 2, '5.1.4')
        break

d.save(RUTA)

d2 = Document(RUTA)
print()
for i in range(176, 186):
    print('P%-4d [%-16s] %s' % (i, d2.paragraphs[i].style.name, d2.paragraphs[i].text.strip()[:130]))
print()
print('=== Listado de Tablas (5.x) ===')
for r in d2.tables[1].rows:
    if r.cells[0].text.strip().startswith('Tabla 5'):
        print('  ' + ' | '.join(c.text.strip()[:70] for c in r.cells))
