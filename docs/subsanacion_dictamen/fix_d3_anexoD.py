# -*- coding: utf-8 -*-
"""Anexo D — obligatoria 8b, y un defecto de fondo que el dictamen no vio.

El dictamen observa que el Anexo D transcribe diez entradas mientras la Seccion
4.2.4 declara veintitres. El problema real es peor: CUATRO de esas diez filas
—las categorias "Cuotas", "Tracking", "Factura A" y "Mayorista"— no existen en
la base de datos. El esquema seed tiene 23 entradas repartidas en 8 categorias:
Pagos, Envios, Devoluciones, Garantia, Soporte, Facturacion, Productos y Pedidos.
Y la Seccion 4.2.4 arrastra esa lista inventada al declarar "10 categorias".

Se reemplaza la tabla completa por las 23 entradas reales, transcriptas desde
init_simple.sql (6) y seed_expand.sql (17), y se corrige el recuento de 4.2.4.
"""
import sys
import os
import re
import copy
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'


# ---------------------------------------------------------------- fuente real
def leer_faqs(ruta):
    txt = open(ruta, encoding='utf-8').read()
    m = re.search(r"INSERT INTO faq_responses \(question, answer, category\) VALUES"
                  r"(.*?)(?:ON CONFLICT|;)\s*\n", txt, re.S)
    if not m:
        return []
    return re.findall(r"\(\s*'((?:[^']|'')*)'\s*,\s*'((?:[^']|'')*)'\s*,\s*'((?:[^']|'')*)'\s*\)",
                      m.group(1))


faqs = leer_faqs('init_simple.sql') + leer_faqs('seed_expand.sql')
assert len(faqs) == 23, 'se esperaban 23 FAQ, se leyeron %d' % len(faqs)
cats = []
for _, _, c in faqs:
    if c not in cats:
        cats.append(c)
print('FAQ leidas del seed : %d' % len(faqs))
print('categorias reales   : %d -> %s' % (len(cats), ', '.join(cats)))
print()


def resumir(ans, tope=118):
    """Primera oracion util de la respuesta, recortada. Sin reescribir contenido."""
    ans = ans.replace("''", "'").strip()
    partes = re.split(r'(?<=[.!?])\s+', ans)
    out = partes[0]
    for p in partes[1:]:
        if len(out) + 1 + len(p) > tope:
            break
        out += ' ' + p
    if len(out) > tope:
        out = out[:tope].rsplit(' ', 1)[0] + '…'
    return out


d = Document(RUTA)

# ------------------------------------------------------- la tabla del Anexo D
tabD = None
for t in d.tables:
    if t.rows[0].cells[0].text.strip() == 'Categoría' and 'Pregunta' in t.rows[0].cells[1].text:
        tabD = t
        break
assert tabD is not None, 'no encontre la tabla del Anexo D'
print('tabla del Anexo D: %d filas (1 encabezado + %d datos)' % (len(tabD.rows), len(tabD.rows) - 1))

# se conserva el encabezado y una fila de datos como molde de formato
molde = copy.deepcopy(tabD.rows[1]._tr)
for fila in list(tabD.rows[1:]):
    fila._tr.getparent().remove(fila._tr)

for pregunta, respuesta, categoria in faqs:
    tr = copy.deepcopy(molde)
    tabD._tbl.append(tr)
    i = len(tabD.rows) - 1
    set_cell(tabD, i, 0, categoria)
    set_cell(tabD, i, 1, pregunta.replace("''", "'"))
    set_cell(tabD, i, 2, resumir(respuesta))
print('reescritas %d filas de datos' % (len(tabD.rows) - 1))

# ------------------------------------------------- 4.2.4: 10 categorias -> 8
n = replace_everywhere(d,
    'FAQ (23 entradas, 10 categorías): Pagos, Envíos, Devoluciones, Garantía, Soporte técnico, '
    'Facturación, Cuotas, Tracking, Factura A, Mayorista.',
    'FAQ (23 entradas, 8 categorías): Pagos, Envíos, Devoluciones, Garantía, Soporte, '
    'Facturación, Productos y Pedidos. El contenido íntegro se transcribe en el Anexo D.')
print('4.2.4 recuento de categorias corregido: %d' % n)

# --------------------------------------------- encabezado del Anexo D y nota
n = replace_everywhere(d, 'Anexo D: FAQ Predefinidas',
                          'Anexo D: FAQ predefinidas (base de conocimiento completa)')
print('encabezado del Anexo D: %d' % n)

anc = None
for i, p in enumerate(d.paragraphs):
    if p.text.strip().startswith('Anexo D: FAQ predefinidas'):
        anc = p
        break
insert_paragraph_after(anc,
    'Se transcribe la base de conocimiento completa: las veintitrés entradas de la tabla '
    'faq_responses, en el orden en que las cargan los guiones init_simple.sql y seed_expand.sql '
    'del repositorio. La columna de respuesta reproduce el inicio del texto almacenado; el texto '
    'íntegro está en esos guiones. Conviene recordar, para leer este anexo en su justo alcance, '
    'lo establecido en la Sección 4.4.3: este contenido está cargado en la base y el flujo lo '
    'recupera, pero no llega al prompt del modelo, de modo que no intervino en las respuestas '
    'medidas en el Capítulo 5.',
    estilo='First Paragraph')
print('nota de alcance insertada')

d.save(RUTA)
print()
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
for t in d2.tables:
    if t.rows[0].cells[0].text.strip() == 'Categoría' and 'Pregunta' in t.rows[0].cells[1].text:
        print()
        print('=== ANEXO D (%d entradas) ===' % (len(t.rows) - 1))
        for r in t.rows:
            print('  %-13s | %-48s | %s' % (r.cells[0].text.strip(),
                                            r.cells[1].text.strip()[:48],
                                            r.cells[2].text.strip()[:62]))
TXT = '\n'.join(p.text for p in d2.paragraphs) + '\n' + '\n'.join(
    c.text for t in d2.tables for r in t.rows for c in r.cells)
print()
for pat in ('Cuotas', 'Tracking', 'Factura A', 'Mayorista', '10 categorías'):
    print('  categoria fantasma "%s": %d ocurrencias %s'
          % (pat, TXT.count(pat), '' if TXT.count(pat) == 0 else '<-- REVISAR'))
