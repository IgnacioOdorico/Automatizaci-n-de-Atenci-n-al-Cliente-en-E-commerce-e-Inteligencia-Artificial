# -*- coding: utf-8 -*-
"""Tabla 5.3 — reconstrucción con las dos poblaciones separadas (Bloqueante #3)."""
import sys
import copy
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
t = d.tables[18]


def clonar_ultima_fila(tabla):
    """Duplica la última fila conservando su formato."""
    nueva = copy.deepcopy(tabla.rows[-1]._tr)
    tabla.rows[-1]._tr.addnext(nueva)
    return tabla.rows[-1]


FILAS = [
    ('Métrica', 'Resultado'),
    ('E1.a — Órdenes enviadas (secuenciales, espaciadas 2 s)', '50'),
    ('E1.a — Órdenes procesadas', '50/50 (100 %)'),
    ('E1.a — Órdenes confirmadas', '35 (70 %)'),
    ('E1.a — Órdenes sin stock', '15 (30 %)'),
    ('E1.a — Errores de pipeline', '0'),
    ('E1.b — Rondas × solicitudes simultáneas', '6 × 20 = 120 órdenes'),
    ('E1.b — Stock inicial por ronda', '5 unidades (PROD-005)'),
    ('E1.b — Órdenes confirmadas por ronda', '5 en las 6 rondas (ninguna sobreventa)'),
    ('E1.b — Productos con stock negativo', '0'),
    ('E1.b — Órdenes sin procesar bajo concurrencia', '49 de 120 (40,8 %)'),
    ('E1.b — MTTD medio bajo concurrencia', '0,092 s (±0,025 s; máx. 0,157 s)'),
    ('E1.b — End-to-end medio bajo concurrencia', '0,192 s (±0,024 s; máx. 0,260 s)'),
    ('Total de órdenes con data_source = \'measured\'', '173 (50 de E1.a + 120 de E1.b + 3 de verificación)'),
]

while len(t.rows) < len(FILAS):
    clonar_ultima_fila(t)

for ri, (a, b) in enumerate(FILAS):
    set_cell(t, ri, 0, a)
    set_cell(t, ri, 1, b)

# --- Tabla 6.1, OE1: la cifra 50/50 ahora tiene contexto ---
t61 = d.tables[25]
set_cell(t61, 1, 2,
    'MTTD: 0,009 s / MTTR: 0,054 s / Total: 0,063 s (n = 50, corrida E1.a). '
    '50/50 órdenes procesadas sin errores. Bajo concurrencia (E1.b, 120 órdenes) '
    'el pipeline sostuvo la atomicidad del descuento de stock sin sobreventa.')

d.save(RUTA)

d2 = Document(RUTA)
print('=== TABLA 5.3 (idx 18) ===')
for ri, r in enumerate(d2.tables[18].rows):
    print('  r%-2d | %-52s | %s' % (ri, r.cells[0].text.strip(), r.cells[1].text.strip()))
print()
print('=== TABLA 6.1 r1 ===')
print('  ' + d2.tables[25].rows[1].cells[2].text.strip())
