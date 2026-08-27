# -*- coding: utf-8 -*-
"""
docxkit — utilidades de edicion para TESIS_FINAL_UTN_v6.docx

Regla de oro aprendida a los golpes: el texto de un parrafo esta partido en
muchos runs. Un replace por run FALLA cuando la cadena cruza el limite de dos
runs. Todo lo de aca abajo trabaja sobre el texto concatenado del parrafo y
despues reescribe los runs afectados, conservando el formato del primero.
"""
import sys
import io

if sys.stdout.encoding is None or 'utf' not in (sys.stdout.encoding or '').lower():
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

import copy
from docx.text.paragraph import Paragraph


# ------------------------------------------------------------------
#  Reemplazo consciente de runs
# ------------------------------------------------------------------
def replace_in_paragraph(p, old, new, count=0):
    """Reemplaza `old` por `new` en el parrafo, aunque cruce varios runs.

    Devuelve la cantidad de reemplazos hechos. count=0 -> todos.
    """
    runs = p.runs
    if not runs:
        return 0
    full = ''.join(r.text for r in runs)
    if old not in full:
        return 0

    # mapa: posicion global -> (indice de run, offset dentro del run)
    bounds = []
    pos = 0
    for i, r in enumerate(runs):
        bounds.append((pos, pos + len(r.text), i))
        pos += len(r.text)

    hechos = 0
    desde = 0          # avanza SIEMPRE: si `new` contiene a `old`, buscar desde 0 cicla infinito
    while True:
        if count and hechos >= count:
            break
        full = ''.join(r.text for r in p.runs)
        idx = full.find(old, desde)
        if idx < 0:
            break
        fin = idx + len(old)
        desde = idx + len(new)   # el próximo intento arranca después de lo ya escrito

        runs = p.runs
        bounds = []
        pos = 0
        for i, r in enumerate(runs):
            bounds.append((pos, pos + len(r.text), i))
            pos += len(r.text)

        afectados = [b for b in bounds if b[0] < fin and b[1] > idx]
        if not afectados:
            break

        pri = afectados[0]
        # prefijo que sobrevive en el primer run afectado
        pref = runs[pri[2]].text[: idx - pri[0]]
        ult = afectados[-1]
        suf = runs[ult[2]].text[fin - ult[0]:]

        runs[pri[2]].text = pref + new + (suf if pri[2] == ult[2] else '')
        for b in afectados[1:]:
            runs[b[2]].text = suf if b[2] == ult[2] else ''
        hechos += 1

    return hechos


def replace_everywhere(doc, old, new, solo_indices=None):
    """Reemplaza en parrafos del cuerpo y en todas las celdas de tabla."""
    n = 0
    for i, p in enumerate(doc.paragraphs):
        if solo_indices is not None and i not in solo_indices:
            continue
        n += replace_in_paragraph(p, old, new)
    if solo_indices is None:
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        n += replace_in_paragraph(p, old, new)
    return n


def set_cell(tabla, fila, col, texto):
    """Reescribe una celda conservando el formato del primer run."""
    celda = tabla.rows[fila].cells[col]
    ps = celda.paragraphs
    p = ps[0]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(texto)
    for extra in ps[1:]:
        extra._element.getparent().remove(extra._element)


def cell_text(tabla, fila, col):
    return tabla.rows[fila].cells[col].text.strip()


# ------------------------------------------------------------------
#  Insercion de parrafos
# ------------------------------------------------------------------
def insert_paragraph_after(par, texto='', estilo=None):
    """Crea un parrafo nuevo inmediatamente despues de `par`."""
    nuevo_el = copy.deepcopy(par._element)
    # limpiar contenido heredado
    for hijo in list(nuevo_el):
        if hijo.tag.endswith('}r') or hijo.tag.endswith('}hyperlink'):
            nuevo_el.remove(hijo)
    par._element.addnext(nuevo_el)
    nuevo = Paragraph(nuevo_el, par._parent)
    if estilo is not None:
        try:
            nuevo.style = estilo
        except KeyError:
            pass
    if texto:
        nuevo.add_run(texto)
    return nuevo


def insertar_bloque(doc, indice_ancla, bloques):
    """Inserta una lista de (texto, estilo) despues del parrafo `indice_ancla`.

    Devuelve la lista de parrafos creados, en orden.
    """
    ancla = doc.paragraphs[indice_ancla]
    creados = []
    actual = ancla
    for texto, estilo in bloques:
        actual = insert_paragraph_after(actual, texto, estilo)
        creados.append(actual)
    return creados


# ------------------------------------------------------------------
#  Verificacion
# ------------------------------------------------------------------
def buscar(doc, patron, etiqueta='', lim=220):
    import re
    print('--- %s  /%s/' % (etiqueta or 'buscar', patron))
    n = 0
    for i, p in enumerate(doc.paragraphs):
        if re.search(patron, p.text, re.I):
            n += 1
            print('   P%-4d %s' % (i, p.text.strip()[:lim].replace('\n', ' ')))
    for ti, t in enumerate(doc.tables):
        for ri, r in enumerate(t.rows):
            linea = ' | '.join(c.text.strip() for c in r.cells)
            if re.search(patron, linea, re.I):
                n += 1
                print('   T%d.r%-2d %s' % (ti, ri, linea[:lim]))
    if n == 0:
        print('   (sin coincidencias)')
    return n
