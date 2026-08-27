# -*- coding: utf-8 -*-
"""P348 no matcheó por reemplazo parcial: se reescribe el párrafo completo."""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
p = d.paragraphs[348]

NUEVO = (
    'El modelo GPT-4o-mini demostró ser adecuado para la clasificación de intenciones en dominios '
    'acotados, manejando correctamente mensajes con errores ortográficos, abreviaciones y varias '
    'preguntas dentro de un mismo mensaje. El dato relevante es que ese 92,7 % de precisión se '
    'alcanzó en régimen zero-shot: sin ajuste fino, sin ejemplos etiquetados en el prompt y sin '
    'inyección de la base de FAQ, con un prompt de sistema de poco más de mil caracteres que se '
    'limita a enumerar las cuatro categorías y a fijar el formato de salida (Anexo H). El '
    'resultado debe leerse por lo tanto como un piso del desempeño alcanzable en la tarea y no '
    'como su techo: las técnicas de prompting sistematizadas por Liu et al. (2023), así como el '
    'cableado del contexto de FAQ que el flujo ya construye, constituyen margen de mejora todavía '
    'sin explotar.')

p.runs[0].text = NUEVO
for r in p.runs[1:]:
    r.text = ''

d.save(RUTA)

d2 = Document(RUTA)
print(d2.paragraphs[348].text)
print()
buscar(d2, r'few.?shot|inyección dinámica', 'few-shot restantes (solo 2.2.2 y bibliografía)')
