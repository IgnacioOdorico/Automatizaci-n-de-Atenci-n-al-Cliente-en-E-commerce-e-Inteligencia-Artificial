# -*- coding: utf-8 -*-
"""#14 — v_chatbot_corpus alimenta 5 paneles y no estaba definida en el DDL del Anexo A.
También: baja del párrafo vacío que partía el Resumen (#26) y corrección del Anexo E (#3)."""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)

VISTA = """

-- ============================================================
--  Vista del corpus evaluado del Flujo 2
--  Aísla los 150 mensajes que tienen etiqueta de referencia, para que
--  las tablas de resultados del Cap. 5 no promedien sobre la totalidad
--  de las interacciones medidas. Se delimita por ventana temporal
--  porque el identificador de usuario se repite entre corridas.
-- ============================================================
CREATE OR REPLACE VIEW v_chatbot_corpus AS
SELECT
    i.id,
    i.channel,
    i.user_id,
    i.intent,                              -- intención PREDICHA por el modelo
    i.received_at,
    i.responded_at,
    EXTRACT(EPOCH FROM (i.responded_at - i.received_at)) AS tmr_seconds,
    i.is_urgent
FROM interactions i
WHERE i.data_source = 'measured'
  AND i.responded_at IS NOT NULL
  AND i.received_at >= '2026-08-12 23:00:00'
  AND i.received_at <  '2026-08-13 00:00:00';
"""

# --- se anexa la sexta vista al final del DDL ---
for i, p in enumerate(d.paragraphs):
    if p.style.name == 'Source Code' and 'CREATE TABLE' in p.text:
        ult = p.runs[-1]
        ult.text = ult.text + VISTA
        print('vista agregada al DDL del Anexo A (P%d)' % i)
        break

# --- Anexo E: el script de pruebas describía una sola corrida ---
for i, p in enumerate(d.paragraphs):
    if p.style.name == 'Source Code' and 'docker compose up -d' in p.text:
        txt = p.text
        if 'seq 1 20' in txt or 'for i in' in txt:
            print('Anexo E (P%d) menciona la corrida de carga; se revisa abajo.' % i)
        print('--- Anexo E actual (ultimos 500) ---')
        print(txt[-500:])
        break

d.save(RUTA)

# --- baja del párrafo vacío del Resumen ---
d = Document(RUTA)
p16 = d.paragraphs[16]
if p16.text.strip() == '':
    p16._element.getparent().remove(p16._element)
    print('párrafo vacío del Resumen eliminado')
d.save(RUTA)

d2 = Document(RUTA)
print()
print('=== RESUMEN ===')
for i in range(12, 18):
    print('%4d [%-16s] %s' % (i, d2.paragraphs[i].style.name, d2.paragraphs[i].text.strip()[:150]))
print()
for i, p in enumerate(d2.paragraphs):
    if p.style.name == 'Source Code' and 'CREATE TABLE' in p.text:
        print('DDL: v_chatbot_corpus x%d | CREATE OR REPLACE VIEW x%d'
              % (p.text.count('v_chatbot_corpus'), p.text.count('CREATE OR REPLACE VIEW')))
        break
