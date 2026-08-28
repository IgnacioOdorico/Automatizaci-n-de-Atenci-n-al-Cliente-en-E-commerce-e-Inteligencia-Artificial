# -*- coding: utf-8 -*-
"""Verificacion del tercer dictamen: 10 obligatorias + 8 recomendadas.

Cada control comprueba el ESTADO FINAL del documento, no que un script haya
corrido. Un control que pasa por casualidad no sirve, asi que varios chequean
a la vez la ausencia del defecto y la presencia de lo que lo reemplaza.
"""
import sys
import io
import re
import zipfile
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.oxml.ns import qn

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
P = [p.text.strip() for p in d.paragraphs]
TXT_P = '\n'.join(P)
TXT_T = '\n'.join(' | '.join(c.text for c in r.cells) for t in d.tables for r in t.rows)
TODO = TXT_P + '\n' + TXT_T
HEAD = [(p.style.name, p.text.strip()) for p in d.paragraphs if p.style.name.startswith('Heading')]
XML = zipfile.ZipFile(RUTA).read('word/document.xml').decode('utf8')
SETT = zipfile.ZipFile(RUTA).read('word/settings.xml').decode('utf8')

fallos = []
n_ok = 0


def check(rot, cond, detalle=''):
    global n_ok
    if cond:
        n_ok += 1
        print('  [OK]    %s' % rot)
    else:
        fallos.append(rot + ((' — ' + detalle) if detalle else ''))
        print('  [FALLA] %s %s' % (rot, detalle))


print('=' * 78)
print(' OBLIGATORIAS')
print('=' * 78)

# --- 1
check('1a  ABSTRACT presente', any(t == 'ABSTRACT' for _, t in HEAD))
check('1b  keywords en ingles', 'Keywords: business process automation' in TODO)
check('1c  declaracion de originalidad',
      any(t == 'DECLARACIÓN DE ORIGINALIDAD' for _, t in HEAD))
check('1d  el abstract reporta las cifras clave',
      all(x in TODO for x in ('0.009 seconds', '92.7 %', 'Mann-Whitney U = 0')))

# --- 2
check('2a  no queda "nueve con stock"', 'nueve con stock' not in TODO)
check('2b  composicion real declarada (7+3 validas, 2 descartadas)',
      'siete correspondieron a la rama con stock' in TODO and 'ORD-E4-002' in TXT_P)

# --- 3
check('3   Tabla 6.1 no invoca un compromiso numerico ausente de 1.5.2',
      'El objetivo comprometía 5 tablas y 5 vistas' not in TODO)

# --- 4
check('4a  la promesa incumplida de 3.2 ya no esta',
      'con evaluación cualitativa de coherencia de las respuestas' not in TODO)
check('4b  se declara explicitamente lo no medido',
      'No se evaluó la calidad del contenido de las respuestas' in TODO)
check('4c  la triangulacion que se afirma es la que existe',
      'triangulación de fuentes de evidencia' in TODO and '51,28' in TODO)
check('4d  el Capitulo 7 lleva la rubrica del procedimiento faltante',
      'rúbrica de tres niveles' in TODO)

# --- 5
check('5a  no se afirma "resuelve de forma autónoma"',
      'resuelve de forma autónoma la totalidad' not in TODO)
check('5b  el indicador se llama tasa de no escalada',
      'Tasa de no escalada' in TODO and 'No mide que esa respuesta fuera correcta' in TODO)

# --- 6
oe1 = ''
for t in d.tables:
    for r in t.rows:
        if r.cells[0].text.strip() == 'OE1':
            oe1 = ' | '.join(c.text for c in r.cells)
check('6a  OE1 incorpora la perdida del 40,8 %', '40,8 %' in oe1)
check('6b  OE1 acota su veredicto al regimen secuencial', 'régimen secuencial' in oe1)

# --- 7
check('7a  sin vocabulario inferencial sin inferencia',
      'reduce significativamente el tiempo' not in TODO
      and 'reducción significativa de los tiempos' not in TODO
      and 'reduce significativamente los tiempos' not in TODO)
check('7b  contraste no parametrico ejecutado y reportado',
      'Mann-Whitney' in TODO and '2,7 × 10⁻¹¹' in TODO)
check('7c  contraste parametrico ejecutado y reportado',
      'Welch' in TODO and '19,05' in TODO and '1,4 × 10⁻⁸' in TODO)
check('7d  Tabla 5.10 sostiene H1 sobre la prueba',
      any('Mann-Whitney' in ' '.join(c.text for c in r.cells)
          for t in d.tables for r in t.rows))
check('7e  se declara el limite del contraste (asimetria de constructo)',
      'no resuelve la asimetría' in TODO)

# --- 8
check('8a  Tabla 4.8 remite al Anexo J', 'valor obtenido en el Anexo H' not in TODO
      and 'valor obtenido en el Anexo J' in TODO)
tabD = [t for t in d.tables
        if t.rows[0].cells[0].text.strip() == 'Categoría'
        and 'Pregunta' in t.rows[0].cells[1].text]
check('8b  Anexo D transcribe las 23 entradas',
      len(tabD) == 1 and len(tabD[0].rows) - 1 == 23,
      'filas de datos: %d' % (len(tabD[0].rows) - 1 if tabD else -1))
check('8c  el recuento de categorias coincide con la base (8, no 10)',
      '8 categorías' in TODO and '10 categorías' not in TODO)
check('8d  no quedan categorias inexistentes en la base',
      not any(x in TODO for x in ('Cuotas,', 'Tracking,', 'Factura A,', 'Mayorista')))

# --- 9
check('9a  un solo indice: el campo automatico',
      len(re.findall(r'<w:sdt>', XML)) >= 1 and 'TOC \\o' in XML)
check('9b  el indice manual fue dado de baja',
      len([p for p in d.paragraphs
           if p.style.name == 'Compact'
           and p.text.strip().startswith(('Capítulo ', '1.1 ', '2.1 '))]) == 0)
# ojo: <w:docPartGallery w:val="Table of Contents"/> es el identificador interno con
# el que Word reconoce el control como indice — debe quedar. Solo se controla el texto visible.
check('9c  el indice esta en castellano (texto visible)',
      not re.search(r'<w:t[^>]*>[^<]*Table of Contents', XML)
      and re.search(r'<w:t[^>]*>ÍNDICE<', XML) is not None)
check('9d  Word recalcula el indice al abrir', 'updateFields' in SETT)
check('9e  la portada salio del arbol de contenidos',
      not any(t == 'PORTADA' for st, t in HEAD))
check('9f  no quedan encabezados sin numerar',
      not [t for st, t in HEAD if st == 'Heading 3' and not t[:1].isdigit()],
      str([t for st, t in HEAD if st == 'Heading 3' and not t[:1].isdigit()]))
check('9g  5.1.1 a 5.1.4 completos',
      all(any(t.startswith(x) for _, t in HEAD) for x in ('5.1.1', '5.1.2', '5.1.3', '5.1.4')))
check('9h  4.1.1 y 4.5.1 numerados',
      all(any(t.startswith(x) for _, t in HEAD) for x in ('4.1.1', '4.5.1')))
check('9i  las figuras del cuerpo tienen una sola numeracion',
      'Figura A1' not in TODO and 'Figura A2' not in TODO
      and 'Figura 10:' in TODO and 'Figura 11:' in TODO)
check('9j  el Anexo F ya no repite el listado del frontispicio',
      len([t for t in d.tables if t.rows[0].cells[0].text.strip() == 'Figura'
           and len(t.rows) == 12]) == 1)

# --- 10
check('10  Tabla 3.1 no declara 3 canales', '3 canales' not in TODO
      and 'dos canales (WhatsApp simulado y Telegram real)' in TODO)

print()
print('=' * 78)
print(' RECOMENDADAS')
print('=' * 78)

check('R1a estado del arte con 13 antecedentes',
      len([t for t in d.tables if t.rows[0].cells[0].text.strip() == 'Antecedente'][0].rows) - 2 == 13)
check('R1b literatura arbitrada argentina incorporada',
      'Alderete, M. V., Jones, C., & Motta, J. J. (2017)' in TXT_P
      and 'Alderete, M. V., & Porris, M. S. (2023)' in TXT_P)
check('R1c literatura arbitrada latinoamericana incorporada',
      all(x in TXT_P for x in ('Aguirre Mayorga', 'Ramos De Santis', 'Pachas-Santos',
                               'Bravo Maruri', 'Fondevila-Gascón')))
check('R1d bases regionales declaradas en el procedimiento',
      'SciELO, Redalyc, Dialnet' in TODO)
check('R1e nueva subseccion regional',
      any(t.startswith('2.5.4 Antecedentes latinoamericanos') for _, t in HEAD)
      and any(t.startswith('2.5.5 Vacío identificado') for _, t in HEAD))
check('R2  operacionalizacion del MTTD declarada',
      'no la capa HTTP que lo recibe' in TODO)
check('R3  IC del factor por Fieller, con la simplificacion declarada',
      'Fieller' in TODO and '686× a 875×' in TODO and 'inferior al 0,4 %' in TODO)
check('R4  exactitud (global) y precision (por clase) diferenciadas',
      'se reserva «exactitud (accuracy)» para esta métrica global' in TODO)
check('R5  el canal simulado ya no se llama canal WhatsApp a secas',
      'canal simulado en formato WhatsApp Cloud API' in TODO
      and 'simulación del canal WhatsApp' not in TODO)
check('R6  salvedad sobre los nodos 5 y 6',
      'no inciden en la salida del sistema en la configuración efectivamente medida' in TODO)
check('R7a la Seccion 2.4 declara su funcion',
      'Corresponde declarar de antemano la función de esta sección' in TODO)
check('R7b el Capitulo 6 retoma el marco teorico',
      'escala de madurez de proceso de van der Aalst' in TODO
      and 'capacidad de respuesta (responsiveness)' in TXT_P.split('CAPÍTULO 6')[-1])
check('R8a locator de CACE preciso',
      'Estudio Anual de Comercio Electrónico 2024' in TODO
      and 'cace.org.ar/estadisticas' not in TODO)
check('R8b fecha de consulta de Jurafsky y Martin',
      'Recuperado el 28 de agosto de 2026, de https://web.stanford.edu/~jurafsky/slp3/' in TODO)

print()
print('=' * 78)
print(' EXTRA: defectos detectados por cuenta propia')
print('=' * 78)
check('E1  6.4 enuncia sus limitaciones, no las delega',
      'Se enuncian aquí, para que el lector' in TODO)
check('E2  sin residuos de la afirmacion de few-shot',
      'ampliar los ejemplos de few-shot' not in TODO)
check('E3  el trabajo se autodenomina de una sola manera',
      'esta tesis' not in TXT_P)

print()
print('=' * 78)
print(' INTEGRIDAD DEL ARCHIVO')
print('=' * 78)
check('Z1  zip integro', zipfile.ZipFile(RUTA).testzip() is None)
check('Z2  11 imagenes', sum(1 for r in d.part.rels.values() if 'image' in r.reltype) == 11)
check('Z3  36 tablas', len(d.tables) == 36)
check('Z4  12 capitulos/encabezados de nivel 1',
      len([t for st, t in HEAD if st == 'Heading 1']) >= 11)

# Z5/Z6: la imagen de la Figura 1 vivia DENTRO del encabezado "Servicios Docker" y
# se perdio al renumerarlo (run.text='' borra el <w:drawing> del run). Estos dos
# controles impiden que esa clase de regresion vuelva a pasar inadvertida.
from PIL import Image
_malas = []
for _i, _p in enumerate(d.paragraphs):
    _m = re.match(r'Figura (\d+):', _p.text.strip())
    if not _m:
        continue
    _xml = d.paragraphs[_i - 1]._element.xml
    _r = re.search(r'r:embed="(rId\d+)"', _xml)
    if not _r:
        _malas.append('Figura %s sin imagen' % _m.group(1))
        continue
    _im = Image.open(io.BytesIO(d.part.rels[_r.group(1)].target_part.blob))
    _cx, _cy = map(int, re.search(r'<wp:extent cx="(\d+)" cy="(\d+)"', _xml).groups())
    if abs((_cx / _cy) - (_im.size[0] / _im.size[1])) >= 0.01:
        _malas.append('Figura %s deformada' % _m.group(1))
check('Z5  las 11 figuras tienen imagen y conservan su proporcion',
      not _malas and sum(1 for p in d.paragraphs if 'blip' in p._element.xml) == 11,
      str(_malas))
check('Z6  ningun encabezado lleva una imagen colgando',
      not [p for p in d.paragraphs
           if p.style.name.startswith('Heading') and 'blip' in p._element.xml])

print()
print('=' * 78)
if fallos:
    print(' RESULTADO: %d/%d — %d FALLAS' % (n_ok, n_ok + len(fallos), len(fallos)))
    print('=' * 78)
    for f in fallos:
        print('  - %s' % f)
    sys.exit(1)
print(' RESULTADO: %d/%d — SIN FALLAS' % (n_ok, n_ok))
print('=' * 78)
