# -*- coding: utf-8 -*-
"""Anexo E reescrito (los comandos no funcionaban) y puerto de PostgreSQL corregido.

Verificado contra los workflows y la base:
  * path real del chatbot: whatsapp-business (el anexo decía chatbot-whatsapp -> 404)
  * el webhook del Flujo 1 espera order_number, customer_name, customer_email,
    customer_phone, product_sku y quantity. El anexo mandaba product_id y omitía
    order_number, que es NOT NULL UNIQUE -> 500.
  * el normalizador del chatbot acepta payload PLANO ({from, message}), no la
    estructura anidada de la Cloud API. El ejemplo original tenía la forma correcta.
  * el host publica PostgreSQL en 5433 (mapeo "5433:5432"), no en 5432.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)

ANEXO_E = """# ---------- Levantar el entorno ----------
docker compose up -d
docker compose ps

# ---------- Verificar la base ----------
# El host publica PostgreSQL en 5433 (mapeo 5433:5432) para no colisionar
# con una instalación nativa que ocupe el 5432. Dentro de la red de Docker
# los servicios se ven entre sí como postgres:5432.
docker exec -it tesis_postgres psql -U n8n_user -d ecommerce_tesis -c "\\dt"

# ---------- Flujo 1: orden con stock suficiente ----------
curl -X POST http://localhost:5678/webhook/orden-nueva \\
  -H "Content-Type: application/json" \\
  -d '{"order_number":"ORD-TEST-001","customer_name":"Test","customer_email":"test@example.com","customer_phone":"5492615551234","product_sku":"PROD-018","quantity":1}'

# ---------- Flujo 1: rama sin stock ----------
# Se piden más unidades de las disponibles para forzar la rama no_stock.
curl -X POST http://localhost:5678/webhook/orden-nueva \\
  -H "Content-Type: application/json" \\
  -d '{"order_number":"ORD-TEST-002","customer_name":"Test","customer_email":"test@example.com","customer_phone":"5492615551234","product_sku":"PROD-005","quantity":999}'

# ---------- Flujo 2: chatbot ----------
# El path del webhook es whatsapp-business. El nodo normalizador acepta un
# payload plano con los campos from y message.
curl -X POST http://localhost:5678/webhook/whatsapp-business \\
  -H "Content-Type: application/json" \\
  -d '{"from":"5492615551234","message":"¿Cuáles son los métodos de pago?"}'

# ---------- Corrida de carga secuencial (50 órdenes, espaciadas 2 s) ----------
for i in $(seq 1 50); do
  curl -s -X POST http://localhost:5678/webhook/orden-nueva \\
    -H "Content-Type: application/json" \\
    -d "{\\"order_number\\":\\"ORD-CARGA-$i\\",\\"customer_name\\":\\"Carga $i\\",\\"customer_email\\":\\"carga$i@example.com\\",\\"customer_phone\\":\\"5492615551234\\",\\"product_sku\\":\\"PROD-018\\",\\"quantity\\":1}"
  sleep 2
done

# ---------- Prueba de concurrencia (20 solicitudes simultáneas) ----------
# Contra un producto con stock deliberadamente escaso, para verificar que no
# se confirmen más órdenes que unidades disponibles.
for i in $(seq 1 20); do
  curl -s -X POST http://localhost:5678/webhook/orden-nueva \\
    -H "Content-Type: application/json" \\
    -d "{\\"order_number\\":\\"ORD-CONC-$i\\",\\"customer_name\\":\\"Conc $i\\",\\"customer_email\\":\\"conc$i@example.com\\",\\"customer_phone\\":\\"5492615551234\\",\\"product_sku\\":\\"PROD-005\\",\\"quantity\\":1}" &
done
wait

# ---------- Verificar métricas y la invariante de stock ----------
docker exec -it tesis_postgres psql -U n8n_user -d ecommerce_tesis \\
  -c "SELECT * FROM v_metrics_summary;" \\
  -c "SELECT COUNT(*) AS stock_negativo FROM products WHERE stock < 0;"

# ---------- Ver correos generados ----------
# Abrir http://localhost:8025 en el navegador

# ---------- Ver tableros ----------
# Abrir http://localhost:3000 (usuario admin; contraseña definida en el archivo .env)"""

for i, p in enumerate(d.paragraphs):
    if p.style.name == 'Source Code' and 'docker compose up -d' in p.text:
        p.runs[0].text = ANEXO_E
        for r in p.runs[1:]:
            r.text = ''
        print('Anexo E reescrito en P%d' % i)
        break

# --- puerto de PostgreSQL: el host publica 5433 ---
for ti, t in enumerate(d.tables):
    for ri, r in enumerate(t.rows):
        celdas = [c.text.strip() for c in r.cells]
        if celdas and celdas[0] == 'PostgreSQL' and '5432' in celdas:
            set_cell(t, ri, celdas.index('5432'), '5433 → 5432 (host → contenedor)')
            print('puerto corregido en T%d.r%d' % (ti, ri))

d.save(RUTA)

d2 = Document(RUTA)
print()
for ti, t in enumerate(d2.tables):
    for r in t.rows:
        if r.cells[0].text.strip() == 'PostgreSQL':
            print('T%d: %s' % (ti, ' | '.join(c.text.strip()[:46] for c in r.cells)))
