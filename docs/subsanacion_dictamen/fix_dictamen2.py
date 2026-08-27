# -*- coding: utf-8 -*-
"""Segundo dictamen (APROBADO CON OBSERVACIONES MENORES) — sus 6 recomendaciones
más una que el informe no detectó: Ley 25.326 quedó fuera del orden alfabético.

REC 1b (capitalización de capítulos) ya estaba resuelta: los 9 encabezados de
Heading 1 están uniformemente en versalitas. No se toca.
"""
import sys
import os
import copy
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def reescribir(i, texto):
    p = d.paragraphs[i]
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ''


# ===== REC 2 — declarar el origen del supuesto de distribución =====
n = replace_in_paragraph(d.paragraphs[idx('Flujo 2 — Clasificación de intents')],
    'ESTADO_PEDIDO 35 (23,3%) y GENERAL 25 (16,7%).',
    'ESTADO_PEDIDO 35 (23,3 %) y GENERAL 25 (16,7 %). Corresponde precisar el origen de ese '
    'criterio: la frecuencia esperada por tipo de consulta es una estimación del propio equipo, '
    'construida a partir del catálogo de FAQ definido para el caso y de los escenarios de uso '
    'previstos. No proviene de datos de tráfico observados en un e-commerce real, a los que este '
    'trabajo no tuvo acceso, y por lo tanto se declara como supuesto de diseño y no como dato '
    'empírico.')
print('REC 2 (distribución del corpus): %d' % n)

# ===== REC 3 — no afirmar reproducibilidad sin ensayo de reproducción =====
n = replace_in_paragraph(d.paragraphs[idx('Sobre la arquitectura Docker Compose')],
    'la infraestructura basada en cuatro contenedores demostró ser reproducible en lo esencial.',
    'la infraestructura basada en cuatro contenedores está diseñada para ser reproducible.')
print('REC 3 (reproducibilidad): %d' % n)

# ===== REC 4 — anclar el umbral del 85 % en la literatura revisada =====
n = replace_in_paragraph(d.paragraphs[idx('La evaluación de la precisión de clasificación')],
    'Se adopta un umbral mínimo de 85% de accuracy como criterio de aceptación definido por el '
    'equipo, apropiado para un prototipo académico de clasificación de intenciones en un dominio '
    'acotado.',
    'Se adopta un umbral mínimo del 85 % de accuracy como criterio de aceptación definido por el '
    'equipo. Ese valor no deriva de un estándar publicado, y conviene explicitar el razonamiento '
    'que lo sustenta. La literatura revisada en la Sección 2.5 muestra que la clasificación de '
    'intenciones sin ajuste fino alcanza desempeño útil: Parikh et al. (2023) documentan que el '
    'prompting zero-shot basado en descripciones de intención resulta efectivo, y Luo et al. '
    '(2023) obtienen desempeño competitivo con cantidades mínimas de datos. Esos trabajos evalúan '
    'sobre corpus normalizados con decenas de intenciones, mientras que el dominio de este '
    'trabajo se limita a cuatro categorías, lo que vuelve la tarea comparativamente más sencilla. '
    'Por esa razón se fija un umbral exigente para el nivel del prototipo, sin pretender que sea '
    'equiparable a los valores reportados sobre aquellos corpus.')
print('REC 4 (umbral del 85 %%): %d' % n)

# ===== REC 5 — matizar HubSpot como reporte de industria =====
n = replace_everywhere(d,
    'el 82% de los clientes espera una resolución inmediata de sus consultas (HubSpot, 2024).',
    'el 82 % de los clientes espera una resolución inmediata de sus consultas (HubSpot, 2024). '
    'Corresponde señalar que esa cifra proviene de un reporte de industria elaborado sobre una '
    'encuesta comercial y no de una publicación con revisión por pares, de modo que se la '
    'consigna como indicio del contexto de expectativas y no como evidencia arbitrada.')
print('REC 5 (HubSpot): %d' % n)

# ===== REC 6 — OE2 formulado sobre lo efectivamente implementado =====
reescribir(idx('Implementar un chatbot cuya arquitectura contempla tres canales'),
    'Implementar un chatbot multicanal sobre dos canales —WhatsApp, con envío simulado a un '
    'capturador SMTP local, y Telegram, con envío real— que normalice el mensaje entrante, '
    'clasifique la intención con GPT-4o-mini (FAQ, ESTADO_PEDIDO, RECLAMO y GENERAL) y responda '
    'de forma automatizada registrando el TMR; y verificar que ese TMR sea inferior a 10 segundos '
    'y que la precisión de clasificación supere el 85 % (H2a y H2b). La extensión a un tercer '
    'canal por correo electrónico queda planteada desde el inicio como línea futura, y su diseño '
    'se documenta en el Capítulo 7.')
print('REC 6 (OE2): reformulado')

# ===== REC 1a — headings vacíos que ensucian el índice automático =====
i = idx('4.2.2 Diagrama entidad-relación')
fig = d.paragraphs[i + 1]                      # este SÍ lleva la imagen de la Figura 2
if 'blip' in fig._element.xml:
    fig.style = d.paragraphs[idx('Figura 8:') - 1].style   # mismo estilo que las otras figuras
    print('REC 1a: párrafo de la Figura 2 reestilado a %s' % fig.style.name)
for p in list(d.paragraphs):
    if p.style.name.startswith('Heading') and not p.text.strip() and 'blip' not in p._element.xml:
        p._element.getparent().remove(p._element)
        print('REC 1a: heading vacío eliminado')

# ===== EXTRA — Ley 25.326 estaba fuera del orden alfabético =====
i_ley = idx('Ley 25.326 de Protección')
ley_xml = copy.deepcopy(d.paragraphs[i_ley]._element)
d.paragraphs[i_ley]._element.getparent().remove(d.paragraphs[i_ley]._element)
d.paragraphs[idx('Laudon, K. C.')]._element.addnext(ley_xml)
print('EXTRA: Ley 25.326 reubicada entre Laudon y Liu')

d.save(RUTA)

# ===== verificación =====
d2 = Document(RUTA)
print()
vac = [i for i, p in enumerate(d2.paragraphs)
       if p.style.name.startswith('Heading') and not p.text.strip()]
print('headings vacíos restantes: %d' % len(vac))
print()
print('=== bibliografía, tramo L ===')
for p in d2.paragraphs:
    t = p.text.strip()
    if t[:6] in ('Landis', 'Laudon', 'Ley 25', 'Liu, P', 'Luo, H', 'McTear'):
        print('  ' + t[:80])
