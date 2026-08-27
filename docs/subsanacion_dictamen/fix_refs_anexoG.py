# -*- coding: utf-8 -*-
"""Dos hallazgos de la auditoría referencial:

1) §7.1 citaba 'workflows/Flujo 2 — Chatbot IA PRODUCCION.json', archivo que
   NO existe. El real es 'Flujo 2 — Chatbot Omnicanal IA PRODUCCION.json'.
2) El Anexo G repetía los epígrafes "Figura A1:" y "Figura A2:" sobre párrafos
   que no llevan imagen: son instrucciones de importación, no epígrafes. Al
   estar así formateados, duplicaban las figuras del Capítulo 7.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene el documento abierto. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def set_par(i, texto):
    p = d.paragraphs[i]
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ''


n = replace_everywhere(d,
    'workflows/Flujo 2 — Chatbot IA PRODUCCION.json',
    'workflows/Flujo 2 — Chatbot Omnicanal IA PRODUCCION.json')
print('ruta de workflow corregida: %d' % n)

for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t.startswith('Figura A1: Canvas del Flujo 1 PRODUCCIÓN'):
        set_par(i, 'Para importar el workflow del Flujo 1 en su variante de producción '
                   '(Figura A1) en n8n: Workflows → Import from file.')
        print('Anexo G, A1 reformulado')
    elif t.startswith('Figura A2: Canvas del Flujo 2 PRODUCCIÓN'):
        set_par(i, 'El workflow del Flujo 2 en su variante de producción (Figura A2) se importa '
                   'por el mismo procedimiento, y requiere credenciales de Telegram Bot, Gmail y '
                   'WhatsApp Business API para poder activarse.')
        print('Anexo G, A2 reformulado')

d.save(RUTA)

d2 = Document(RUTA)
print()
for i, p in enumerate(d2.paragraphs):
    t = p.text.strip()
    if 'PRODUCCION.json' in t or t.startswith(('Para importar el workflow', 'El workflow del Flujo 2')):
        print('P%-4d %s' % (i, t[:165]))
