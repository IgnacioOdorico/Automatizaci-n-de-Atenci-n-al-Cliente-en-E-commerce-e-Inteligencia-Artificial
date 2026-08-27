# -*- coding: utf-8 -*-
"""Hallazgos #21, #25, #26 y el rótulo del tablero (#9)."""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
P = d.paragraphs


def set_par(i, texto, estilo=None):
    p = P[i]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(texto)
    if estilo:
        try:
            p.style = estilo
        except KeyError:
            pass


# ============ #25 — el 181 % es nominal, no crecimiento real ============
replace_in_paragraph(P[69],
    'En efecto, en 2024 la facturación del sector creció un 181% interanual y, para el 54% de '
    'las empresas relevadas, el canal online ya representa más del 10% de sus ventas —frente al '
    '50% en 2023— (Cámara Argentina de Comercio Electrónico [CACE], 2025).',
    'En efecto, en 2024 la facturación nominal del sector, medida en pesos corrientes, creció un '
    '181 % interanual; dado que ese período transcurrió bajo inflación de tres dígitos, la cifra '
    'no es interpretable por sí sola como expansión real del volumen comerciado y se consigna '
    'únicamente como indicador de la magnitud nominal del mercado. El dato que sí describe la '
    'penetración del canal es de naturaleza no monetaria: para el 54 % de las empresas relevadas '
    'el canal online ya representa más del 10 % de sus ventas, frente al 50 % en 2023 (Cámara '
    'Argentina de Comercio Electrónico [CACE], 2025).')

# ============ #21 — cita colgante a OpenAI Status ============
replace_in_paragraph(P[270],
    'el TMR del Flujo 2 está dominado por la latencia de la API externa, que oscila típicamente '
    'entre 1 y 3 segundos según OpenAI Status (2025) y no depende del entorno local.',
    'el TMR del Flujo 2 está dominado por la latencia de la API externa y no por el entorno '
    'local. Esa latencia no se toma de una fuente externa sino que se estima a partir de las '
    'propias mediciones de este trabajo: el TMR observado se ubica entre 1,47 s sobre el canal '
    'con entrega local y 3,07 s sobre el canal Telegram real, y la diferencia entre ambos acota '
    'el componente atribuible a la red del canal.')

replace_in_paragraph(P[337],
    'el TMR está dominado por la latencia de la API de OpenAI (~1–3 s), que no depende del '
    'hardware local.',
    'el TMR está dominado por la latencia de la llamada de inferencia a la API de OpenAI, que no '
    'depende del hardware local: el rango observado en este trabajo va de 1,47 s a 3,07 s según '
    'el canal.')

# ============ #26 — el Resumen estaba partido en mitad de oración ============
set_par(15,
    'Los resultados obtenidos muestran un MTTD promedio de 0,009 segundos y un MTTR promedio de '
    '0,054 segundos, para un tiempo total end-to-end de 0,063 segundos sobre 50 órdenes medidas. '
    'Ese valor resulta aproximadamente 780 veces menor que el baseline de atención manual, '
    'cronometrado por el propio equipo sobre diez órdenes con una media de 49,13 segundos '
    '(IC 95 %: 43,3 s a 55,0 s); el factor debe leerse como una cota superior, por provenir el '
    'término automatizado de un entorno de laboratorio. El TMR promedio del chatbot fue de 1,47 '
    'segundos sobre un corpus de 150 interacciones y de 3,07 segundos sobre 45 interacciones del '
    'canal Telegram real, con disponibilidad continua, y la precisión de clasificación alcanzó el '
    '92,7 % (IC 95 % de Wilson: 87,3 % a 95,9 %). Estos valores permiten confirmar la hipótesis '
    'de que la automatización reduce significativamente los tiempos operativos del ciclo '
    'post-venta respecto del proceso manual de referencia.')
set_par(16, '')

# ============ #9 — el tablero se rotula multicanal ============
n = replace_everywhere(d, 'Chatbot Omnicanal', 'Chatbot Multicanal')
print('Rótulos de tablero actualizados: %d' % n)

# ============ #26 — ortografía ============
for mal, bien in (('notificaciónes', 'notificaciones'),
                  ('restricciónes', 'restricciones'),
                  ('Notificaciónes', 'Notificaciones')):
    k = replace_everywhere(d, mal, bien)
    print('  %-16s -> %-16s  %d' % (mal, bien, k))

d.save(RUTA)

# ===== Verificación =====
d2 = Document(RUTA)
for i in (15, 16, 69, 270, 337):
    print('===== P%d =====' % i)
    print(d2.paragraphs[i].text[:800])
    print()
buscar(d2, r'notificaci[oó]nes|restricci[oó]nes|Omnicanal|OpenAI Status', 'residuos (deben ser 0)')
