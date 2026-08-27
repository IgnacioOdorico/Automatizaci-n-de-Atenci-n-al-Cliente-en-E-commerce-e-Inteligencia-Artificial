# -*- coding: utf-8 -*-
"""Reemplaza imágenes dentro del .docx por número de figura y reajusta el marco.

Figura 1: el diagrama de arquitectura mostraba grafana/grafana:13.0.0 (etiqueta
          que NO existe en el registro) y "5 vistas" (son 6 con v_chatbot_corpus).
Figura 5: flechas de bypass FAQ/GENERAL redibujadas como curvas exteriores, para
          que no se lean como un recuadro que agrupa las ramas.
"""
import zipfile
import shutil
import re
import io
import os
import sys
from PIL import Image
from docx import Document

DOCX = 'docs/TESIS_FINAL_UTN_v6.docx'
NUEVAS = {                       # nº de figura -> archivo PNG
    '1': 'docs/figuras_v6/diagrama_arquitectura.png',
    '5': 'docs/figuras_v6/diagrama_flujo2.png',
}

if os.path.exists(DOCX + '.lock') or any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: hay un archivo de bloqueo en docs/ — cerrá Word primero.')

shutil.copy(DOCX, DOCX + '.bak-figs2')

# qué media le corresponde a cada figura
d = Document(DOCX)
media_de_figura = {}
for i, p in enumerate(d.paragraphs):
    m = re.match(r'Figura (\d+):', p.text.strip())
    if m and m.group(1) in NUEVAS:
        rid = re.search(r'r:embed="(rId\d+)"', d.paragraphs[i-1]._element.xml).group(1)
        media_de_figura[m.group(1)] = (rid, str(d.part.rels[rid].target_part.partname).lstrip('/'))

zin = zipfile.ZipFile(DOCX, 'r')
doc = zin.read('word/document.xml').decode('utf-8')

blobs = {}
for fig, ruta in NUEVAS.items():
    rid, media = media_de_figura[fig]
    blobs[media] = open(ruta, 'rb').read()
    ratio = Image.open(ruta).size[0] / Image.open(ruta).size[1]
    partes = re.split(r'(<w:drawing>.*?</w:drawing>)', doc, flags=re.S)
    for k, parte in enumerate(partes):
        if parte.startswith('<w:drawing>') and ('r:embed="%s"' % rid) in parte:
            e = re.search(r'<wp:extent cx="(\d+)" cy="(\d+)"', parte)
            cx, cy = int(e.group(1)), int(e.group(2))
            partes[k] = re.sub(r'cx="%d" cy="%d"' % (cx, cy),
                               'cx="%d" cy="%d"' % (cx, int(round(cx / ratio))), parte)
    doc = ''.join(partes)
    print('Figura %s -> %s  (%s)' % (fig, media, ruta.split('/')[-1]))

TMP = DOCX + '.tmpf'
zout = zipfile.ZipFile(TMP, 'w', zipfile.ZIP_DEFLATED)
for item in zin.infolist():
    data = zin.read(item.filename)
    if item.filename in blobs:
        data = blobs[item.filename]
    elif item.filename == 'word/document.xml':
        data = doc.encode('utf-8')
    zout.writestr(item, data)
zin.close()
zout.close()
os.replace(TMP, DOCX)

# ===== verificación =====
print()
d2 = Document(DOCX)
print('zip:', 'integro' if zipfile.ZipFile(DOCX).testzip() is None else 'CORRUPTO')
print('parrafos %d | tablas %d | imagenes %d\n'
      % (len(d2.paragraphs), len(d2.tables),
         sum(1 for r in d2.part.rels.values() if 'image' in r.reltype)))
for i, p in enumerate(d2.paragraphs):
    t = p.text.strip()
    if re.match(r'Figura [1-9]:', t):
        xml = d2.paragraphs[i-1]._element.xml
        rid = re.search(r'r:embed="(rId\d+)"', xml).group(1)
        im = Image.open(io.BytesIO(d2.part.rels[rid].target_part.blob))
        cx, cy = map(int, re.search(r'<wp:extent cx="(\d+)" cy="(\d+)"', xml).groups())
        ok = abs((cx/cy) - (im.size[0]/im.size[1])) < 0.01
        print('  %-10s %-12s marco %5.2f x %5.2f cm  %s'
              % (t[:9], str(im.size), cx/360000, cy/360000, 'OK' if ok else 'DEFORMADA'))
