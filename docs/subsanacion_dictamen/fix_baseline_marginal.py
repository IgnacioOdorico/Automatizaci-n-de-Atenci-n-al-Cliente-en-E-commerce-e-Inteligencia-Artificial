# -*- coding: utf-8 -*-
"""CORRECCIÓN DE FONDO — qué mide realmente el baseline de 49,13 s.

Verificado contra la base viva: las 12 filas e4_manual comparten el MISMO
received_at (2026-08-14 16:46:53.931585), porque preparar_e4.sql las inserta
en un solo lote. Por lo tanto notified_at - received_at NO es la latencia de
una orden: es el tiempo de espera en la cola, y crece monótonamente de 72,4 s
a 533,9 s (media 313,37; sd 154,57).

Los 49,13 s del cronómetro son el CYCLE TIME MARGINAL por orden (T1+T2+T3).
Validación cruzada: los incrementos entre órdenes consecutivas en la base dan
51,28 s (sd 5,22). La diferencia de 2,15 s son los huecos entre terminar una
orden y arrancar la siguiente. Dos instrumentos independientes, 4 % de brecha.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


# ============================================================
#  §3.5.5 — qué mide el número, dicho antes de darlo
# ============================================================
i = idx('Criterio de análisis. Se reportan media, mediana')
insertar_bloque(d, i, [
    ('Qué magnitud mide el valor obtenido. Esta precisión es determinante para la comparación '
     'del Capítulo 5 y por eso se declara aquí. El cronómetro registra el tiempo de trabajo '
     'efectivo que el operador dedica a una orden, es decir el costo marginal de procesar una '
     'orden adicional. No mide la latencia que percibe el cliente, que en un proceso manual '
     'depende de cuántas órdenes haya por delante en la cola y no de la velocidad del operador. '
     'Ambas magnitudes son legítimas y responden a preguntas distintas: la primera describe la '
     'capacidad de proceso, la segunda la experiencia de espera.', 'Body Text'),

    ('Las doce órdenes se insertaron en la base mediante una única sentencia de preparación, de '
     'modo que todas comparten la misma marca de recepción. En consecuencia, la diferencia entre '
     'notified_at y received_at que queda registrada para cada una no es su latencia individual '
     'sino su tiempo de permanencia en la cola, que crece de manera monótona a medida que el '
     'operador avanza. Ese registro se conserva y se reporta en la Sección 5.1.4 porque describe '
     'un fenómeno real —el encolamiento—, pero no es el valor que se contrasta contra el pipeline '
     'automatizado. El término de comparación es el tiempo marginal cronometrado.', 'Body Text'),
])

d.save(RUTA)

# ============================================================
#  §5.1.4 — resultados: marginal, validación cruzada y cola
# ============================================================
d = Document(RUTA)
i = idx('El tiempo medio de procesamiento manual de una orden resultó de 49,13 s')
p = d.paragraphs[i]
NUEVO = (
    'El tiempo marginal de procesamiento manual de una orden resultó de 49,13 s, con un intervalo '
    'de confianza al 95 % de [43,30 s; 54,95 s] calculado por la distribución t de Student con '
    'nueve grados de libertad. El coeficiente de variación fue del 16,6 %, dispersión baja que '
    'indica que el procedimiento se ejecutó de manera estable a lo largo de la serie. Conforme se '
    'estableció en la Sección 3.5.5, este valor mide el costo de procesar una orden adicional y no '
    'la latencia percibida por el cliente.')
p.runs[0].text = NUEVO
for r in p.runs[1:]:
    r.text = ''

i = idx('Corresponden dos advertencias sobre la interpretación de este valor')
insertar_bloque(d, i - 1, [
    ('El valor admite una validación cruzada que conviene explicitar, porque proviene de un '
     'instrumento independiente del cronómetro. Las marcas temporales que el operador fue '
     'escribiendo en la base durante la medición permiten calcular el intervalo entre '
     'notificaciones de órdenes consecutivas, que es la misma magnitud marginal medida por otra '
     'vía: ese cálculo arroja 51,28 s de media (desvío 5,22 s). La diferencia de 2,15 s respecto '
     'de los 49,13 s cronometrados corresponde a los intervalos muertos entre terminar una orden '
     'y comenzar la siguiente, que el cronómetro no contabiliza y el reloj de la base sí. Dos '
     'instrumentos independientes coinciden dentro del 4 %, lo que respalda la validez de la '
     'medición.', 'Body Text'),

    ('El registro en la base aporta además un resultado complementario sobre el encolamiento. '
     'Como las doce órdenes ingresaron en un mismo lote (Sección 3.5.5), el tiempo transcurrido '
     'entre la recepción del lote y la notificación de cada orden crece de forma monótona: 72,4 s '
     'para la primera atendida y 533,9 s —ocho minutos y catorce segundos— para la última, con una '
     'media de 313,37 s y un desvío de 154,57 s (Tabla I.2). Corresponde subrayar que esta cifra '
     'no es una propiedad del proceso manual sino del tamaño del lote elegido, y que por lo tanto '
     'no se la utiliza como término de comparación: un lote mayor la aumentaría sin que el '
     'operador trabajase más lento. Su valor es ilustrativo y apunta a la razón de fondo por la '
     'que la automatización importa en este dominio: el pipeline automatizado procesó 50 órdenes '
     'consecutivas sin que el tiempo de la última se degradara respecto de la primera, mientras '
     'que en el proceso manual la espera del último cliente creció un factor de siete frente a la '
     'del primero.', 'Body Text'),
])

d.save(RUTA)
print('§3.5.5 y §5.1.4 corregidas')

d2 = Document(RUTA)
for i in range(idx('5.1.4 Baseline de atención manual'),
               idx('5.1.4 Baseline de atención manual') + 9):
    print('%4d [%-14s] %s' % (i, d2.paragraphs[i].style.name, d2.paragraphs[i].text.strip()[:120]))
