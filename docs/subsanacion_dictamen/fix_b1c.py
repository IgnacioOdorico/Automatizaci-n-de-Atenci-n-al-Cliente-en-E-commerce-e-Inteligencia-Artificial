# -*- coding: utf-8 -*-
"""Bloqueante #3 — reconciliación de la prueba de carga.

Verdad de campo (verificada contra la base viva y los manifiestos de E1):
  E1.a  50 órdenes secuenciales espaciadas 2 s → 35 confirmed / 15 no_stock / 0 error
  E1.b  6 rondas × 20 requests simultáneas contra stock=5 → 5 confirmadas SIEMPRE,
        0 sobreventa, 0 stock negativo, 49/120 sin procesar
  Las ORD-LOAD-001..020 que el documento llamaba "prueba de carga" son SEED sintético.
"""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
P = d.paragraphs


def set_par(i, texto, estilo=None):
    p = P[i]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(texto)
    if estilo:
        try:
            p.style = estilo
        except KeyError:
            pass


# ============ §4.6.2 — diseño de las dos pruebas ============
set_par(277, '4.6.2 Pruebas de carga y de concurrencia', 'Heading 3')

set_par(278,
    'Se diseñaron dos ensayos distintos sobre el Flujo 1, con objetivos diferentes y '
    'poblaciones separadas. El primero (identificado como E1.a en el repositorio) evalúa el '
    'comportamiento del pipeline bajo un régimen de llegada sostenido pero no simultáneo: se '
    'envían 50 órdenes secuenciales al webhook, espaciadas 2 segundos entre sí, con una '
    'composición de 35 órdenes con stock suficiente y 15 que fuerzan deliberadamente la rama '
    'de stock insuficiente. Esta corrida es la que provee las mediciones de MTTD, MTTR y '
    'tiempo end-to-end reportadas en la Sección 5.1.')

set_par(279,
    'El segundo ensayo (E1.b) evalúa una propiedad distinta y no trivial: la atomicidad de la '
    'transacción de descuento de stock bajo concurrencia real. Se disparan 20 solicitudes HTTP '
    'simultáneas contra un producto cuyo stock inicial se fija deliberadamente en 5 unidades, y '
    'el ensayo se repite durante 6 rondas, totalizando 120 órdenes. El criterio de falla está '
    'definido de antemano y es binario: si en alguna ronda el sistema confirma más de 5 órdenes, '
    'existe sobreventa y por lo tanto una condición de carrera en el descuento de inventario.',
    'Body Text')

set_par(280,
    'Ambas corridas se ejecutan mediante scripts versionados que dejan un manifiesto JSON con '
    'la semilla utilizada, el número de nodos del flujo al momento de medir, el commit de Git y '
    'las marcas temporales en UTC. El plan de órdenes es determinístico: la misma semilla sobre '
    'el mismo stock inicial reproduce exactamente la misma secuencia de SKU y cantidades.',
    'Body Text')

set_par(281,
    'Las órdenes generadas por ambos ensayos se registran con data_source = \'measured\', lo que '
    'las separa tanto de los datos seed precargados como del baseline manual de la Sección 3.5.5. '
    'Las vistas que alimentan los tableros de resultados filtran por ese valor, de modo que '
    'ninguna cifra reportada en el Capítulo 5 mezcla datos medidos con datos sintéticos.',
    'Body Text')

set_par(282,
    'Los correos de confirmación y de aviso de stock insuficiente generados durante ambas '
    'corridas se verifican en la bandeja del capturador SMTP (ver Figura 8).', 'Body Text')

set_par(283, '', 'Body Text')

# ============ §5.1 — resultados de carga y concurrencia ============
set_par(296, 'Pruebas de carga y de concurrencia', 'Heading 3')

set_par(297,
    'La corrida de carga secuencial (E1.a) procesó las 50 órdenes enviadas sin un solo error de '
    'pipeline y con cobertura completa de las tres marcas temporales: las 50 órdenes registraron '
    'received_at, processed_at y notified_at, incluidas las 15 que siguieron la rama de stock '
    'insuficiente. Sobre esa población se calculan las métricas de la Tabla 5.2.')

set_par(298, 'Tabla 5.3: Resultados de las pruebas de carga y de concurrencia.', 'Body Text')

set_par(299,
    'La prueba de concurrencia (E1.b) arroja dos resultados que conviene reportar por separado, '
    'porque apuntan en direcciones distintas. El primero es concluyente y favorable: en las seis '
    'rondas ejecutadas, contra un stock inicial de 5 unidades y con 20 solicitudes simultáneas en '
    'cada una, el sistema confirmó exactamente 5 órdenes en todas y cada una de las rondas. No se '
    'registró sobreventa en ningún caso, ni quedó ningún producto con stock negativo. La '
    'restricción de integridad a nivel de esquema y el carácter transaccional de la sentencia de '
    'descuento resultaron suficientes para impedir la condición de carrera, que era exactamente '
    'la propiedad que el ensayo se proponía refutar.')

d.save(RUTA)

d2 = Document(RUTA)
for i in range(277, 300):
    t = d2.paragraphs[i].text.strip()
    print('P%-4d [%-16s] %s' % (i, d2.paragraphs[i].style.name, t[:130]))
