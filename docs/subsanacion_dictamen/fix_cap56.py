# -*- coding: utf-8 -*-
"""Capítulos 5 y 6 — hallazgos #6, #7, #8, #23 y las cuatro ocurrencias del 780×."""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
P = d.paragraphs


def set_par(i, texto, estilo=None):
    p = P[i]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(texto)
    if estilo:
        try:
            p.style = estilo
        except KeyError:
            pass


# ===== #8h — tasa de resolución: el texto decía ~70 %, la tabla ~73 % =====
set_par(304,
    'TMR promedio: 1,47 s sobre los 150 mensajes del corpus. Este valor confirma H2a '
    '(TMR < 10 s) con amplio margen.')

set_par(305,
    'Tasa de resolución automática: 73,3 % (110 de 150 mensajes). Si se excluyen los '
    'reclamos —que escalan a ticket por diseño del sistema y no por falla del clasificador—, '
    'la tasa asciende al 100 %: el chatbot resuelve de forma autónoma la totalidad de las '
    'consultas que no requieren intervención humana.')

# ===== #8 y #6f — las precisiones del texto contradecían la Tabla 5.8 =====
set_par(320,
    'La categoría con mayor precisión fue FAQ (97,8 %), seguida de RECLAMO (95,0 %), '
    'ESTADO_PEDIDO (88,9 %) y GENERAL (86,2 %). El patrón se invierte al observar la '
    'exhaustividad: GENERAL alcanza el 100 % (25 de 25), es decir que el modelo no dejó sin '
    'identificar ningún mensaje que perteneciera realmente a esa categoría, pero incorporó a '
    'ella cuatro mensajes provenientes de otras clases. Esa asimetría —exhaustividad máxima '
    'junto con la precisión más baja del conjunto— es consistente con el carácter abierto de '
    'la categoría: al no estar acotada por patrones específicos, GENERAL opera de hecho como '
    'clase residual y absorbe los mensajes que el modelo no consigue asignar con confianza a '
    'las otras tres (ver Figura 9).')

# ===== #6 y #4 — contrastación de hipótesis =====
set_par(327,
    '• H1 se confirma en sus dos términos. En el término comparativo, que es el sustantivo: '
    'el tiempo end-to-end del pipeline (0,063 s; n = 50) resulta un factor de aproximadamente '
    '780 veces menor que el baseline de atención manual medido en 49,13 s (IC 95 % del factor: '
    '687× a 872×). En el criterio operativo secundario: ese mismo tiempo representa el 0,2 % '
    'del umbral de 30 segundos. La Sección 5.4 precisa por qué el factor debe leerse como una '
    'cota superior y no como el valor esperable en producción.')

set_par(328,
    '• H2a se confirma con amplio margen: el TMR del corpus (1,47 s) representa el 14,7 % del '
    'umbral de 10 s, y el del canal Telegram real (3,07 s) el 30,7 %.')

set_par(329,
    '• H2b se confirma con un accuracy global de 92,7 % (139/150), cuyo intervalo de confianza '
    'de Wilson al 95 % —[87,3 %; 95,9 %]— tiene el límite inferior por encima del umbral del '
    '85 %. El procedimiento de cálculo se detalla en el Anexo H.')

# ===== #23 y #6f — "A diferencia de la versión anterior" y "desempeño desigual" =====
replace_in_paragraph(P[330], 'A diferencia de la versión anterior, ninguna',
                     'Ninguna')
replace_in_paragraph(P[330], '[88,2%; 96,3%]', '[87,3 %; 95,9 %]')
replace_in_paragraph(P[330], 'Este desempeño desigual es coherente',
                     'La asimetría entre precisión y exhaustividad de GENERAL es coherente')

# ===== #1 — el 780× ahora con su reserva en la misma oración =====
set_par(332,
    'Los resultados confirman la premisa central del trabajo: la automatización reduce '
    'significativamente los tiempos operativos del ciclo post-venta. La reducción en el tiempo '
    'de procesamiento de órdenes es de un factor aproximado de 780 veces respecto del baseline '
    'de atención manual medido en la Sección 5.1.4 (49,13 s → 0,063 s); ese factor debe leerse '
    'como una cota superior y no como el valor esperable en un despliegue productivo, por dos '
    'razones que actúan en la misma dirección. La primera es que el término automatizado '
    'proviene de un entorno de laboratorio en el que la notificación se entrega a un capturador '
    'SMTP alojado en el mismo host, sin tránsito de correo real ni latencia de servicios de '
    'terceros. La segunda es que el término manual se midió sobre un operador que ya conocía el '
    'procedimiento y que trabajó a partir de una plantilla fija, dos decisiones deliberadamente '
    'conservadoras que reducen el numerador del cociente. Corresponde además señalar que ambos '
    'términos no cubren exactamente el mismo alcance: los 0,063 s miden el intervalo '
    'received_at → notified_at, es decir escrituras sobre la base de datos local, mientras que el '
    'baseline manual cubre la tarea humana completa de lectura, verificación, actualización y '
    'redacción. En atención al cliente, el chatbot responde en 1,47 s en promedio sobre el '
    'corpus evaluado y en 3,07 s sobre el canal Telegram real, con disponibilidad continua.')

# ===== #7 — la limitación (d) negaba el análisis que el trabajo sí hizo =====
set_par(339,
    '(d) Ausencia de contraste inferencial contra un grupo de control: los intervalos de '
    'confianza que este trabajo reporta —Wilson al 95 % para el accuracy de clasificación, y t '
    'de Student para el baseline manual y para las latencias del pipeline— describen la '
    'precisión de las estimaciones propias, pero no constituyen una prueba de hipótesis en '
    'sentido inferencial estricto. El trabajo no ejecuta un contraste formal contra un grupo de '
    'control asignado al azar, y el baseline manual de la Sección 3.5.5 se midió sobre un único '
    'operador, lo que impide estimar la variabilidad entre operadores y descartar el efecto del '
    'evaluador. Una comparación inferencial en sentido pleno exigiría varios operadores '
    'independientes y un diseño experimental que excede el alcance declarado de este trabajo.')

# ===== #1 — Capítulo 6 =====
set_par(343,
    'Los resultados demuestran que la automatización reduce los tiempos de forma sustancial: el '
    'tiempo end-to-end del pipeline de órdenes (0,063 s; n = 50) es aproximadamente 780 veces '
    'menor que el baseline de atención manual medido en 49,13 s sobre diez órdenes cronometradas '
    '(Sección 3.5.5), factor que la Sección 5.4 acota explícitamente como cota superior; y el TMR '
    'del chatbot es de 1,47 s en promedio sobre el corpus evaluado, con disponibilidad continua. '
    'Las tres hipótesis de trabajo (H1, H2a y H2b) fueron confirmadas por los datos '
    'experimentales (ver Sección 5.3).')

set_par(347,
    'Los cinco objetivos específicos fueron cumplidos. OE3 y OE4 se cumplieron por encima de lo '
    'comprometido —7 tablas y 6 vistas frente a las 5 y 5 previstas, y 13 paneles frente a los 7 '
    'enunciados—, en tanto que OE2 se cumplió sobre dos canales efectivamente implementados y '
    'medidos, quedando el canal de correo como diseño documentado en el Capítulo 7.')

# ===== #8 (sobregeneralización) — el ensayo de reproducción nunca se hizo =====
set_par(351,
    'Sobre la arquitectura Docker Compose: la infraestructura basada en cuatro contenedores '
    'demostró ser reproducible en lo esencial. El entorno se levanta con un único comando '
    '(docker compose up -d), las imágenes están fijadas por versión explícita, los datos '
    'persisten entre reinicios mediante volúmenes nombrados, y el esquema, los flujos y los '
    'scripts de medición están versionados en el repositorio. Corresponde aclarar, no obstante, '
    'que no se realizó un ensayo de reproducción con un evaluador externo: el tiempo efectivo de '
    'puesta en marcha por un tercero no fue medido y por lo tanto no se afirma.')

set_par(354,
    'Las principales limitaciones del presente trabajo —entorno de prueba local, datos '
    'simulados, ausencia de un contraste inferencial contra un grupo de control, baseline manual '
    'medido sobre un único operador y dependencia de la API de OpenAI— fueron documentadas en '
    'detalle en la Sección 5.4.1 y en el análisis de amenazas a la validez del Capítulo 3. Estas '
    'limitaciones no invalidan las conclusiones dentro del alcance declarado (prototipo funcional '
    'en entorno controlado), pero deben tenerse en cuenta al extrapolar los resultados a entornos '
    'de producción reales.')

d.save(RUTA)

# ===== Verificación =====
d2 = Document(RUTA)
for i in (304, 305, 320, 327, 328, 329, 332, 339, 343, 347, 351, 354):
    print('===== P%d =====' % i)
    print(d2.paragraphs[i].text)
    print()
print()
buscar(d2, r'88,2|96,3 ?%|versión anterior|desempeño desigual|~70', 'residuos (deben ser 0)')
