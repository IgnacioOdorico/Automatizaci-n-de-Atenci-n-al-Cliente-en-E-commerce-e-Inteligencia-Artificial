# -*- coding: utf-8 -*-
"""Tercer dictamen — recomendadas 7 y 8, la observacion de 6.4, y dos defectos
propios detectados al auditar el documento:

  - El Anexo F repetia integro el Listado de Figuras del frontispicio (misma
    clase de duplicacion que el indice, que el dictamen si observo). El anexo
    dice listar "los workflows activos y probados en el entorno de desarrollo",
    de modo que le corresponden solo las capturas: 4, 6, 7, 8 y 9.
  - Los listados remitian a 4.1, 4.5 y 5.1, secciones que ahora tienen
    subapartados numerados.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def par(pref):
    return d.paragraphs[idx(pref)]


print('=' * 74)
print(' Anexo F: se elimina la duplicacion del Listado de Figuras')
print('=' * 74)
# el listado del frontispicio es T0; el del Anexo F es el segundo con ese encabezado
listados = [t for t in d.tables if t.rows[0].cells[0].text.strip() == 'Figura'
            and t.rows[0].cells[2].text.strip() == 'Sección']
assert len(listados) == 2, 'esperaba 2 listados de figuras, hay %d' % len(listados)
anexoF = listados[1]
CAPTURAS = {'Figura 4', 'Figura 6', 'Figura 7', 'Figura 8', 'Figura 9'}
borradas = 0
for fila in list(anexoF.rows[1:]):
    if fila.cells[0].text.strip() not in CAPTURAS:
        fila._tr.getparent().remove(fila._tr)
        borradas += 1
print('filas eliminadas del Anexo F: %d  (quedan %d capturas)' % (borradas, len(anexoF.rows) - 1))
n = replace_everywhere(d,
    'Las siguientes figuras corresponden a los workflows activos y probados con datos reales en '
    'el entorno de desarrollo local.',
    'Se listan aquí únicamente las capturas de pantalla del entorno de desarrollo local, es '
    'decir las figuras que registran el sistema tal como quedó ejecutándose. Las restantes '
    'figuras del trabajo son diagramas de elaboración propia y se consignan en el Listado de '
    'Figuras del frontispicio, que no se repite.')
print('nota de alcance del Anexo F: %d' % n)

print()
print('=' * 74)
print(' Listados: remisiones a secciones que ahora estan numeradas')
print('=' * 74)
for t in d.tables:
    if t.rows[0].cells[0].text.strip() in ('Figura', 'Tabla') and len(t.rows[0].cells) == 3 \
            and t.rows[0].cells[2].text.strip() == 'Sección':
        for i, fila in enumerate(t.rows):
            if i == 0:
                continue
            sec = fila.cells[2].text.strip()
            nuevo = {'4.1': '4.1.1', '4.5': '4.5.1', '5.1': None}.get(sec, sec)
            if sec == '5.1':
                rot = fila.cells[0].text.strip()
                nuevo = {'Tabla 5.1': '5.1.1', 'Tabla 5.2': '5.1.2',
                         'Tabla 5.3': '5.1.3'}.get(rot, '5.1')
            if nuevo and nuevo != sec:
                set_cell(t, i, 2, nuevo)
                print('  %-11s %s -> %s' % (fila.cells[0].text.strip(), sec, nuevo))
# la descripcion de la Tabla 4.8 en el listado quedo corta frente a su epigrafe
n = replace_everywhere(d, 'Paneles del dashboard de Grafana',
                          'Paneles de los dashboards de Grafana')
print('  descripcion de la Tabla 4.8 sincronizada: %d' % n)

print()
print('=' * 74)
print(' REC 7 — la Seccion 2.4 declara su funcion; el Capitulo 6 cierra el circuito')
print('=' * 74)
n = replace_everywhere(d,
    '2.4 Infraestructura: contenedores, bases de datos relacionales y visualización',
    '2.4 Infraestructura: fundamentos conceptuales y justificación del stack')
print('titulo de 2.4: %d' % n)
insert_paragraph_after(par('2.4 Infraestructura: fundamentos'),
    'Corresponde declarar de antemano la función de esta sección, porque difiere de la de las '
    'tres anteriores. Las Secciones 2.1 a 2.3 construyen los conceptos que el trabajo necesita '
    'para formular sus hipótesis; esta, en cambio, cumple una función doble: expone las nociones '
    'conceptuales que gobiernan la infraestructura —reproducibilidad computacional, modelo '
    'relacional y observabilidad operativa— y, junto a cada una, la decisión de herramienta que '
    'el trabajo tomó apoyándose en ella. Esa segunda mitad es de naturaleza metodológica antes '
    'que teórica, y podría haberse ubicado en el Capítulo 3. Se la mantiene aquí para que cada '
    'elección quede junto al concepto que la fundamenta; el detalle operativo del stack se '
    'consigna en la Sección 3.4 y su configuración concreta en el Capítulo 4.',
    estilo='First Paragraph')
print('nota de funcion de 2.4: insertada')

insert_paragraph_after(par('Sobre las métricas MTTD/MTTR/TMR:'),
    'Sobre el retorno al marco teórico: las tres nociones que el Capítulo 2 puso a trabajar '
    'admiten una lectura a la luz de lo medido. La primera es la escala de madurez de proceso de '
    'van der Aalst (2016), donde la automatización es un nivel intermedio entre la documentación '
    'y la gobernanza: el sistema construido alcanza el nivel de automatización —el ciclo se '
    'ejecuta sin intervención humana y queda instrumentado— pero no el de gobernanza, porque no '
    'incorpora control de admisión, política de reintentos ni gestión del caso de excepción, y '
    'la prueba de concurrencia lo mostró con precisión: 40,8 % de las órdenes quedaron sin '
    'procesar sin que ningún mecanismo lo advirtiera. La segunda es la dimensión de capacidad de '
    'respuesta (responsiveness) del modelo SERVQUAL (Zeithaml et al., 1990), que el TMR '
    'operacionaliza: los 1,47 s del canal simulado y los 3,07 s de Telegram saturan esa '
    'dimensión, pero SERVQUAL mide cinco y este trabajo instrumentó una sola. La fiabilidad y la '
    'empatía —que dependen de que la respuesta sea correcta y pertinente, no de que sea rápida— '
    'quedaron fuera de la medición, y es exactamente allí donde se ubica el riesgo declarado en '
    'la Sección 5.2.1 sobre el contenido de las respuestas de tipo FAQ. La tercera es la '
    'transferibilidad de MTTD y MTTR discutida en la Sección 2.1.3: la transferencia se sostuvo '
    'a nivel estructural, como se anticipó, pero la medición confirmó también su límite '
    'semántico. Un MTTD de 0,009 segundos no informa sobre detección de nada, porque no hay '
    'falla que detectar en un flujo nominal; informa sobre la distancia entre dos escrituras. '
    'Métricas propias del dominio comercial —order cycle time, order fulfillment lead time— '
    'describirían el mismo intervalo sin arrastrar una promesa que el fenómeno no tiene.',
    estilo='Body Text')
print('cierre del circuito teorico en 6.3: insertado')

print()
print('=' * 74)
print(' 6.4 — las limitaciones se enuncian, no se delegan')
print('=' * 74)
p64 = par('Las principales limitaciones del presente trabajo')
n = replace_in_paragraph(p64,
    'Las principales limitaciones del presente trabajo —entorno de prueba local, datos '
    'simulados, ausencia de un contraste inferencial contra un grupo de control, baseline manual '
    'medido sobre un único operador y dependencia de la API de OpenAI— fueron documentadas en '
    'detalle en la Sección 5.4.1 y en el análisis de amenazas a la validez del Capítulo 3.',
    'Las limitaciones del trabajo se documentan en detalle en la Sección 5.4.1 y en el análisis '
    'de amenazas a la validez del Capítulo 3. Se enuncian aquí, para que el lector que llega a '
    'las conclusiones las encuentre y no deba reconstruirlas: (i) el entorno es de laboratorio '
    'local y los datos son simulados, de modo que el término automatizado de toda comparación '
    'proviene de condiciones más favorables que las de un despliegue productivo; (ii) el '
    'baseline manual se midió sobre un único operador que ya conocía el procedimiento, lo que '
    'acota su representatividad aun cuando la decisión reduzca el factor reportado; (iii) el '
    'contraste entre ambas series se ejecutó sobre poblaciones que miden costo marginal de '
    'procesamiento y no latencia percibida, asimetría de constructo que la significación '
    'estadística no resuelve; (iv) el corpus de evaluación del clasificador es de 150 mensajes '
    'sobre cuatro categorías, construido por el propio equipo, y su distribución responde a un '
    'supuesto de diseño y no a tráfico observado; (v) no se evaluó la corrección del contenido '
    'de las respuestas entregadas al cliente, omisión que resulta material en las consultas de '
    'tipo FAQ porque el contexto de la base de conocimiento no llega al prompt; (vi) bajo '
    'concurrencia el sistema pierde el 40,8 % de las órdenes sin perder integridad, de modo que '
    'las cifras de desempeño valen para régimen secuencial; y (vii) el Flujo 2 depende de un '
    'proveedor externo de inferencia, lo que introduce a la vez un punto único de falla y la '
    'transferencia internacional de datos analizada en la Sección 2.6.')
print('6.4 enumerada: %d' % n)

print()
print('=' * 74)
print(' REC 8 — locators de CACE y de Jurafsky y Martin')
print('=' * 74)
n = replace_everywhere(d,
    'Cámara Argentina de Comercio Electrónico. (2025). Estudio anual de comercio electrónico en '
    'Argentina. CACE. https://www.cace.org.ar/estadisticas',
    'Cámara Argentina de Comercio Electrónico. (2025). Estudio Anual de Comercio Electrónico '
    '2024. CACE. Recuperado el 28 de agosto de 2026, de '
    'https://cace.org.ar/pages/biblioteca-de-estudios')
print('CACE: %d' % n)
n = replace_everywhere(d,
    'Jurafsky, D., & Martin, J. H. (2024). Speech and language processing (3.ª ed., borrador). '
    'Stanford University. https://web.stanford.edu/~jurafsky/slp3/',
    'Jurafsky, D., & Martin, J. H. (2024). Speech and language processing: An introduction to '
    'natural language processing, computational linguistics, and speech recognition with '
    'language models (3.ª ed., borrador en revisión permanente). Stanford University. '
    'Recuperado el 28 de agosto de 2026, de https://web.stanford.edu/~jurafsky/slp3/')
print('Jurafsky y Martin: %d' % n)

d.save(RUTA)
print()
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
print()
print('=== listados de figuras ===')
for t in d2.tables:
    if t.rows[0].cells[0].text.strip() == 'Figura':
        print('  %d filas: %s' % (len(t.rows) - 1,
                                  ', '.join(r.cells[0].text.strip() for r in t.rows[1:])))
TXT = '\n'.join(p.text for p in d2.paragraphs) + '\n' + '\n'.join(
    c.text for t in d2.tables for r in t.rows for c in r.cells)
print()
for rot, pat in [('cace.org.ar/estadisticas', 'cace.org.ar/estadisticas'),
                 ('biblioteca-de-estudios', 'biblioteca-de-estudios'),
                 ('Recuperado el 28 de agosto de 2026', 'Recuperado el 28 de agosto de 2026'),
                 ('van der Aalst en el Cap. 6', 'escala de madurez de proceso de van der Aalst'),
                 ('SERVQUAL en el Cap. 6', 'dimensión de capacidad de respuesta (responsiveness)'),
                 ('6.4 enumerada', 'Se enuncian aquí, para que el lector')]:
    print('  %-34s %d' % (rot, TXT.count(pat)))
print()
print('parrafos: %d | tablas: %d' % (len(d2.paragraphs), len(d2.tables)))
