# -*- coding: utf-8 -*-
"""Tercer dictamen — obligatoria 4: la evaluacion cualitativa prometida en 3.2.

El dictamen ofrece dos salidas: ejecutar la evaluacion o eliminar la afirmacion.
Se toma la segunda, y no de manera pasiva: en lugar de borrar la frase, se
declara que el diseño fue cuantitativo, se dice por que la evaluacion de
contenido no se ejecuto y donde queda registrado el riesgo que eso deja abierto.

Y se repone en su lugar una afirmacion de triangulacion que SI es verdadera, y
que el trabajo ya tenia sin haberla nombrado: el baseline manual se midio con
dos instrumentos independientes que convergen dentro del 4 % (cronometro por
fases: 49,13 s; incrementos registrados en la base: 51,28 s), y el ground truth
del clasificador se valido con un anotador independiente (kappa de Cohen 0,919).
Eso es triangulacion de fuentes de evidencia en el sentido de Yin (2018).
"""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def reescribir(pref, texto):
    p = d.paragraphs[idx(pref)]
    p.runs[0].text = texto
    for r in p.runs[1:]:
        r.text = ''
    return p


p = reescribir('Se combinaron métricas cuantitativas',
    'El diseño de medición es cuantitativo: MTTD, MTTR y TMR sobre marcas temporales registradas '
    'automáticamente, y exactitud de clasificación contra un conjunto de etiquetas de referencia. '
    'Corresponde declarar con precisión qué se evaluó y qué no, porque el alcance de lo medido '
    'condiciona la lectura de todo el Capítulo 5. No se evaluó la calidad del contenido de las '
    'respuestas que el chatbot entrega al cliente: no se definió una rúbrica de coherencia ni de '
    'corrección factual, no se seleccionó una muestra de respuestas para juzgar y no se reporta '
    'ningún resultado sobre esa dimensión. La omisión es consciente y su consecuencia se declara '
    'en las Secciones 5.2.1 y 6.4: el trabajo puede afirmar que el sistema clasifica la intención '
    'con determinada exactitud y que responde en determinado tiempo, pero no puede afirmar que la '
    'respuesta entregada sea correcta, y la afirmación es especialmente frágil en las consultas '
    'de tipo FAQ, donde el contexto de la base de conocimiento no llega al modelo (Sección 4.4.3). '
    'La evaluación de la corrección del contenido queda planteada como línea futura en el '
    'Capítulo 7, con la rúbrica y el procedimiento que requeriría.')
insert_paragraph_after(p,
    'Lo que el trabajo sí practica es triangulación de fuentes de evidencia, en el sentido en que '
    'Yin (2018) la recomienda para el estudio de caso: que un dato relevante no descanse sobre un '
    'único instrumento. Se aplica en dos puntos, ambos verificables en el Capítulo 5. El primero '
    'es el baseline de atención manual, medido simultáneamente por dos vías independientes —un '
    'cronómetro por fases operado sobre la tarea, que arroja 49,13 s por orden, y los incrementos '
    'entre notificaciones consecutivas registrados por la propia base de datos, que arrojan '
    '51,28 s—, cuya convergencia dentro del 4 % es la que sostiene la validez de la cifra que '
    'después se compara contra el pipeline. El segundo es el conjunto de etiquetas de referencia '
    'del clasificador, construido por el equipo y contrastado a ciegas por un evaluador '
    'independiente sobre una submuestra aleatoria, con un κ de Cohen de 0,919 (Sección 3.5.3). En '
    'ambos casos la segunda fuente pudo haber refutado a la primera y no lo hizo; ese es el valor '
    'del procedimiento y también su límite, porque ninguna de las dos triangula la dimensión '
    'cualitativa que este trabajo dejó sin medir.',
    estilo='Body Text')
print('3.2 reescrita y triangulacion real declarada')

# la linea futura del Capitulo 7 debe recoger la rubrica que 3.2 remite
n = replace_everywhere(d,
    'Incorporación de ejemplos etiquetados al prompt y cableado del contexto de FAQ que el flujo '
    'ya construye, para cuantificar la mejora sobre el desempeño zero-shot medido en este '
    'trabajo.',
    'Incorporación de ejemplos etiquetados al prompt y cableado del contexto de FAQ que el flujo '
    'ya construye, para cuantificar la mejora sobre el desempeño zero-shot medido en este '
    'trabajo. Ese cableado debe medirse junto con lo que hoy falta: una evaluación de la '
    'corrección del contenido de las respuestas, que este trabajo no ejecutó (Sección 3.2). El '
    'procedimiento que se propone es el siguiente: extraer de la tabla interactions las '
    'respuestas efectivamente entregadas, restringir la muestra a los mensajes clasificados como '
    'FAQ, y hacerlas juzgar por dos evaluadores independientes contra las veintitrés entradas de '
    'la base de conocimiento transcriptas en el Anexo D, sobre una rúbrica de tres niveles '
    '—respuesta consistente con la política de la tienda, respuesta genérica pero no '
    'contradictoria, y respuesta que contradice la política vigente—, reportando el acuerdo entre '
    'ambos evaluadores con el mismo coeficiente κ que se empleó para el conjunto de etiquetas de '
    'referencia. La comparación de esa medición antes y después del cableado del contexto es la '
    'que permitiría atribuirle una mejora, y no la mera inspección del prompt.')
print('linea futura del Capitulo 7 con la rubrica: %d' % n)

d.save(RUTA)
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
TXT = '\n'.join(p.text for p in d2.paragraphs)
print()
for rot, pat, esperado in [
    ('promesa vieja de coherencia', 'con evaluación cualitativa de coherencia de las respuestas', 0),
    ('declaracion de lo no medido', 'No se evaluó la calidad del contenido de las respuestas', 1),
    ('triangulacion real', 'triangulación de fuentes de evidencia', 1),
    ('rubrica en el Cap. 7', 'rúbrica de tres niveles', 1),
]:
    n = TXT.count(pat)
    print('  %-30s %d  %s' % (rot, n, 'OK' if n == esperado else '<-- REVISAR'))
print()
i = next(k for k, p in enumerate(d2.paragraphs) if p.text.strip().startswith('3.2 Enfoque'))
for k in range(i, i + 7):
    t = d2.paragraphs[k].text.strip()
    if t:
        print('P%-4d %s' % (k, t[:190]))
