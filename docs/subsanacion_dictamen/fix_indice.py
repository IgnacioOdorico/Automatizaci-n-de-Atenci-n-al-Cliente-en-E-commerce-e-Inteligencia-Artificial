# -*- coding: utf-8 -*-
"""#26 — el índice omitía el glosario, los listados y los anexos, y no reflejaba §2.5/§2.6.
El índice es una lista manual de párrafos Compact, no un campo TOC de Word."""
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def set_par(doc, i, texto):
    p = doc.paragraphs[i]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(texto)


# --- 1) Anexos desagregados, glosario y listados (después de "Capítulo 9: Anexos") ---
ANEXOS = [
    ('Anexo A: Resumen del schema de base de datos', 'Compact'),
    ('Anexo B: Configuración Docker Compose', 'Compact'),
    ('Anexo C: Catálogo de productos seed', 'Compact'),
    ('Anexo D: FAQ predefinidas', 'Compact'),
    ('Anexo E: Comandos de testing', 'Compact'),
    ('Anexo F: Lista de figuras — implementación de desarrollo', 'Compact'),
    ('Anexo G: Propuesta de arquitectura de producción', 'Compact'),
    ('Anexo H: Prompt de sistema del clasificador de intenciones', 'Compact'),
    ('Anexo I: Baseline de atención manual — tiempos cronometrados', 'Compact'),
    ('Anexo J: Corpus de evaluación y procedimiento de cálculo estadístico', 'Compact'),
    ('Listado de Figuras', 'Compact'),
    ('Listado de Tablas principales', 'Compact'),
    ('Glosario de siglas y acrónimos', 'Compact'),
]
insertar_bloque(d, 58, ANEXOS)

# --- 2) el índice mantiene su propia caja: se revierte a mayúscula inicial ---
set_par(d, 39, 'Capítulo 4: Desarrollo del Estudio')

# --- 3) §2.3 cambió de título ---
set_par(d, 30, '2.3 Comunicación multicanal en e-commerce')

# --- 4) las dos secciones nuevas del Capítulo 2 ---
insertar_bloque(d, 31, [
    ('2.5 Estado del arte', 'Compact'),
    ('2.6 Tratamiento de datos personales y marco legal aplicable', 'Compact'),
])

d.save(RUTA)

d2 = Document(RUTA)
print('=== ÍNDICE FINAL ===')
for i in range(18, 78):
    t = d2.paragraphs[i].text.strip()
    if t:
        print('%4d %s' % (i, t))
    if t == 'Glosario de siglas y acrónimos':
        break
