# -*- coding: utf-8 -*-
"""Lote final: hallazgos #14, #15, #16, #17-menor, #22 y las líneas futuras faltantes."""
import sys
import copy
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
P = d.paragraphs


def set_par(i, texto, estilo=None):
    p = P[i]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(texto)
    if estilo:
        try:
            p.style = estilo
        except KeyError:
            pass


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def clonar(tabla, ri):
    nueva = copy.deepcopy(tabla.rows[ri]._tr)
    tabla.rows[ri]._tr.addnext(nueva)


# ===== #22 — Tabla 3.2: versiones fijas, no "latest" =====
t32 = d.tables[5]
set_cell(t32, 1, 1, 'n8nio/n8n:2.12.2')
set_cell(t32, 2, 1, 'postgres:15.13-alpine')
set_cell(t32, 3, 0, 'Docker Compose')
set_cell(t32, 3, 1, 'Formato de archivo Compose v2 (especificación vigente)')
set_cell(t32, 5, 1, 'axllent/mailpit:v1.29.6')
set_cell(t32, 6, 1, 'grafana/grafana:13.0.0')
set_cell(t32, 4, 1, 'gpt-4o-mini (API de OpenAI)')

# ===== #15 — Tabla 3.1: conteos alineados con lo entregado =====
t31 = d.tables[4]
set_cell(t31, 1, 1, 'Modelado de 7 tablas y 6 vistas de métricas en PostgreSQL')
set_cell(t31, 4, 1, 'Workflow de 18 nodos funcionales: 3 canales, IA, tickets')
set_cell(t31, 5, 1, '13 paneles con métricas MTTD, MTTR y TMR en dos tableros')
set_cell(t31, 6, 1, '10 pruebas funcionales, 1 corrida de carga y 1 prueba de concurrencia')

# ===== #14 — Tabla 4.3: la sexta vista =====
t43 = d.tables[9]
clonar(t43, 5)
set_cell(t43, 6, 0, 'v_chatbot_corpus')
set_cell(t43, 6, 1, 'Aísla el corpus de 150 mensajes con etiquetas de referencia, por ventana temporal')
set_cell(t43, 6, 2, 'TMR e intención predicha sobre el corpus evaluado')

# ===== #17-menor — alinear F1–F5 con PF-01–PF-05 =====
t15 = d.tables[15]
set_cell(t15, 4, 1, 'SKU inválido')
set_cell(t15, 4, 2, 'POST con product_sku inexistente en el catálogo')
set_cell(t15, 4, 3, 'Rechazo con error controlado')
set_cell(t15, 5, 1, 'Orden duplicada')
set_cell(t15, 5, 2, 'POST con un order_number ya registrado')
set_cell(t15, 5, 3, 'Rechazo por restricción de unicidad')

# ===== Anexo G: las rutas reales de los workflows de producción =====
tg = d.tables[33]
set_cell(tg, 1, 2, 'workflows/Flujo 1 — Pipeline de Procesamiento de Órdenes PRODUCCION.json')
set_cell(tg, 2, 2, 'workflows/Flujo 2 — Chatbot Omnicanal IA PRODUCCION.json')

# ===== Listado de Tablas: altas de la 2.1 y la I.1 =====
t1 = d.tables[1]
for ri, r in enumerate(t1.rows):
    if r.cells[0].text.strip() == 'Tabla 3.1':
        clonar(t1, ri - 1)
        set_cell(t1, ri, 0, 'Tabla 2.1')
        set_cell(t1, ri, 1, 'Posición de este trabajo respecto de los antecedentes revisados')
        set_cell(t1, ri, 2, '2.5.4')
        break
clonar(t1, len(t1.rows) - 1)
u = len(t1.rows) - 1
set_cell(t1, u, 0, 'Tabla I.1')
set_cell(t1, u, 1, 'Tiempos cronometrados del procesamiento manual, por orden y por fase')
set_cell(t1, u, 2, 'Anexo I')

# ===== #15 — OE3 en §1.5.2 =====
set_par(95,
    'Diseñar un esquema de base de datos en PostgreSQL con las tablas, vistas de métricas e '
    'índices necesarios para instrumentar el ciclo post-venta y sostener el cálculo automático '
    'de MTTD, MTTR y TMR.')

# ===== #16 — el canal Telegram cubre latencia, no exactitud =====
i = idx('El TMR global del canal Telegram')
replace_in_paragraph(d.paragraphs[i],
    'lo que confirma que H2a se sostiene también sobre un canal en producción.',
    'lo que confirma que H2a se sostiene también sobre un canal en producción. Corresponde '
    'precisar el alcance de esta validación: las 45 interacciones de Telegram son adicionales al '
    'corpus de 150 mensajes y no un subconjunto suyo, y no cuentan con etiquetas de referencia. '
    'En consecuencia, la validación del segundo canal cubre la latencia de respuesta pero no la '
    'exactitud de clasificación, que se reporta únicamente sobre el corpus etiquetado.')

d.save(RUTA)

# ===== §7.1 y §7.2 — las líneas que el diagnóstico exigía =====
d = Document(RUTA)
i = idx('Extensión del pipeline a envíos y logística')
insertar_bloque(d, i, [
    ('Medición del baseline manual con un diseño de mayor alcance: replicar el protocolo de la '
     'Sección 3.5.5 sobre varios operadores independientes y sin entrenamiento previo, lo que '
     'habilitaría un contraste inferencial contra grupo de control y convertiría el factor de '
     'mejora en un resultado con validez estadística plena.', 'Compact'),
    ('Implementación efectiva de la omnicanalidad: incorporar un identificador de cliente '
     'unificado y un identificador de conversación a la tabla interactions, y agregar al Flujo 2 '
     'un nodo de recuperación del historial previo, de modo que el contexto se preserve entre '
     'canales según la definición de la Sección 2.3.1.', 'Compact'),
    ('Cumplimiento del régimen de protección de datos personales: implementar el registro del '
     'consentimiento del titular, la política de retención y supresión, la minimización de los '
     'datos remitidos al proveedor de inferencia y la resolución de la base legal de la '
     'transferencia internacional, conforme la Ley 25.326 analizada en la Sección 2.6.', 'Compact'),
    ('Incorporación de ejemplos etiquetados al prompt y cableado del contexto de FAQ que el flujo '
     'ya construye, para cuantificar la mejora sobre el desempeño zero-shot medido en este '
     'trabajo.', 'Compact'),
])

# se agrega al final de las recomendaciones de producción, justo antes del título 7.2
i72 = idx('7.2 Líneas futuras')
insertar_bloque(d, i72 - 1, [
    ('Control de admisión y encolado de solicitudes: la prueba de concurrencia de la Sección 5.1 '
     'mostró que, si bien el pipeline no pierde integridad transaccional bajo carga simultánea, '
     'el 40,8 % de las órdenes queda sin procesar. Un despliegue productivo debería interponer '
     'una cola de mensajes entre el webhook y el pipeline, de modo que las solicitudes se '
     'encolen en lugar de descartarse, e incorporar métricas de profundidad de cola y de '
     'reintentos.', 'Compact'),
])

d.save(RUTA)

d3 = Document(RUTA)
print('=== Tabla 3.2 ===')
for r in d3.tables[5].rows:
    print('  ' + ' | '.join(c.text.strip()[:52] for c in r.cells))
print()
print('=== Tabla 4.3 ===')
for r in d3.tables[9].rows:
    print('  ' + ' | '.join(c.text.strip()[:52] for c in r.cells))
print()
print('=== 7.1 final y 7.2 ===')
j = None
for i, p in enumerate(d3.paragraphs):
    if p.text.strip().startswith('7.2 Líneas futuras'):
        j = i
        break
for k in range(j - 2, j + 10):
    print('%4d [%-10s] %s' % (k, d3.paragraphs[k].style.name, d3.paragraphs[k].text.strip()[:135]))
