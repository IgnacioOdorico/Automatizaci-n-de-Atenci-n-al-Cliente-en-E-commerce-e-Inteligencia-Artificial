# -*- coding: utf-8 -*-
"""Bloque 1 — los cuatro hallazgos Bloqueantes del dictamen."""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)
P = d.paragraphs


def set_par(i, texto, estilo=None):
    p = P[i]
    if p.runs:
        p.runs[0].text = texto
        for r in p.runs[1:]:
            r.text = ''
    else:
        p.add_run(texto)
    if estilo:
        try:
            p.style = estilo
        except KeyError:
            pass


# ===================== #4 — Hipótesis comparativas =====================
set_par(85,
    'H1: El pipeline automatizado de procesamiento de órdenes reduce significativamente '
    'el tiempo end-to-end del ciclo —desde la recepción del pedido hasta la notificación '
    'al cliente— respecto del tiempo que insume procesar la misma tarea de forma manual, '
    'medido como baseline propio según el protocolo de la Sección 3.5.5. Como criterio '
    'operativo secundario se fija que dicho tiempo se mantenga por debajo de 30 segundos.')

set_par(86,
    'H2a: El chatbot basado en GPT-4o-mini responde al cliente en un tiempo medio de '
    'respuesta (TMR) inferior a 10 segundos para consultas de tipo FAQ, estado de pedido '
    'y consultas generales.', 'Body Text')

set_par(87,
    'H2b: El chatbot clasifica correctamente la intención del mensaje en al menos el 85 % '
    'de los casos. Este umbral se fija como criterio del equipo —no deriva de un estándar '
    'publicado ni de un requerimiento de negocio externo— y se justifica en la Sección 2.2.3.')

set_par(88,
    'Las tres hipótesis se contrastan en el Capítulo 5 a partir de los datos recolectados '
    'durante las pruebas funcionales, la corrida de carga secuencial y la prueba de '
    'concurrencia. Corresponde señalar de antemano una precisión epistemológica: los '
    'umbrales de 30 y de 10 segundos operan como criterios de aceptación fijados por el '
    'equipo y no constituyen, por sí mismos, afirmaciones falsables. Dado el diseño del '
    'sistema —persistencia sobre una base de datos local y notificación a un capturador '
    'SMTP alojado en el mismo host— era previsible que ambos se cumplieran con amplio '
    'margen. La afirmación sustantiva y efectivamente refutable de H1 es la comparación '
    'contra el baseline de atención manual medido en este trabajo, y es en esos términos '
    'que se la somete a prueba en la Sección 5.3.')

# ===================== #2 — El residuo de "9.48 segundos" =====================
set_par(295,
    'El tiempo total end-to-end medido fue de 0,063 s en promedio (desvío estándar 0,005 s; '
    'mediana 0,062 s; mínimo 0,057 s; máximo 0,095 s; n = 50), lo que significa que desde la '
    'recepción de una orden hasta la escritura de la marca de notificación al cliente '
    'transcurren menos de siete centésimas de segundo, incluyendo la verificación de stock, '
    'la actualización del inventario y el envío del correo al servidor SMTP. El valor máximo '
    'observado (0,095 s) corresponde a la primera orden de la corrida y se atribuye al arranque '
    'en frío del motor de flujos; se reporta sin excluirlo de la muestra.')

d.save(RUTA)

# ===================== Verificación =====================
d2 = Document(RUTA)
print('=== HIPÓTESIS ===')
for i in (84, 85, 86, 87, 88):
    print('P%-4d [%s]' % (i, d2.paragraphs[i].style.name))
    print('      ' + d2.paragraphs[i].text[:400])
print()
print('=== P295 ===')
print(d2.paragraphs[295].text)
print()
buscar(d2, r'9[.,]48', '9.48 residual')
buscar(d2, r'\bordenes\b|recepcion|notificacion|Seccion|hipotesis', 'sin tilde (deben ser 0)')
