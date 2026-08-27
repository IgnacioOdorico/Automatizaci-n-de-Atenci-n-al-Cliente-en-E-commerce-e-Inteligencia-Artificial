# -*- coding: utf-8 -*-
"""Verificación de los 26 hallazgos del dictamen contra el documento final."""
import sys
import re
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

d = Document('docs/TESIS_FINAL_UTN_v6.docx')
TXT = '\n'.join(p.text for p in d.paragraphs)
for t in d.tables:
    for r in t.rows:
        TXT += '\n' + ' | '.join(c.text for c in r.cells)

OK, FALLA = [], []


def chk(num, sev, desc, cond):
    (OK if cond else FALLA).append((num, sev, desc))


def hay(pat):
    return re.search(pat, TXT, re.I) is not None


def n(pat):
    return len(re.findall(pat, TXT, re.I))


# --- BLOQUEANTES ---
chk(1, 'BLOQ', 'Baseline: protocolo §3.5.5 + resultados §5.1.4 + Anexo I',
    hay(r'3\.5\.5 Medición del baseline') and hay(r'5\.1\.4 Baseline de atención manual')
    and hay(r'Anexo I: Baseline') and hay(r'49,13 s'))
chk(2, 'BLOQ', 'Residuo "9.48 segundos" eliminado', not hay(r'9[.,]48'))
chk(3, 'BLOQ', 'Carga reconciliada: E1.a (50) y E1.b (120) diferenciadas',
    hay(r'E1\.a') and hay(r'E1\.b') and hay(r'6 × 20 = 120') and hay(r'40,8 ?%'))
chk(4, 'BLOQ', 'H1 comparativa; H2a y H2b declaradas en §1.4.2',
    hay(r'H1: El pipeline automatizado.{0,200}respecto del tiempo')
    and hay(r'H2a: El chatbot') and hay(r'H2b: El chatbot'))

# --- MAYORES ---
chk(5, 'MAY', '§2.5 Estado del arte con procedimiento y tabla comparativa',
    hay(r'2\.5 Estado del arte') and hay(r'Procedimiento de búsqueda')
    and hay(r'Tabla 2\.1') and hay(r'Parikh') and hay(r'Ngai'))
chk(6, 'MAY', 'IC de Wilson corregido a [87,3; 95,9] en TODAS sus ocurrencias',
    n(r'87,3 ?%') >= 3 and not hay(r'88,2') and not hay(r'96,3 ?%'))
chk(7, 'MAY', '§5.4.1(d) reescrita: ya no niega el análisis realizado',
    hay(r'Ausencia de contraste inferencial') and not hay(r'No se calcularon intervalos de confianza'))
chk(8, 'MAY', 'Precisiones del texto alineadas con la Tabla 5.9',
    hay(r'FAQ \(97,8 ?%\)') and not hay(r'FAQ \(96,3 ?%\)'))
chk(9, 'MAY', 'Omnicanalidad reencuadrada como multicanal',
    hay(r'2\.3 Comunicación multicanal') and hay(r'El sistema desarrollado es multicanal')
    # se admite "Chatbot Omnicanal" SOLO como parte del nombre real del archivo versionado
    and n(r'Chatbot Omnicanal') == n(r'Chatbot Omnicanal IA PRODUCCION\.json'))
chk(10, 'MAY', 'Soberanía de datos acotada (no cubre la inferencia)',
    hay(r'no se extiende a la inferencia del Flujo 2'))
chk(11, 'MAY', 'Prompt, corpus y datos crudos anexados',
    hay(r'Anexo H: Prompt de sistema') and hay(r'Anexo J: Corpus de evaluación')
    and hay(r'Tabla I\.1') and hay(r'ORD-E4-003'))
chk(12, 'MAY', 'WhatsApp: se retira la afirmación de sandbox',
    hay(r'no se ejecutó contra la API real ni contra su entorno sandbox')
    and hay(r'canal simulado vía SMTP local'))
chk(13, 'MAY', '§2.6 Ley 25.326 con artículos y deber de informar',
    hay(r'2\.6 Tratamiento de datos personales') and hay(r'Ley 25\.326')
    and hay(r'artículo 12') and hay(r'artículo 5'))
chk(14, 'MAY', 'v_chatbot_corpus documentada en Tabla 4.3 y en el DDL',
    n(r'v_chatbot_corpus') >= 7 and hay(r'CREATE OR REPLACE VIEW v_chatbot_corpus'))

# --- MENORES ---
chk(15, 'MEN', 'Conteos reconciliados (7 tablas, 6 vistas, 13 paneles, 18 nodos, 150 msj)',
    hay(r'7 tablas') and hay(r'6 vistas') and hay(r'13 paneles')
    and hay(r'18 nodos') and not hay(r'107 mensajes') and not hay(r'~21 nodos'))
chk(16, 'MEN', 'Tabla 5.5: columna rotulada como distribución predicha',
    hay(r'Mensajes clasificados') and not hay(r'Mensajes enviados \| 45'))
chk(17, 'MEN', 'F1–F5 alineadas con PF-01–PF-05',
    hay(r'SKU inválido') and hay(r'Orden duplicada') and not hay(r'Payload inválido'))
chk(18, 'MEN', 'received/notified ya no se usan como estados',
    hay(r'ambas son diferencias entre marcas temporales y no')
    and not hay(r'la transición received → confirmed calcula'))
chk(19, 'MEN', 'Contradicción media/mediana resuelta en §3.6.4',
    not hay(r'corresponden a la mediana de múltiples ejecuciones'))
chk(20, 'MEN', 'Liu et al. ya no se invoca como fuente de resultados comparables',
    not hay(r'consistente con los resultados reportados por Liu'))
chk(21, 'MEN', 'Cita colgante a OpenAI Status eliminada', not hay(r'OpenAI Status'))
chk(22, 'MEN', 'Versiones fijas en Tabla 3.2; "v3.8" corregido',
    hay(r'n8nio/n8n:2\.12\.2') and hay(r'grafana/grafana:13\.0\.7')
    and not hay(r'\| latest') and not hay(r'Docker Compose \| v3\.8'))
chk(23, 'MEN', 'Códigos internos y "versión anterior" fuera del texto expuesto',
    not hay(r'versión anterior') and not hay(r'orders \(E1\.a\)') and not hay(r'Valor de E2'))
chk(24, 'MEN', 'Referencia al "log de eventos" inexistente corregida',
    hay(r'no incorpora una tabla de bitácora de eventos')
    and not hay(r'registrados en el log de eventos'))
chk(25, 'MEN', 'CACE: el 181 % se declara nominal',
    hay(r'facturación nominal del sector, medida en pesos corrientes'))
chk(26, 'MEN', 'Ortografía y Resumen contiguo',
    not hay(r'notificaciónes') and not hay(r'restricciónes')
    and hay(r'Estos valores permiten confirmar la hipótesis'))

# --- extra: hallazgo propio ---
chk('X', 'PROP', 'Afirmación falsa de few-shot corregida a zero-shot',
    hay(r'régimen zero-shot') and not hay(r'inyección dinámica de FAQ'))

print('=' * 78)
print(' VERIFICACIÓN DEL DICTAMEN — %d/%d' % (len(OK), len(OK) + len(FALLA)))
print('=' * 78)
for num, sev, desc in sorted(OK, key=lambda x: str(x[0]).zfill(2)):
    print('  OK    [%-4s] #%-3s %s' % (sev, num, desc))
if FALLA:
    print()
    for num, sev, desc in FALLA:
        print('  FALLA [%-4s] #%-3s %s' % (sev, num, desc))
print()
print('Documento: %d párrafos, %d tablas' % (len(d.paragraphs), len(d.tables)))
