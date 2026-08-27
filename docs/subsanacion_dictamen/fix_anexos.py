# -*- coding: utf-8 -*-
"""Anexos H, I y J — cierran el hallazgo #11 (prompt, corpus y datos crudos no anexados)
y la recomendación 6 (procedimiento de cálculo del IC de Wilson)."""
import sys
sys.path.insert(0, '.')
from docxkit import *
from docx import Document

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)

PROMPT = """Sos un asistente virtual de atención al cliente de "TechStore", un e-commerce
argentino de tecnología.

## IDENTIDAD
- Nombre: Asistente TechStore
- Tono: cálido, humano, profesional — como si fuera una persona real de atención
  al cliente
- Idioma: respondé SIEMPRE en el mismo idioma que el cliente. Si escribe en
  inglés, respondé en inglés. Si escribe en español, usá español argentino
  (voseás: vos, tenés, podés)
- NO sos un chatbot genérico — representás a TechStore con orgullo
- Usá el nombre del cliente si está disponible — hace que la conversación se
  sienta más personal

## TU TAREA
Analizá el mensaje del cliente, clasificá su intención, y generá una respuesta
genuinamente útil y humana.
Respondé ÚNICAMENTE con un JSON válido. Sin texto antes ni después del JSON.
Sin markdown.

## FORMATO DE RESPUESTA (JSON estricto)
{
  "intent": "FAQ | ESTADO_PEDIDO | RECLAMO | GENERAL",
  "order_id": null,
  "urgente": false,
  "respuesta": "texto cálido, claro y personalizado para el cliente"
}"""

WILSON = """Datos:      x = 139 aciertos,  n = 150 mensajes,  p̂ = x/n = 0,926667
Nivel:      95 %  ->  z = 1,959964

                p̂ + z²/(2n)                    z            /  p̂(1-p̂)     z²
  centro =  ---------------- ;   semiancho = --------- · \\/  --------- + -----
                 1 + z²/n                     1 + z²/n           n         4n²

  z² = 3,841459        z²/n = 0,025610        z²/(2n) = 0,012805

  centro    = (0,926667 + 0,012805) / 1,025610 = 0,916019
  semiancho = (1,959964 / 1,025610) · √(0,000453 + 0,0000427) = 0,042548

  IC 95 % de Wilson = [0,873471 ; 0,958567] = [87,3 % ; 95,9 %]

Comprobación con otros métodos, a título informativo:
  Wald  (normal, sin corrección) ....... [88,5 % ; 96,8 %]
  El límite inferior supera el umbral del 85 % en cualquiera de los métodos,
  por lo que la conclusión sobre H2b no depende de esta elección."""

BLOQUE = [
    ('Anexo H: Prompt de sistema del clasificador de intenciones', 'Heading 2'),
    ('Se transcribe a continuación el texto íntegro del prompt de sistema que recibe el modelo '
     'GPT-4o-mini en el nodo de inferencia del Flujo 2, tal como está configurado en el workflow '
     'versionado en el repositorio. Es el prompt con el que se obtuvo el accuracy del 92,7 % '
     'reportado en la Sección 5.2.2. Se publica porque de él depende íntegramente ese resultado: '
     'sin el prompt, la medición de precisión no es auditable ni reproducible.', 'First Paragraph'),
    (PROMPT, 'Source Code'),
    ('Obsérvese que el prompt no contiene ejemplos etiquetados de entrada y salida: el '
     'clasificador opera en régimen zero-shot. Tampoco recibe el contenido de la base de FAQ; el '
     'nodo que construye ese contexto existe en el flujo pero su salida no se referencia desde '
     'este prompt, extremo que se declara en la Sección 4.4.3.', 'Body Text'),

    ('Anexo I: Baseline de atención manual — tiempos cronometrados', 'Heading 2'),
    ('Se transcriben los tiempos crudos de las doce órdenes procesadas manualmente según el '
     'protocolo de la Sección 3.5.5, incluidas las dos descartadas por interrupción del operador, '
     'que se consignan con su marca de descarte y no se eliminan del registro. Los valores están '
     'en segundos y provienen sin edición del archivo exportado por el cronómetro. El resultado '
     'primario de la Sección 5.1.4 se calcula sobre las diez órdenes válidas.', 'First Paragraph'),
    ('Tabla I.1: Tiempos cronometrados del procesamiento manual, por orden y por fase.', 'Body Text'),
    ('El texto completo de las notificaciones redactadas por el operador durante la medición, así '
     'como el archivo original exportado por el cronómetro y el script de análisis, se encuentran '
     'versionados en el repositorio del trabajo bajo el directorio del experimento.', 'Body Text'),

    ('Anexo J: Corpus de evaluación y procedimiento de cálculo estadístico', 'Heading 2'),
    ('El corpus de 150 mensajes utilizado para evaluar la clasificación, sus etiquetas de '
     'referencia, las predicciones del modelo y las etiquetas del evaluador independiente se '
     'encuentran versionados en el repositorio del trabajo en formato separado por comas, junto '
     'con los tiempos de anotación registrados por el instrumento de etiquetado. El corpus fue '
     'incorporado al repositorio antes de producirse el etiquetado, de modo que el sello temporal '
     'de esa incorporación acredita el carácter ciego del procedimiento descripto en la Sección '
     '3.5.3.', 'First Paragraph'),
    ('Se detalla a continuación el procedimiento de cálculo del intervalo de confianza de Wilson '
     'reportado para el accuracy global, a fin de que el valor publicado sea verificable.',
     'Body Text'),
    (WILSON, 'Source Code'),
    ('El coeficiente κ de Cohen reportado en la Sección 5.2.2 se calculó sobre la submuestra de 50 '
     'mensajes etiquetada por el evaluador independiente, con un acuerdo observado Po = 0,940 (47 '
     'coincidencias sobre 50) y un acuerdo esperado por azar Pe = 0,259 derivado de las '
     'distribuciones marginales de ambos anotadores, de donde κ = (Po − Pe) / (1 − Pe) = 0,919. '
     'La interpretación se realiza según la escala de Landis y Koch (1977).', 'Body Text'),
]

creados = insertar_bloque(d, len(d.paragraphs) - 1, BLOQUE)
print('Anexos insertados: %d párrafos' % len(creados))

# --- Tabla I.1 con los tiempos crudos ---
FILAS = [
    ('#', 'Orden', 'Rama', 'T1', 'T2', 'T3', 'Total', 'Estado'),
    ('1', 'ORD-E4-001', '—', '—', '—', '—', '—', 'Descartada'),
    ('2', 'ORD-E4-002', 'sin stock', '8,033', '—', '—', '—', 'Descartada'),
    ('3', 'ORD-E4-003', 'con stock', '7,530', '8,598', '53,674', '69,801', 'Válida'),
    ('4', 'ORD-E4-004', 'con stock', '11,386', '9,089', '34,453', '54,928', 'Válida'),
    ('5', 'ORD-E4-005', 'sin stock', '6,708', '5,598', '30,449', '42,754', 'Válida'),
    ('6', 'ORD-E4-006', 'con stock', '8,236', '4,698', '36,577', '49,510', 'Válida'),
    ('7', 'ORD-E4-007', 'con stock', '5,506', '5,768', '38,081', '49,354', 'Válida'),
    ('8', 'ORD-E4-008', 'con stock', '7,958', '4,857', '31,787', '44,602', 'Válida'),
    ('9', 'ORD-E4-009', 'sin stock', '7,695', '4,448', '34,960', '47,104', 'Válida'),
    ('10', 'ORD-E4-010', 'con stock', '5,306', '4,447', '35,884', '45,638', 'Válida'),
    ('11', 'ORD-E4-011', 'con stock', '5,806', '5,087', '32,009', '42,902', 'Válida'),
    ('12', 'ORD-E4-012', 'sin stock', '4,365', '5,448', '34,863', '44,676', 'Válida'),
]
idx = None
for i, p in enumerate(d.paragraphs):
    if p.text.strip().startswith('Tabla I.1:'):
        idx = i
        break
t = d.add_table(rows=len(FILAS), cols=8)
try:
    t.style = d.tables[18].style
except Exception:
    t.style = 'Table Grid'
for ri, fila in enumerate(FILAS):
    for ci, v in enumerate(fila):
        set_cell(t, ri, ci, v)
d.paragraphs[idx]._element.addnext(t._element)

# --- el IC de Wilson vive ahora en el Anexo J ---
replace_everywhere(d, 'El procedimiento de cálculo se detalla en el Anexo H.',
                   'El procedimiento de cálculo se detalla en el Anexo J.')
# --- Anexo A: conteo de vistas ---
replace_everywhere(d, 'la base de datos incluye 5 vistas de métricas',
                   'la base de datos incluye 6 vistas de métricas')

d.save(RUTA)

d2 = Document(RUTA)
print()
for i in range(444, len(d2.paragraphs)):
    print('%4d [%-16s] %s' % (i, d2.paragraphs[i].style.name,
                              d2.paragraphs[i].text.strip()[:110].replace('\n', ' / ')))
print()
for ti, t in enumerate(d2.tables):
    if t.rows[0].cells[1].text.strip() == 'Orden':
        print('Tabla I.1 en índice %d, %d filas' % (ti, len(t.rows)))
        for r in t.rows[:4]:
            print('   ' + ' | '.join(c.text.strip() for c in r.cells))
