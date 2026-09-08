# -*- coding: utf-8 -*-
"""Auditoria de 2.a instancia — BLOQUE 2 del plan de correccion.

Cubre M-03, M-04, M-05, M-06 y B-05 a B-08.

Decisiones tomadas y por que:
  M-03  Se retira la columna "Estado" de las Tablas 4.9 y 4.10 (el capitulo de
        desarrollo describe el diseño de la prueba, no su ejecucion) y se abre
        una §5.2.1 con los resultados de C1 a C5. Las subsecciones existentes se
        renumeran 5.2.1 -> 5.2.2 y 5.2.2 -> 5.2.3, con sus 15 referencias
        cruzadas. Los resultados se reportan en prosa y no en tabla nueva: una
        tabla intercalada obligaria a renumerar las Tablas 5.5 a 5.10 y sus
        referencias, con un riesgo desproporcionado al beneficio.
  M-04  La diferencia 1,47 s vs 3,07 s se presenta como observacion y no como
        acotacion del componente de red: las dos series no son apareadas.
  B-05  Se aplica la convencion que el propio documento ya usa en las otras 31
        entradas: nombre de revista y volumen en cursiva; titulo en cursiva para
        libros, informes y documentacion; "arXiv" en cursiva para preprints.
        La entrada de la ley no lleva cursiva, que es lo correcto para material
        normativo.
  B-06  MTTR bajo concurrencia = 0,192 - 0,092 = 0,100 s. La identidad
        E2E = MTTD + MTTR vale fila por fila, de modo que tambien vale para las
        medias. El desvio no es derivable de las cifras publicadas y no se
        inventa.
"""
import sys
import os
import re
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docxkit import *
from docx import Document

if any(f.startswith('~$') for f in os.listdir('docs')):
    sys.exit('ERROR: Word tiene abierto un documento en docs/. Cerralo primero.')

RUTA = 'docs/TESIS_FINAL_UTN_v6.docx'
d = Document(RUTA)


def idx(pref):
    for i, p in enumerate(d.paragraphs):
        if p.text.strip().startswith(pref):
            return i
    raise KeyError(pref)


def par(pref):
    return d.paragraphs[idx(pref)]


def tabla_de(epigrafe):
    from docx.text.paragraph import Paragraph
    ti = 0
    marcado = False
    for ch in d.element.body.iterchildren():
        tag = ch.tag.split('}')[1]
        if tag == 'p':
            if Paragraph(ch, d).text.strip().startswith(epigrafe + ':'):
                marcado = True
        elif tag == 'tbl':
            if marcado:
                return d.tables[ti]
            ti += 1
    raise KeyError(epigrafe)


def borrar_columna(t, titulo):
    """Elimina la columna cuyo encabezado es `titulo`, con su gridCol."""
    from docx.oxml.ns import qn
    cols = [c.text.strip() for c in t.rows[0].cells]
    if titulo not in cols:
        return False
    j = cols.index(titulo)
    for r in t.rows:
        tcs = r._tr.findall(qn('w:tc'))
        if j < len(tcs):
            r._tr.remove(tcs[j])
    grid = t._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        gc = grid.findall(qn('w:gridCol'))
        if j < len(gc):
            grid.remove(gc[j])
    return True


print('=' * 76)
print(' M-03 — el capitulo de diseño reportaba resultados de prueba')
print('=' * 76)
for ep in ('Tabla 4.9', 'Tabla 4.10'):
    t = tabla_de(ep)
    ok = borrar_columna(t, 'Estado')
    print('  columna "Estado" retirada de %s: %s  (%d columnas)'
          % (ep, ok, len(t.rows[0].cells)))

# renumeracion: primero la de mayor numero, para no pisarse
n = replace_everywhere(d, '5.2.2', '5.2.3')
print('  5.2.2 -> 5.2.3 : %d referencias' % n)
n = replace_everywhere(d, '5.2.1', '5.2.2')
print('  5.2.1 -> 5.2.2 : %d referencias' % n)

p52 = par('5.2 Resultados del Flujo 2')
bloque = [
    ('5.2.1 Pruebas funcionales del chatbot', 'Heading 3'),
    ('Se ejecutaron cinco pruebas funcionales sobre el Flujo 2, con el mismo criterio con que se '
     'ejecutaron las del Flujo 1: se verificó la respuesta entregada por el canal, el estado de '
     'la base de datos y los artefactos generados. Su diseño se detalla en la Tabla 4.10; aquí se '
     'reporta su ejecución. Las cinco resultaron aprobadas. C1, sobre una consulta de tipo '
     'frecuente, fue clasificada como FAQ y respondida por el canal de origen. C2, sobre estado '
     'de pedido, fue clasificada como ESTADO_PEDIDO y disparó la consulta a la tabla de órdenes, '
     'devolviendo el estado del pedido referido en el mensaje. C3, sobre un reclamo, fue '
     'clasificada como RECLAMO y generó el ticket correspondiente. C4, sobre un reclamo redactado '
     'en términos de urgencia, fue clasificada como RECLAMO, generó ticket y quedó registrada con '
     'la marca de urgencia en alto. C5, sobre una consulta abierta, fue clasificada como GENERAL '
     'y respondida por el modelo sin escalar. En los cinco casos la interacción quedó registrada '
     'con sus dos marcas temporales, de modo que el TMR de cada una es reconstruible.',
     'First Paragraph'),
    ('Corresponde una precisión sobre el alcance de estas pruebas, para que no se las confunda '
     'con la evaluación del clasificador. Son pruebas funcionales de un caso por intención: '
     'verifican que el flujo se recorra completo y que cada rama se active, no que la '
     'clasificación sea correcta en general ni que el contenido de la respuesta sea adecuado. '
     'Lo primero se mide sobre el corpus de 150 mensajes en la Sección 5.2.3; lo segundo no se '
     'midió en este trabajo, según se declara en la Sección 3.2.', 'Body Text'),
]
ancla = p52
for texto, estilo in bloque:
    ancla = insert_paragraph_after(ancla, texto, estilo=estilo)
print('  §5.2.1 con los resultados de C1 a C5: insertada')

# OE5 ya no puede decir que las diez estan "documentadas" sin mas
t61 = [t for t in d.tables if t.rows[0].cells[0].text.strip() == 'Obj.'][0]
f5 = [i for i, r in enumerate(t61.rows) if r.cells[0].text.strip() == 'OE5'][0]
set_cell(t61, f5, 2,
         '10 pruebas funcionales (5 por flujo), con su diseño en las Tablas 4.9 y 4.10 y sus '
         'resultados en las Secciones 5.1.1 y 5.2.1; 1 corrida de carga secuencial de 50 órdenes '
         'y 1 prueba de concurrencia de 6 rondas × 20 solicitudes simultáneas, con datos crudos y '
         'manifiesto de ejecución versionados.')
print('  celda de OE5 actualizada')

print()
print('=' * 76)
print(' M-04 — la diferencia entre canales no es una acotacion del componente de red')
print('=' * 76)
n = replace_in_paragraph(par('Las pruebas se ejecutaron sobre una notebook'),
    'Esa latencia no se toma de una fuente externa sino que se estima a partir de las propias '
    'mediciones de este trabajo: el TMR observado se ubica entre 1,47 s sobre el canal con '
    'entrega local y 3,07 s sobre el canal Telegram real, y la diferencia entre ambos acota el '
    'componente atribuible a la red del canal.',
    'El TMR observado se ubica entre 1,47 s sobre el canal con entrega local y 3,07 s sobre el '
    'canal Telegram real. Conviene precisar qué se puede y qué no se puede concluir de esa '
    'diferencia de 1,6 s. No es una acotación del componente atribuible a la red del canal, '
    'porque las dos series no son apareadas: provienen de conjuntos distintos —150 mensajes del '
    'corpus etiquetado y 45 interacciones de Telegram declaradas como muestra adicional y no como '
    'subconjunto— que difieren en composición de intenciones, en longitud de los mensajes y en '
    'momento de ejecución, y los tres factores inciden sobre el tiempo de inferencia. La '
    'diferencia se reporta por lo tanto como observación entre dos condiciones de medición, no '
    'como descomposición del tiempo. Aislar el componente de red exigiría enviar un mismo '
    'subconjunto de mensajes por ambos canales, lo que queda planteado en el Capítulo 7.')
print('  §4.6 reformulada: %d' % n)

p_hom = par('Homogeneizar el filtro de procedencia')
insert_paragraph_after(p_hom,
    'Aislar el componente de red del tiempo de respuesta: enviar un mismo subconjunto de mensajes '
    'por el canal simulado y por Telegram, de modo que las dos series sean apareadas y la '
    'diferencia entre ambas sea atribuible al trayecto de entrega y no a la composición de la '
    'muestra (Sección 4.6).',
    estilo=p_hom.style.name)
print('  ítem de series apareadas agregado al Capítulo 7')

print()
print('=' * 76)
print(' M-05 — el panel por canal no muestra el segundo canal')
print('=' * 76)
n = replace_everywhere(d,
    'Figura 9: Dashboard “Chatbot Multicanal” en Grafana, con el TMR promedio, la precisión de '
    'clasificación y la distribución de intenciones sobre el corpus evaluado.',
    'Figura 9: Dashboard “Chatbot Multicanal” en Grafana, con el TMR promedio, la precisión de '
    'clasificación y la distribución de intenciones sobre el corpus evaluado. El panel de '
    'interacciones por día y canal se alimenta de la vista del corpus, que está acotada a la '
    'ventana temporal de esa corrida: por eso muestra únicamente las 150 interacciones del canal '
    'simulado y no las 45 de Telegram, que se ejecutaron fuera de esa ventana y se reportan en la '
    'Tabla 5.6. La serie de correo aparece en cero porque el canal está previsto en el esquema y '
    'no implementado.')
print('  epígrafe de la Figura 9: %d' % n)

print()
print('=' * 76)
print(' M-06 — las 150 llamadas se ejecutaron dentro de una ventana de una hora')
print('=' * 76)
p_c = par('(c) Tamaño del conjunto de prueba')
insert_paragraph_after(p_c,
    '(c-bis) Ventana temporal de la corrida del chatbot: las 150 llamadas del corpus se '
    'ejecutaron dentro de un intervalo de sesenta minutos de un mismo día, según acota la propia '
    'vista que aísla la población (Anexo A). La decisión es coherente con la mitigación declarada '
    'en la Sección 3.6.4 —acotar la ventana reduce la variabilidad del servicio externo dentro de '
    'la medición— pero tiene un costo que corresponde consignar: el desvío de ±0,25 s que '
    'acompaña al tiempo medio de respuesta describe la dispersión dentro de esa hora y no la '
    'variabilidad del proveedor de inferencia entre días y franjas horarias, que es justamente la '
    'fuente de incertidumbre que la amenaza a la validez identifica. El valor reportado debe '
    'leerse en consecuencia como el desempeño en una ventana favorable y no como el esperable a '
    'lo largo del día.',
    estilo=p_c.style.name)
print('  (c-bis) incorporada a §5.4.1')

print()
print('=' * 76)
print(' B-05 a B-08')
print('=' * 76)

# --- B-05: cursivas bibliograficas, con la convencion que el propio documento usa
CURSIVAS = {
    'Aguirre Mayorga':   'Cuadernos de Administración, 35',
    'Alderete, M. V., Jones': 'Redes. Revista de Estudios Sociales de la Ciencia y la Tecnología, 23',
    'Alderete, M. V., & Porris': 'Ciencias Administrativas',
    'Bravo Maruri':      'GADE: Revista Científica, 5',
    'Cámara Argentina':  'Estudio Anual de Comercio Electrónico 2024',
    'Fondevila-Gascón':  'Correspondencias & Análisis',
    'Jurafsky':          'Speech and language processing: An introduction to natural language '
                         'processing, computational linguistics, and speech recognition with '
                         'language models',
    'Luo, H.':           'arXiv',
    'Ngai':              'Electronic Commerce Research and Applications, 50',
    'Pachas-Santos':     'Computación y Sistemas, 27',
    'Parikh':            'Proceedings of the 61st Annual Meeting of the Association for '
                         'Computational Linguistics',
    'Pypłacz':           'Procedia Computer Science, 225',
    'Ramos De Santis':   'Retos. Revista de Ciencias de la Administración y Economía, 14',
    'Zhang, D., Pee':    'International Journal of Information Management, 57',
}


def italizar(p, trozo):
    """Parte los runs para dejar `trozo` en cursiva conservando el resto."""
    full = ''.join(r.text for r in p.runs)
    i = full.find(trozo)
    if i < 0:
        return False
    pre, med, pos = full[:i], trozo, full[i + len(trozo):]
    base = p.runs[0]
    for r in list(p.runs[1:]):
        r._element.getparent().remove(r._element)
    base.text = pre
    base.italic = False
    r2 = p.add_run(med)
    r2.italic = True
    r3 = p.add_run(pos)
    r3.italic = False
    for r in (r2, r3):
        r.font.name = base.font.name
        r.font.size = base.font.size
    return True


hechos = 0
for p in d.paragraphs:
    t = p.text.strip()
    for clave, trozo in CURSIVAS.items():
        if t.startswith(clave) and not any(r.italic for r in p.runs):
            if italizar(p, trozo):
                hechos += 1
            break
print('  B-05 entradas con cursiva aplicada: %d de %d' % (hechos, len(CURSIVAS)))

# --- B-06: el factor de degradacion y el MTTR bajo concurrencia
n = replace_everywhere(d,
    'las métricas se degradaron por un factor aproximado de diez (Sección 5.1.3)',
    'las métricas se degradaron: el tiempo de detección por un factor de diez '
    '(0,009 s → 0,092 s) y el extremo a extremo por un factor de tres '
    '(0,063 s → 0,192 s), según la Sección 5.1.3')
print('  B-06 factor de degradación precisado en OE1: %d' % n)

t53 = tabla_de('Tabla 5.3')
fila_e2e = [i for i, r in enumerate(t53.rows)
            if r.cells[0].text.strip().startswith('E1.b — End-to-end')][0]
import copy
from docx.oxml.ns import qn
tr = copy.deepcopy(t53.rows[fila_e2e]._tr)
t53.rows[fila_e2e]._tr.addprevious(tr)
j = [k for k, r in enumerate(t53.rows) if r._tr is tr][0]
set_cell(t53, j, 0, 'E1.b — MTTR medio bajo concurrencia')
set_cell(t53, j, 1, '0,100 s (derivado: end-to-end − MTTD; el desvío no es derivable de las '
                    'cifras publicadas)')
print('  B-06 MTTR bajo concurrencia agregado a la Tabla 5.3')

# --- B-07: el epigrafe de la Figura 8 debe decir de que corrida son los correos
n = replace_everywhere(d,
    'Figura 8: Bandeja de Mailpit con los correos generados por el Flujo 1 durante las pruebas:',
    'Figura 8: Bandeja de Mailpit con los correos generados por el Flujo 1 durante una corrida de '
    'verificación de cinco órdenes (ORD-FIG4-001 a 005), preparada para exhibir ambas ramas de '
    'notificación y distinta de la corrida de carga de cincuenta órdenes:')
print('  B-07 epígrafe de la Figura 8: %d' % n)

# --- B-08: dos residuos menores
n = replace_everywhere(d, 'un catálogo de 20 productos electrónicos distribuidos en 11 categorías',
                          'un catálogo de 20 productos de tecnología y mobiliario distribuidos en '
                          '11 categorías')
print('  B-08 catálogo de §3.2: %d' % n)
for t in d.tables:
    for i, r in enumerate(t.rows):
        if r.cells[0].text.strip() == '17' and r.cells[1].text.strip() == 'Enviar Respuesta':
            set_cell(t, i, 1, 'Enviar Respuesta (canal simulado)')
            print('  B-08 nodo 17 de la Tabla 4.7 renombrado')

d.save(RUTA)
print()
print('guardado.')

# ================================ verificacion ================================
d2 = Document(RUTA)
P2 = [p.text.strip() for p in d2.paragraphs]
TODO = '\n'.join(P2) + '\n' + '\n'.join(
    c.text for t in d2.tables for r in t.rows for c in r.cells)
HEAD = [(p.style.name, p.text.strip()) for p in d2.paragraphs if p.style.name.startswith('Heading')]
fallos = []


def check(rot, cond, det=''):
    print('  [%s] %s %s' % ('OK   ' if cond else 'FALLA', rot, det))
    if not cond:
        fallos.append(rot)


print()
print('=== control ===')
dis = [t for t in d2.tables if t.rows[0].cells[0].text.strip() == '#'
       and 'Resultado esperado' in ' | '.join(c.text for c in t.rows[0].cells)]
cols = [[c.text.strip() for c in t.rows[0].cells] for t in dis]
check('M-03 las tablas de diseño ya no llevan "Estado"',
      len(dis) == 2 and all('Estado' not in c for c in cols), str(cols))
check('M-03 §5.2.1 reporta las pruebas del chatbot',
      any(t.startswith('5.2.1 Pruebas funcionales del chatbot') for _, t in HEAD))
check('M-03 5.2.2 y 5.2.3 renumeradas',
      any(t.startswith('5.2.2 Tiempos de respuesta') for _, t in HEAD)
      and any(t.startswith('5.2.3 Precisión') for _, t in HEAD),
      str([t for s, t in HEAD if t.startswith('5.2')]))
check('M-03 OE5 remite a ambas secciones de resultados',
      'Secciones 5.1.1 y 5.2.1' in TODO)
check('M-04 la diferencia ya no se presenta como acotación',
      'acota el componente atribuible a la red del canal' not in TODO
      and 'las dos series no son apareadas' in TODO)
check('M-05 el epígrafe de la Figura 9 explica la ausencia de Telegram',
      'no las 45 de Telegram, que se ejecutaron fuera de esa ventana' in TODO)
check('M-06 §5.4.1 declara la ventana de una hora',
      '(c-bis) Ventana temporal' in TODO and 'sesenta minutos' in TODO)
i8 = next(k for k, t in enumerate(P2) if t.startswith('CAPÍTULO 8'))
j8 = next(k for k, t in enumerate(P2) if t.startswith('CAPÍTULO 9'))
sin = [P2[k][:40] for k in range(i8 + 1, j8)
       if P2[k] and not any(r.italic for r in d2.paragraphs[k].runs)]
check('B-05 solo la entrada normativa queda sin cursiva',
      len(sin) <= 1 and (not sin or sin[0].startswith('Ley 25.326')), str(sin))
check('B-06 el factor de degradación distingue las dos métricas',
      'el tiempo de detección por un factor de diez' in TODO
      and 'el extremo a extremo por un factor de tres' in TODO)
check('B-06 la Tabla 5.3 informa el MTTR bajo concurrencia',
      'E1.b — MTTR medio bajo concurrencia' in TODO)
check('B-07 el epígrafe de la Figura 8 declara la corrida de verificación',
      'ORD-FIG4-001 a 005' in TODO)
check('B-08 catálogo sin "20 productos electrónicos"',
      '20 productos electrónicos' not in TODO)
check('B-08 nodo 17 desambiguado', 'Enviar Respuesta (canal simulado)' in TODO)

print()
if fallos:
    print('FALLAS: %d' % len(fallos))
    for f in fallos:
        print('  - %s' % f)
    sys.exit(1)
print('BLOQUE 2 — SIN FALLAS')
