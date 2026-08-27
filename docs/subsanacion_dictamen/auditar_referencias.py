# -*- coding: utf-8 -*-
"""Auditoría de integridad referencial del documento.

Lo que un jurado hace al leer: sigue una referencia y verifica que llegue.
Esto comprueba que TODA referencia cruzada resuelva a algo que existe.
"""
import sys
import io
import re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document

d = Document('docs/TESIS_FINAL_UTN_v6.docx')
P = [p.text.strip() for p in d.paragraphs]
TXT_P = '\n'.join(P)
TXT_T = '\n'.join(' | '.join(c.text for c in r.cells) for t in d.tables for r in t.rows)
TODO = TXT_P + '\n' + TXT_T

problemas = []

# ---------- 1. epígrafes existentes ----------
cap_tablas = {}
for i, t in enumerate(P):
    m = re.match(r'^Tabla ([\dA-Z]+\.\d+):', t)
    if m:
        cap_tablas.setdefault(m.group(1), []).append(i)
cap_figs = {}
for i, t in enumerate(P):
    m = re.match(r'^Figura ([\dA-Z]+):', t)
    if m:
        cap_figs.setdefault(m.group(1), []).append(i)

print('epígrafes de tabla : %d' % len(cap_tablas))
print('epígrafes de figura: %d' % len(cap_figs))
print('objetos tabla      : %d' % len(d.tables))
print()

for k, v in cap_tablas.items():
    if len(v) > 1:
        problemas.append('Tabla %s tiene %d epígrafes duplicados' % (k, len(v)))
for k, v in cap_figs.items():
    if len(v) > 1:
        problemas.append('Figura %s tiene %d epígrafes duplicados' % (k, len(v)))

# ---------- 2. referencias en el texto ----------
refs_tabla = set(re.findall(r'Tabla ([\dA-Z]+\.\d+)', TODO))
huerfanas = sorted(refs_tabla - set(cap_tablas))
if huerfanas:
    problemas.append('Referencias a tablas SIN epígrafe: %s' % ', '.join(huerfanas))
sin_citar = sorted(set(cap_tablas) - refs_tabla)
if sin_citar:
    problemas.append('Tablas con epígrafe pero NUNCA citadas: %s' % ', '.join(sin_citar))

refs_fig = set(re.findall(r'[Ff]igura ([\dA]\d*)\b', TODO))
huer_f = sorted(refs_fig - set(cap_figs))
if huer_f:
    problemas.append('Referencias a figuras SIN epígrafe: %s' % ', '.join(huer_f))

# ---------- 3. secciones ----------
headings = set()
for p in d.paragraphs:
    if p.style.name.startswith('Heading'):
        m = re.match(r'^(\d+(?:\.\d+)*)', p.text.strip())
        if m:
            headings.add(m.group(1))
refs_sec = set(re.findall(r'(?:Secci[oó]n|§)\s*(\d+(?:\.\d+)+)', TODO))
sec_rotas = sorted(refs_sec - headings)
if sec_rotas:
    problemas.append('Referencias a secciones INEXISTENTES: %s' % ', '.join(sec_rotas))

# ---------- 4. anexos ----------
anexos = set(re.findall(r'^Anexo ([A-Z]):', '\n'.join(P), re.M))
refs_anexo = set(re.findall(r'Anexo ([A-Z])\b', TODO))
anx_rotos = sorted(refs_anexo - anexos)
if anx_rotos:
    problemas.append('Referencias a anexos INEXISTENTES: %s' % ', '.join(anx_rotos))

# ---------- 5. listados vs realidad ----------
def filas_listado(header0):
    for t in d.tables:
        if t.rows[0].cells[0].text.strip() == header0:
            return [r.cells[0].text.strip() for r in t.rows[1:]]
    return []

list_tab = [x.replace('Tabla ', '') for x in filas_listado('Tabla')]
list_fig = [x.replace('Figura ', '') for x in filas_listado('Figura')]
falta_en_listado = sorted(set(cap_tablas) - set(list_tab))
sobra_en_listado = sorted(set(list_tab) - set(cap_tablas))
if falta_en_listado:
    problemas.append('Tablas con epígrafe que FALTAN en el Listado: %s' % ', '.join(falta_en_listado))
if sobra_en_listado:
    problemas.append('Listado de Tablas menciona tablas INEXISTENTES: %s' % ', '.join(sobra_en_listado))
falta_f = sorted(set(cap_figs) - set(list_fig))
if falta_f:
    problemas.append('Figuras que FALTAN en el Listado: %s' % ', '.join(falta_f))

# ---------- 6. cifras clave: que no haya contradicciones ----------
CIFRAS = {
    'MTTD 0,009':      r'0,009',
    'MTTR 0,054':      r'0,054',
    'E2E 0,063':       r'0,063',
    'baseline 49,13':  r'49,13',
    'accuracy 92,7':   r'92,7',
    'TMR corpus 1,47': r'1,47',
    'TMR telegram 3,07': r'3,07',
    'kappa 0,919':     r'0,919',
}
print('=== PRESENCIA DE LAS CIFRAS CANÓNICAS ===')
for nom, pat in CIFRAS.items():
    print('  %-20s %d ocurrencias' % (nom, len(re.findall(pat, TODO))))

VIEJAS = {'1,79': r'\b1,79\b', '7,69': r'\b7,69\b', '2,38': r'\b2,38\b',
          '90,7': r'\b90,7\b', '190x': r'190 veces', '31,6x': r'31,6',
          '5 a 30 minutos': r'5 a 30 minutos', '107': r'\b107\b'}
print()
print('=== CIFRAS DE VERSIONES VIEJAS (deben ser 0) ===')
for nom, pat in VIEJAS.items():
    n = len(re.findall(pat, TODO))
    print('  %-16s %d %s' % (nom, n, '' if n == 0 else '  <-- REVISAR'))
    if n:
        problemas.append('Cifra vieja "%s" aparece %d vez/veces' % (nom, n))

# ---------- resultado ----------
print()
print('=' * 74)
if problemas:
    print(' PROBLEMAS DE INTEGRIDAD REFERENCIAL: %d' % len(problemas))
    print('=' * 74)
    for x in problemas:
        print('  - %s' % x)
else:
    print(' INTEGRIDAD REFERENCIAL: SIN PROBLEMAS')
    print('=' * 74)
